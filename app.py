import re
import os
import time
from datetime import datetime, date, timedelta
import streamlit as st
import pandas as pd
import qrcode
import cv2
import numpy as np
from PIL import Image
import textwrap
import json
import io
from fpdf import FPDF
import matplotlib.pyplot as plt
import base64
import numpy as np
import plotly.express as px
from ultralytics import YOLO
from reportlab.lib.pagesizes import letter
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
import csv
from github import Github, GithubException
import requests
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.image import MIMEImage
# Import per l'estrazione del testo dai file locali
from pypdf import PdfReader
from docx import Document

# ==================================================================
# --- 0. CARICARE MODELLO ADDESTRAMENTO CONTROLLO DPI---
# ==================================================================

@st.cache_resource
def carica_modello():
  # Carica i pesi del modello personalizzato salvati durante il training
  return YOLO("models/best.pt")


model = carica_modello()

# ==========================================
# 0.1. DEFINIZIONE DELLE MANSIONI E DEI DPI
# ==========================================
MANSIONI_DPI = {
    "Addetto inchiostri": ["otoprotector", "googles", "gloves", "boots"],
    "Addetto cliché e fustelle": ["boots", "gloves", "otoprotector", "googles"],
    "Addetto carrellisti magazzino bobine": ["boots", "gloves", "googles", "no-vest", "vest"],
}

# Gestione libreria PyGithub per il salvataggio remoto su repository GitHub
try:
    from github import Github
    GITHUB_AVAILABLE = True
except ImportError:
    GITHUB_AVAILABLE = False


def salva_file_su_github(path_repo, contenuto_bytes, messaggio_commit):
    """
    Salva o aggiorna un file direttamente sulla repository GitHub remota.
    Richiede st.secrets["GITHUB_TOKEN"] e st.secrets["REPO_NAME"] (o variabili d'ambiente).
    """
    token = st.secrets.get("GITHUB_TOKEN") or os.environ.get("GITHUB_TOKEN")
    repo_name = st.secrets.get("REPO_NAME") or os.environ.get("REPO_NAME")

    if not token or not repo_name:
        st.warning("⚠️ GITHUB_TOKEN o REPO_NAME non configurati nei Secrets/variabili d'ambiente. Salvataggio su GitHub non eseguito.")
        return False

    if not GITHUB_AVAILABLE:
        st.error("❌ Libreria 'PyGithub' non trovata. Aggiungi 'PyGithub' al file requirements.txt.")
        return False

    try:
        g = Github(token)
        repo = g.get_repo(repo_name)

        try:
            # Aggiornamento se il file esiste già su GitHub
            contents = repo.get_contents(path_repo)
            repo.update_file(
                path=path_repo,
                message=messaggio_commit,
                content=contenuto_bytes,
                sha=contents.sha
            )
        except Exception:
            # Creazione nuovo file su GitHub se non esiste
            repo.create_file(
                path=path_repo,
                message=messaggio_commit,
                content=contenuto_bytes
            )
        return True
    except Exception as e:
        st.error(f"Errore durante la connessione o scrittura su GitHub: {e}")
        return False

#================================================================
# --- UTILITY GITHUB ---
def save_to_github(file_path_in_repo, file_content, commit_message):
    """Salva o aggiorna un file (bytes o str) direttamente sul repository GitHub."""
    try:
        token = st.secrets.get("GITHUB_TOKEN")
        repo_name = st.secrets.get("REPO_NAME")
        if not token or not repo_name:
            st.error("Credenziali GitHub non configurate correttamente nei secrets.")
            return False
        
        g = Github(token)
        repo = g.get_repo(repo_name)
        
        try:
            contents = repo.get_contents(file_path_in_repo)
            repo.update_file(contents.path, commit_message, file_content, contents.sha)
        except GithubException as e:
            if e.status == 404:
                repo.create_file(file_path_in_repo, commit_message, file_content)
            else:
                raise e
        return True
    except Exception as e:
        st.error(f"Errore durante il salvataggio su GitHub: {e}")
        return False

# --- UTILITY PER DATEDATAFRAME ---
def load_events():
    """Carica gli eventi dai file CSV se presenti."""
    events = []
    for filename in ["segnalazioni_near_miss.csv", "analisi_near_miss.csv"]:
        if os.path.exists(filename):
            try:
                df = pd.read_csv(filename)
                events.extend(df.iloc[:, 0].dropna().astype(str).tolist())
            except Exception:
                pass
    return sorted(list(set(events))) if events else ["Nessun evento disponibile"]


# ==================================================================
# --- 1. CONFIGURAZIONI INIZIALI E DATABASE LOCALI ---
# ==================================================================
st.set_page_config(
    page_title="Piattaforma Centralizzata HSE System", 
    layout="wide", 
    initial_sidebar_state="collapsed"
)

FILE_NEAR_MISS = "near_miss.csv"
FILE_ANALISI_NM = "analisi_near_miss.csv"
FILE_SCADENZARIO = "scadenzario.xlsx"
FILE_CSV_MANUTENZIONE = "manutenzione.csv"
DIR_CONFORMITA = "documenti_conformita"
DIR_IMMAGINI_ANALISI = "Immagini_Analisi_Near_Miss"
DIR_ANALISI = os.path.dirname(os.path.abspath(__file__))

if not os.path.exists(DIR_CONFORMITA):
    os.makedirs(DIR_CONFORMITA)

if not os.path.exists(DIR_IMMAGINI_ANALISI):
    os.makedirs(DIR_IMMAGINI_ANALISI, exist_ok=True)


COLONNE_SCADENZARIO = [
    "Adempimento", 
    "Regolamento", 
    "Cosa produrre", 
    "Rilasciato", 
    "In vigore durante (mesi)", 
    "Scadenza", 
    "Responsabile", 
    "Note"
]
# Salva normativa di conformità legislativa - Aggiungi normativa
def salva_normativa(chiave, titolo, url, descrizione):
    db = {}
    if os.path.exists("normative.json"):
        try:
            with open("normative.json", "r") as f:
                db = json.load(f)
        except Exception:
            db = {}
    
    db[chiave.strip()] = {
        "titolo": titolo.strip(),
        "url": url.strip(),
        "desc": descrizione.strip()
    }
    
    with open("normative.json", "w") as f:
        json.dump(db, f, indent=4)


# ==================================================================
# --- FUNZIONE DI CALCOLO FORZATO SCADENZA ---
# ==================================================================
def applica_calcolo_scadenza_dataframe(df):
    """
    Scorre ogni singola riga del DataFrame e forza la sovrascrittura
    della colonna 'Scadenza' calcolandola da 'Rilasciato' e 'In vigore durante (mesi)'.
    """
    if df is None or df.empty:
        return df
        
    def _calcola_singola_riga(row):
        try:
            val_rilasciato = str(row.get("Rilasciato", "")).strip()
            val_mesi = row.get("In vigore durante (mesi)", 0)
            
            try:
                mesi = int(float(val_mesi))
            except (ValueError, TypeError):
                mesi = 0

            if not val_rilasciato or mesi <= 0:
                return ""

            # Parsing data con priorità al formato italiano
            dt = pd.to_datetime(val_rilasciato, dayfirst=True, errors='coerce')
            if pd.isna(dt):
                return ""

            # Calcolo avanzamento mesi e cambio anno
            tot_mesi = dt.month + mesi
            nuovo_anno = dt.year + (tot_mesi - 1) // 12
            nuovo_mese = (tot_mesi - 1) % 12 + 1

            import calendar
            max_giorni = calendar.monthrange(nuovo_anno, nuovo_mese)[1]
            nuovo_giorno = min(dt.day, max_giorni)

            return datetime(nuovo_anno, nuovo_mese, nuovo_giorno).strftime("%d/%m/%Y")
        except Exception:
            return ""

    # Sovrascrive direttamente la colonna Scadenza riga per riga
    df["Scadenza"] = df.apply(_calcola_singola_riga, axis=1)
    return df

def estrai_info_checklist(percorso_file):
    nome_base = os.path.basename(percorso_file)
    titolo_ricavato = os.path.splitext(nome_base)[0].replace("_", " ")
    testo = estrai_testo_da_file(percorso_file)
    linee = [l.strip() for l in testo.split("\n") if l.strip()]
    
    regolamento = "Non identificato (Consultare file)"
    documenti_produrre = "Vedere testo integrale"
    settore_aziendale = "Aziendale Generale"
    
    settori_mappa = {
        "produzione": "Produzione / Officina",
        "logistica": "Logistica / Magazzino",
        "uffici": "Uffici / Amministrazione",
        "manutenzione": "Manutenzione / Impianti",
        "cantiere": "Cantieri / Esterni",
        "laboratorio": "Laboratorio / Qualità",
        "commerciale": "Commerciale / Vendite"
    }
    
    testo_lower_completo = testo.lower()
    for chiave, nome_settore in settori_mappa.items():
        if chiave in testo_lower_completo:
            settore_aziendale = nome_settore
            break

    for linea in linee:
        l_low = linea.lower()
        if "d.lgs" in l_low or "decreto" in l_low or "legge" in l_low or "regolamento" in l_low:
            if len(linea) < 100:
                regolamento = linea
                break
                
    for linea in linee:
        l_low = linea.lower()
        if "allegato" in l_low or "produrre" in l_low or "documentazione" in l_low or "certificato" in l_low:
            if len(linea) < 120 and len(linea) > 20:
                documenti_produrre = linea
                break
                
    return titolo_ricavato, regolamento, documenti_produrre, settore_aziendale

def calcola_data_scadenza(row):
    scadenza_attuale = str(row.get("Scadenza", "")).strip()
    if scadenza_attuale and scadenza_attuale.lower() not in ["nan", "none", "", "<na>", "nat"]:
        return scadenza_attuale
        
    rilasciato = str(row.get("Rilasciato", "")).strip()
    mesi = str(row.get("In vigore durante (mesi)", "")).strip()
    
    if not rilasciato or rilasciato.lower() in ["nan", "none", ""] or not mesi or mesi.lower() in ["nan", "none", "", "0"]:
        return ""
        
    try:
        rilasciato_pulito = rilasciato.replace("-", "/").split()[0]
        data_rilascio = datetime.strptime(rilasciato_pulito, "%d/%m/%Y")
        
        mesi_da_aggiungere = int(float(mesi))
        nuovo_mese = data_rilascio.month + mesi_da_aggiungere
        nuovo_anno = data_rilascio.year + (nuovo_mese - 1) // 12
        nuovo_mese = (nuovo_mese - 1) % 12 + 1
        
        nuovo_giorno = min(data_rilascio.day, [31, 29 if nuovo_anno % 4 == 0 and (nuovo_anno % 100 != 0 or nuovo_anno % 400 == 0) else 28, 31, 30, 31, 30, 31, 31, 30, 31, 30, 31][nuovo_mese - 1])
        data_scadenza = datetime(nuovo_anno, nuovo_mese, nuovo_giorno)
        return data_scadenza.strftime("%d/%m/%Y")
    except:
        return ""

def evidenzia_righe_scadenza(row):
    val = row.get("Scadenza", "")
    if pd.isna(val) or not val or str(val).strip() in ["", "nan", "None", "<NA>"]:
        return [''] * len(row)
    
    anno_corrente = datetime.now().year
    anno_successivo = anno_corrente + 1
    
    try:
        parti = str(val).strip().replace("-", "/").split("/")
        if len(parti) == 3:
            anno_scadenza = int(parti[-1].split()[0])
            if anno_scadenza == anno_corrente:
                return ['background-color: #ffcccc; color: #cc0000; font-weight: bold;'] * len(row)
            elif anno_scadenza == anno_successivo:
                return ['background-color: #fff2cc; color: #b68500; font-weight: bold;'] * len(row)
    except:
        pass
    return [''] * len(row)

# ==================================================================
# --- FUNZIONE HELPER: RICERCA NEI DOCUMENTI (Fix NameError) CONSULTAZIONE
# ==================================================================
def cerca_nei_documenti(query, cartella_target):
    """
    Effettua una ricerca testuale basata su parole chiave nei file
    presenti nella cartella indicata. Supporta file .pdf, .docx, .txt, .csv.
    """
    query_pulita = query.strip().lower()
    if len(query_pulita) < 3:
        return "specifica"

    risultati = []
    
    if not os.path.exists(cartella_target):
        return risultati

    parole_chiave = query_pulita.split()

    for nome_file in os.listdir(cartella_target):
        percorso_completo = os.path.join(cartella_target, nome_file)
        if not os.path.isfile(percorso_completo):
            continue

        testo_estratto = ""
        ext = os.path.splitext(nome_file)[1].lower()

        try:
            # Lettura basata sul tipo di file
            if ext == ".txt":
                with open(percorso_completo, "r", encoding="utf-8", errors="ignore") as f:
                    testo_estratto = f.read()
            elif ext == ".csv":
                df_temp = pd.read_csv(percorso_completo, errors="ignore")
                testo_estratto = df_temp.to_string()
            elif ext == ".docx":
                import docx
                doc = docx.Document(percorso_completo)
                testo_estratto = "\n".join([p.text for p in doc.paragraphs])
            elif ext == ".pdf":
                try:
                    import pypdf
                    reader = pypdf.PdfReader(percorso_completo)
                    testo_estratto = "\n".join([page.extract_text() or "" for page in reader.pages])
                except Exception:
                    pass
        except Exception:
            continue

        if not testo_estratto.strip():
            continue

        testo_lower = testo_estratto.lower()
        corrispondenze = sum(1 for pk in parole_chiave if pk in testo_lower)

        if corrispondenze > 0:
            score = int((corrispondenze / len(parole_chiave)) * 100)
            risultati.append({
                "fonte": nome_file,
                "nome_file": nome_file,
                "percorso_completo": percorso_completo,
                "score": score,
                "testo": testo_estratto
            })

    # Ordina i risultati per rilevanza (score decrescente)
    risultati.sort(key=lambda x: x["score"], reverse=True)
    return risultati

# Funzione nativa per la generazione di QR Code (Ritorna un'immagine PIL)
def genera_qr_nativo(data):
    qr = qrcode.QRCode(version=1, box_size=10, border=4)
    qr.add_data(data)
    qr.make(fit=True)
    return qr.make_image(fill_color='black', back_color='white')
def decodifica_qr_opencv(image_file):
    # Converti l'immagine in un formato leggibile da OpenCV
    img = Image.open(image_file).convert('RGB')
    img_array = np.array(img)

    #Decodifica tramite OpenCV
    detector = cv2.QRCodeDetector()
    valore, punti, qr_code = detector.detectAndDecode(img_array)

    return valore if valore else None

# ==================================================================
# --- FUNZIONI HELPER: ESTRAZIONE TESTO E CHECKLIST (Fix NameError) ---
# ==================================================================
def estrai_testo_da_file(percorso_file):
    """
    Estrae il contenuto testuale da file .pdf, .docx, .txt, .csv, .xlsx.
    """
    if not os.path.exists(percorso_file):
        return ""

    ext = os.path.splitext(percorso_file)[1].lower()
    testo = ""

    try:
        if ext == ".txt":
            with open(percorso_file, "r", encoding="utf-8", errors="ignore") as f:
                testo = f.read()
        elif ext == ".csv":
            df = pd.read_csv(percorso_file, errors="ignore")
            testo = df.to_string()
        elif ext in [".xlsx", ".xls"]:
            df = pd.read_excel(percorso_file)
            testo = df.to_string()
        elif ext == ".docx":
            import docx
            doc = docx.Document(percorso_file)
            testo = "\n".join([p.text for p in doc.paragraphs])
        elif ext == ".pdf":
            import pypdf
            reader = pypdf.PdfReader(percorso_file)
            testo = "\n".join([page.extract_text() or "" for page in reader.pages])
    except Exception:
        testo = ""

    return testo


def estrai_info_checklist(percorso_file):
    """
    Analizza il file e ne estrae i metadati principali per la check-list.
    """
    nome_base = os.path.basename(percorso_file)
    testo = estrai_testo_da_file(percorso_file)
    
    # Titolo derivato dal nome del file
    titolo = os.path.splitext(nome_base)[0].replace("_", " ").title()
    
    # Valori di default
    regolamento = "D.Lgs. 81/08 / Normativa HSE"
    doc_produrre = "Verifica documentale e registro aggiornato"
    settore_aziendale = "Generale / Sicurezza sul Lavoro"

    testo_lower = testo.lower()

    # Logica di riconoscimento semplice basata sul contenuto
    if "antincendio" in testo_lower or "vigili del fuoco" in testo_lower:
        settore_aziendale = "Antincendio & Emergenze"
        regolamento = "D.M. 02/09/2021 / D.P.R. 151/11"
        doc_produrre = "CPI / Attestazione Rinnovo Periodico"
    elif "rifiuti" in testo_lower or "ambiente" in testo_lower or "scarichi" in testo_lower:
        settore_aziendale = "Ambiente & Ecologia"
        regolamento = "D.Lgs. 152/2006 (Testo Unico Ambiente)"
        doc_produrre = "MUD / Formulari FIR / Registro Carico Scarico"
    elif "medico" in testo_lower or "sorveglianza sanitaria" in testo_lower:
        settore_aziendale = "Medicina del Lavoro"
        regolamento = "D.Lgs. 81/08 Art. 41"
        doc_produrre = "Giudizi di Idoneità / Piano Sanitario"

    return titolo, regolamento, doc_produrre, settore_aziendale
# ==============================================================
# SALVA NORMATIVA DI CONFORMITA SU NORMATIVE.JSON
# =============================================================
def salva_normativa(chiave, titolo, url, descrizione):
    """
    Salva o aggiorna una norma nel file normative.json su GitHub e in locale.
    """
    if not chiave.strip():
        st.warning("Inserire almeno una parola chiave per la norma.")
        return False

    file_json_name = "normative.json"
    db_norme = {}

    # 1. Tenta il recupero del JSON corrente da GitHub
    if github_token and repo_name:
        url_github_api = f"https://api.github.com/repos/{repo_name}/contents/{file_json_name}"
        headers = {"Authorization": f"token {github_token}"}
        sha_file = None
        try:
            res = requests.get(url_github_api, headers=headers)
            if res.status_code == 200:
                data_json = res.json()
                sha_file = data_json.get("sha")
                content_b64 = data_json.get("content", "")
                import base64
                content_decoded = base64.b64decode(content_b64).decode("utf-8")
                db_norme = json.loads(content_decoded)
        except Exception:
            pass

    # Fallback locale se da GitHub non è stato letto nulla
    if not db_norme and os.path.exists(file_json_name):
        try:
            with open(file_json_name, "r", encoding="utf-8") as f:
                db_norme = json.load(f)
        except Exception:
            db_norme = {}

    # 2. Aggiorna il dizionario
    db_norme[chiave.strip()] = {
        "titolo": titolo.strip(),
        "url": url.strip(),
        "desc": descrizione.strip()
    }

    content_bytes = json.dumps(db_norme, indent=4, ensure_ascii=False).encode("utf-8")

    # 3. Salva in locale
    try:
        with open(file_json_name, "wb") as f_out:
            f_out.write(content_bytes)
    except Exception:
        pass

    # 4. Salva o aggiorna il file su GitHub
    if github_token and repo_name:
        try:
            import base64
            url_github_api = f"https://api.github.com/repos/{repo_name}/contents/{file_json_name}"
            headers = {"Authorization": f"token {github_token}"}
            
            # Recupera lo SHA se non già salvato
            if not sha_file:
                res_check = requests.get(url_github_api, headers=headers)
                if res_check.status_code == 200:
                    sha_file = res_check.json().get("sha")

            payload = {
                "message": f"Aggiornamento database normative.json - Norma: {chiave}",
                "content": base64.b64encode(content_bytes).decode("utf-8")
            }
            if sha_file:
                payload["sha"] = sha_file

            res_put = requests.put(url_github_api, headers=headers, json=payload)
            if res_put.status_code in [200, 201]:
                return True
        except Exception as e:
            st.error(f"Errore durante la sincronizzazione su GitHub: {e}")

    return True

# ==================================================================
# --- MOTORE ROBUSTO ANALISI DI CONFORMITÀ & GAP ANALYSIS ---
# ==================================================================

def estrai_testo_da_bytes_robusto(nome_file, file_bytes):
    """Estrae in modo sicuro il testo da file binari o di testo."""
    if not file_bytes:
        return ""
    ext = os.path.splitext(nome_file)[1].lower()
    testo = ""
    try:
        if ext == ".txt":
            testo = file_bytes.decode("utf-8", errors="ignore")
        elif ext == ".csv":
            df = pd.read_csv(io.BytesIO(file_bytes), errors="ignore")
            testo = df.to_string()
        elif ext in [".xlsx", ".xls"]:
            df = pd.read_excel(io.BytesIO(file_bytes))
            testo = df.to_string()
        elif ext == ".docx":
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            testo = "\n".join([p.text for p in doc.paragraphs])
        elif ext == ".pdf":
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            testo = "\n".join([page.extract_text() or "" for page in reader.pages if page.extract_text()])
    except Exception as e:
        st.warning(f"Impossibile leggere il contenuto del file {nome_file}: {e}")
        testo = ""
    return testo.strip()


def estrai_contenuto_link(url_link):
    """
    Effettua una chiamata HTTP GET al link ufficiale per estrarne il testo
    ed evitare errori di timeout o blocchi SSL.
    """
    if not url_link or not (url_link.startswith("http://") or url_link.startswith("https://")):
        return ""
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
        response = requests.get(url_link, headers=headers, timeout=5, verify=False)
        if response.status_code == 200:
            # Estrazione basilare di testo evitando tag html
            import re
            text_clean = re.sub(r'<[^>]+>', ' ', response.text)
            return text_clean[:3000] # Limita a 3000 caratteri per efficienza
    except Exception:
        pass
    return ""


def genera_piano_adeguamento_e_gap(chiave, info_norma, testo_link):
    """
    Identifica i documenti da produrre e le azioni da intraprendere
    in base alla norma specifica non conforme.
    """
    chiave_lower = chiave.lower()
    desc_norma = info_norma.get("desc", "").lower() + " " + testo_link.lower()

    # Mappatura intelligente basata sui requisiti di legge
    if any(k in chiave_lower or k in desc_norma for k in ["rifiuto", "rifiuti", "ambiente", "fir", "mud"]):
        doc_necessari = [
            "Formulari di Identificazione Rifiuti (FIR)",
            "Registro di Carico e Scarico Rifiuti",
            "Dichiarazione Annuale MUD (Modello Unico di Dichiarazione Ambientale)",
            "Contratti con Trasportatori e Smaltitori Autorizzati (Albo Gestori Ambientali)"
        ]
        azioni_da_fare = [
            "Verificare la classificazione dei codici EER (ex CER) prodotti in azienda.",
            "Predisporre un'area stoccaggio temporaneo norma di legge con segnaletica appropriata.",
            "Designare e formare il responsabile della gestione dei rifiuti aziendali."
        ]

    elif any(k in chiave_lower or k in desc_norma for k in ["antincendio", "vigili", "cpi", "emergenza"]):
        doc_necessari = [
            "Certificato di Prevenzione Incendi (CPI) / Segnalazione Certificata di Inizio Attività (SCIA Incendi)",
            "Piano di Emergenza ed Evacuazione aggiornato",
            "Registro dei Controlli Antincendio (manutenzione estintori, idranti, porte REI)",
            "Attestati di formazione per Addetti alla Gestione Emergenze Antincendio"
        ]
        azioni_da_fare = [
            "Programmare la manutenzione periodica semestrale dei presidi antincendio.",
            "Nomina e formazione teorico-pratica degli addetti antincendio aziendali.",
            "Eseguire e verbalizzare la prova periodica di evacuazione aziendale."
        ]

    elif any(k in chiave_lower or k in desc_norma for k in ["medico", "sanitaria", "sorveglianza", "idoneità"]):
        doc_necessari = [
            "Nomina formale del Medico Competente",
            "Piano di Sorveglianza Sanitaria aziendale",
            "Giudizi di Idoneità alla mansione rilasciati dal Medico Competente",
            "Verbale della riunione periodica di sicurezza (Art. 35 D.Lgs. 81/08)"
        ]
        azioni_da_fare = [
            "Sottoporre i lavoratori esposti a rischi specifici a visita medica preventiva e periodica.",
            "Inviare le relazioni sanitarie all'Inail se richiesto.",
            "Aggiornare le schede di rischio mansione."
        ]

    elif any(k in chiave_lower or k in desc_norma for k in ["formazione", "corso", "accordo stato regioni"]):
        doc_necessari = [
            "Attestati di formazione Generale e Specifica dei lavoratori",
            "Attestato formazione RSPP / RLS / Preposti / Dirigenti",
            "Registro delle presenze e verbali dei corsi svolti"
        ]
        azioni_da_fare = [
            "Verificare le scadenze dei quinquenni di aggiornamento formativo.",
            "Pianificare i corsi mancanti in base all'Accordo Stato-Regioni applicabile."
        ]

    else:
        # Piano di conformità generico derivato dalle informazioni della norma
        doc_necessari = [
            f"Documentazione tecnica / Procedura aziendale specifica per '{info_norma.get('titolo', chiave)}'",
            "Verbale di audit interno o scheda di conformità",
            "Registro di tracciabilità delle verifiche e controlli"
        ]
        azioni_da_fare = [
            f"Consultare il link ufficiale ({info_norma.get('url', 'N.D.')}) per verificare i requisiti specifici.",
            "Predisporre una procedura operativa interna di adeguamento.",
            "Incaricare un consulente/tecnico abilitato per la verifica dei requisiti normativi."
        ]

    return doc_necessari, azioni_da_fare

# ======================================================================================
# NOTIFICA EMAIL PER CONTROLLO DPI
# ======================================================================================
def invia_email_notifica_dpi(mansione, dpi_rilevati, mancanti, esito_conforme, img_bytes=None):
    """
    Invia una e-mail di notifica contenente il testo dell'esito
    e l'immagine con i rilievi YOLO allegata.
    """
    email_destinatario = st.secrets.get("EMAIL", "")
    if not email_destinatario:
        return False, "Nessun indirizzo 'EMAIL' configurato in st.secrets."

    # Configurazione server SMTP da secrets
    smtp_server = st.secrets.get("SMTP_SERVER", "smtp.gmail.com")
    smtp_port = st.secrets.get("SMTP_PORT", 587)
    smtp_user = st.secrets.get("SMTP_USER", email_destinatario)
    smtp_password = st.secrets.get("SMTP_PASSWORD", "")

    oggetto = f"[{'CONFORME' if esito_conforme else 'ALLERTA NON CONFORME'}] Verifica DPI - Mansione: {mansione}"
    
    testo_corpo = f"""
    Notifica di verifica controllo DPI tramite IA (YOLO)

    - Data e Ora: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}
    - Mansione analizzata: {mansione}
    - DPI Rilevati: {', '.join(dpi_rilevati) if dpi_rilevati else 'Nessuno'}
    - Stato Conformità: {'✅ CONFORME' if esito_conforme else '❌ NON CONFORME'}
    """
    
    if not esito_conforme:
        testo_corpo += f"\n- DPI Mancanti Obbligatori: {', '.join(mancanti)}\n"

    testo_corpo += "\nIn allegato trovi l'immagine analizzata dal modello YOLO."

    msg = MIMEMultipart()
    msg['From'] = smtp_user if smtp_user else email_destinatario
    msg['To'] = email_destinatario
    msg['Subject'] = oggetto
    msg.attach(MIMEText(testo_corpo, 'plain'))

    # Allegato dell'immagine elaborata da YOLO
    if img_bytes:
        try:
            img_mime = MIMEImage(img_bytes, name=f"yolo_dpi_{datetime.now().strftime('%Y%m%d_%H%M%S')}.jpg")
            msg.attach(img_mime)
        except Exception as e_img:
            st.warning(f"Impossibile allegare l'immagine alla mail: {e_img}")

    try:
        server = smtplib.SMTP(smtp_server, int(smtp_port))
        server.starttls()
        if smtp_password:
            server.login(smtp_user, smtp_password)
        server.sendmail(msg['From'], email_destinatario, msg.as_string())
        server.quit()
        return True, "Email inviata con successo con allegato"
    except Exception as e:
        return False, str(e)

# Aggiornamento Classificazione Riconoscimento
def get_riconoscimenti_data():
    """Recupera i dati dei riconoscimenti dallo stato di sessione o dal file CSV."""
    try:
        base_dir = os.path.dirname(os.path.abspath(__file__))
    except NameError:
        base_dir = os.getcwd()
        
    file_riconoscimenti_csv = os.path.join(base_dir, "riconoscimenti_punteggi.csv")
    
    if os.path.exists(file_riconoscimenti_csv):
        try:
            df = pd.read_csv(file_riconoscimenti_csv, sep=";")
            df["Punti Segnalazione (+50)"] = pd.to_numeric(df["Punti Segnalazione (+50)"], errors='coerce').fillna(0).astype(int)
            df["Punti Skill Matrix (+25)"] = pd.to_numeric(df["Punti Skill Matrix (+25)"], errors='coerce').fillna(0).astype(int)
            df["Punteggio Totale"] = df["Punti Segnalazione (+50)"] + df["Punti Skill Matrix (+25)"]
            return df.sort_values(by="Punteggio Totale", ascending=False).reset_index(drop=True)
        except Exception:
            pass
            
    return pd.DataFrame()
  
# ==================================================================
# --- 3. INTESTAZIONE DELLA PIATTAFORMA ---
# ==================================================================
st.title("Piattaforma HSE - Near Miss")
st.markdown("### Sistema di Gestione dei Near Miss e Monitoraggio Salute e Sicurezza sul Lavoro")
st.markdown("---")

# --- 1. Inizializzazione dello Stato (Inseriscilo dopo st.set_page_config) ---
if "active_tab" not in st.session_state:
    st.session_state.active_tab = "Home Dashboard"

# --- 2. Creazione della barra di navigazione ---
# Usiamo st.radio in orizzontale per simulare le tab
tab_list = [
    "Home Dashboard", 
    "Segnalazione Near Miss", 
    "Consultazione",
    "Segnalazione Manutenzione",
    "Analisi Segnalazioni Near Miss",
    "Skill Matrix",
    "Analisi - Fase 2",
    "Piano Miglioramento",
    "Stima Costo Economico",
    "Riconoscimento",
    "KPI",
    "Scadenzario Adempimenti",
    "Controllo DPI"
]

# Impostiamo la navigazione
nav = st.radio(
    "Navigazione", 
    options=tab_list, 
    index=tab_list.index(st.session_state.active_tab), 
    horizontal=True, 
    label_visibility="collapsed"
)

# Aggiorniamo lo stato
st.session_state.active_tab = nav
st.markdown("---") # Linea di separazione estetica

# ==================================================================
# --- SEZIONE 1: HOME ---
# ==================================================================
if nav == "Home Dashboard":
    st.header("Quadro Generale di Controllo HSE")
    st.markdown("Benvenuto nel menu principale. Qui trovi i dati riassuntivi estratti in tempo reale dai database.")
    st.markdown(" ")

    try:
        base_dir = os.path.dirname(os.path.abspath(__file__))
    except NameError:
        base_dir = os.getcwd()

    # ---------------------------------------------------------
    # 1. CALCOLO METRICHE DASHBOARD
    # ---------------------------------------------------------
    
    # A. Conteggio Segnalazioni Near Miss (segnalazioni_near_miss.csv + Segnalazione_NM_Manutenzione/manutenzione.csv)
    tot_nm = 0
    path_nm_main = os.path.join(base_dir, "segnalazioni_near_miss.csv")
    path_nm_manut = os.path.join(base_dir, "Segnalazione_NM_Manutenzione", "manutenzione.csv")
    
    if os.path.exists(path_nm_main):
        try:
            df_nm1 = pd.read_csv(path_nm_main, sep=None, engine='python')
            tot_nm += len(df_nm1)
        except Exception:
            pass
            
    if os.path.exists(path_nm_manut):
        try:
            df_nm2 = pd.read_csv(path_nm_manut, sep=None, engine='python')
            tot_nm += len(df_nm2)
        except Exception:
            pass

    # B. Conteggio Analisi fatte (analisi_near_miss.csv)
    tot_an = 0
    path_analisi = os.path.join(base_dir, "analisi_near_miss.csv")
    if os.path.exists(path_analisi):
        try:
            df_an = pd.read_csv(path_analisi, sep=None, engine='python')
            tot_an = len(df_an)
        except Exception:
            pass

    # C. Registro Scadenze (Fallback Scadenzario)
    tot_scad = 0
    file_scad = os.path.join(base_dir, "scadenzario.xlsx") if 'FILE_SCADENZARIO' not in globals() else FILE_SCADENZARIO
    if os.path.exists(file_scad):
        try:
            tot_scad = len(pd.read_excel(file_scad))
        except Exception:
            pass

    # D. Stima Costo Economico Complessivo (Lettura Riga 26 dai file .csv in Stima_Economica/Report)
    tot_stima_economica = 0.0
    dir_stima_report = os.path.join(base_dir, "Stima_Economica", "Report")
    if not os.path.exists(dir_stima_report):
        dir_stima_report = os.path.join(base_dir, "Stima_Economica", "Report_Stima_Economica")

    if os.path.exists(dir_stima_report):
        files_csv_stima = [f for f in os.listdir(dir_stima_report) if f.lower().endswith('.csv')]
        
        for file_stima in files_csv_stima:
            path_file_stima = os.path.join(dir_stima_report, file_stima)
            try:
                # Lettura file CSV
                df_rep = pd.read_csv(path_file_stima, sep=None, engine='python', header=None)
                
                # Verifica presenza riga 26 (indice 25)
                if len(df_rep) >= 26:
                    riga_26 = df_rep.iloc[25]
                    # Cerca valori numerici nella riga 26
                    for val in riga_26:
                        if pd.notna(val):
                            s_val = (
                                str(val)
                                .replace("€", "")
                                .replace(" ", "")
                                .replace(".", "")
                                .replace(",", ".")
                                .strip()
                            )
                            try:
                                v_num = float(s_val)
                                tot_stima_economica += v_num
                            except ValueError:
                                pass
            except Exception:
                pass

    # Rendering Metriche 4 Colonne
    colA, colB, colC, colD = st.columns(4)
    with colA:
        st.metric("Segnalazioni Near Miss Ricevute", tot_nm)
    with colB:
        st.metric("Analisi di Segnalazioni Near Miss Trattate", tot_an)
    with colC:
        st.metric("Adempimenti in Registro Scadenze", tot_scad)
    with colD:
        st.metric(
            label="Totale Stima Economica", 
            value=f"€ {tot_stima_economica:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        ) 

    # ---------------------------------------------------------
    # 2. CLASSIFICA RICONOSCIMENTI
    # ---------------------------------------------------------
    st.markdown("---")
    st.markdown("## 🏆 Classifica Riconoscimenti per segnalazioni Near Miss")
    
    path_riconoscimento = os.path.join(base_dir, "Riconoscimento", "Riconoscimento_Partecipazione_NM.csv")
    df_home_ric = pd.DataFrame()

    if os.path.exists(path_riconoscimento):
        try:
            df_home_ric = pd.read_csv(path_riconoscimento, sep=None, engine='python')
        except Exception:
            pass
            
    # Fallback su funzione helper se presente
    if df_home_ric.empty and 'get_riconoscimenti_data' in globals():
        try:
            df_home_ric = get_riconoscimenti_data()
        except Exception:
            pass

    if not df_home_ric.empty and len(df_home_ric) > 0:
        # Ordinamento preventivo se esiste una colonna punteggio
        col_punteggio = next((c for c in df_home_ric.columns if "totale" in c.lower() or "punti" in c.lower()), None)
        if col_punteggio:
            df_home_ric = df_home_ric.sort_values(by=col_punteggio, ascending=False).reset_index(drop=True)

        col_nom = next((c for c in df_home_ric.columns if "nominativo" in c.lower() or "dipendente" in c.lower() or "nome" in c.lower()), df_home_ric.columns[0])

        # 1. Podio / Primi 3 posti
        col_h1, col_h2, col_h3 = st.columns(3)
        
        with col_h1:
            if len(df_home_ric) > 0:
                pts_1 = df_home_ric.iloc[0][col_punteggio] if col_punteggio else ""
                st.metric("🥇 1° Posto", f"{df_home_ric.iloc[0][col_nom]}", f"{pts_1} Pts" if pts_1 != "" else "")
        with col_h2:
            if len(df_home_ric) > 1:
                pts_2 = df_home_ric.iloc[1][col_punteggio] if col_punteggio else ""
                st.metric("🥈 2° Posto", f"{df_home_ric.iloc[1][col_nom]}", f"{pts_2} Pts" if pts_2 != "" else "")
        with col_h3:
            if len(df_home_ric) > 2:
                pts_3 = df_home_ric.iloc[2][col_punteggio] if col_punteggio else ""
                st.metric("🥉 3° Posto", f"{df_home_ric.iloc[2][col_nom]}", f"{pts_3} Pts" if pts_3 != "" else "")
    
        st.markdown("---")
        st.subheader("Tabella di Classificazione Completa")
    
        # Prepara il DataFrame aggiungendo la posizione in classifica
        df_classifica_completa = df_home_ric.copy()
        df_classifica_completa.insert(0, "Posizione", range(1, len(df_classifica_completa) + 1))
    
        # 2. Visualizzazione Tabella Intera
        st.dataframe(
            df_classifica_completa,
            use_container_width=True,
            hide_index=True
        )
    else:
        st.info("Nessun dato di riconoscimento disponibile nel file 'Riconoscimento/Riconoscimento_Partecipazione_NM.csv'.")
            
    # ---------------------------------------------------------
    # 3. VISUALIZZAZIONE FLOWCHART SEGNALAZIONE NEAR MISS (.PDF)
    # ---------------------------------------------------------
    st.markdown("---")
    st.subheader("Flowchart Segnalazione Near Miss")
    
    file_pdf_path = os.path.join(base_dir, "documenti_conformita", "FLOWCHART SEGNALAZIONE NEAR MISS.pdf")
    
    # Ricerca fallback in caso di difformità nel nome cartella
    if not os.path.exists(file_pdf_path):
        percorsi_alternativi = [
            os.path.join(base_dir, "documenti_conformita", "FLOWCHART SEGNALAZIONE NEAR MISS.PDF"),
            os.path.join(base_dir, "FLOWCHART SEGNALAZIONE NEAR MISS.pdf"),
            os.path.join("documenti_conformita", "FLOWCHART SEGNALAZIONE NEAR MISS.pdf")
        ]
        for p in percorsi_alternativi:
            if os.path.exists(p):
                file_pdf_path = p
                break

    if os.path.exists(file_pdf_path):
        try:
            with open(file_pdf_path, "rb") as f:
                pdf_bytes = f.read()
                
            st.download_button(
                label="📥 Scarica Flowchart Segnalazione Near Miss",
                data=pdf_bytes,
                file_name="FLOWCHART SEGNALAZIONE NEAR MISS.pdf",
                mime="application/pdf",
                use_container_width=True,
                key="btn_download_flowchart_pdf"
            )
                
            import base64
            base64_pdf = base64.b64encode(pdf_bytes).decode('utf-8')
            pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="700px" type="application/pdf"></iframe>'
            st.markdown(pdf_display, unsafe_allow_html=True)
        except Exception as e_pdf:
            st.error(f"Errore durante la lettura del file PDF: {e_pdf}")
    else:
        st.error("Il file PDF 'FLOWCHART SEGNALAZIONE NEAR MISS.pdf' non è stato trovato nella cartella 'documenti_conformita'.")

# ==================================================================
# --- SEZIONE 2: SEGNALAZIONE NEAR MISS ---
# ==================================================================
if nav == "Segnalazione Near Miss":
    st.info(
        "Near miss (mancato infortunio): evento avvenuto nel luogo di lavoro che non ha recato danno fisico al lavoratore, pur avendone il potenziale.\n"
        "Esempi: caduta di materiale imballato durante movimentazione con carrello elevatore; improvvisa fuoriuscita di liquido da tubazione; lavoratore scivola su pavimento bagnato senza riportare danni.\n\n"
        "Non conformità: situazione di pericolo che non genera alcun incidente/infortunio ma rilevabile su procedure operative, attrezzature, ambienti di lavoro, dpi.\n"
        "Esempi: macchinario senza protezione, casco di sicurezza non indossato, area di lavoro priva di percorsi sicuri."
    )

    st.subheader("MODULO S.NM.NC - Segnalazione Near Miss o Non Conformità")

    st.markdown("#### Inserisce immagine (Facoltativo)")
    opzione_immagine = st.radio(
        "Scegli la modalità di inserimento immagine:",
        [
            "Nessuna immagine",
            "Carica file",
            "Scatta foto col cellulare/webcam",
        ],
        key="scelta_media_reattiva",
    )

    immagine_salvata_nome = "Nessuna"
    if opzione_immagine == "Carica file":
        file_img = st.file_uploader(
            "Scegli un file immagine",
            type=["png", "jpg", "jpeg"],
            key="uploader_reattivo",
        )
        if file_img:
            immagine_salvata_nome = file_img.name
    elif opzione_immagine == "Scatta foto col cellulare/webcam":
        foto_scattata = st.camera_input(
            "Scatta una foto della criticità", key="camera_reattiva"
        )
        if foto_scattata:
            immagine_salvata_nome = (
                f"scatto_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
            )

    st.markdown("---")

    with st.form("form_segnalazione_near_miss", clear_on_submit=True):
        col_tipo, col_segnalatore = st.columns(2)
        with col_tipo:
            tipo_evento = st.radio(
                "Tipo evento", ["Near Miss", "Non Conformità"]
            )
        with col_segnalatore:
            segnalatore = st.text_input(
                "Segnalatore (inserire mansione o nome cognome)"
            )

        col_sesso, col_eta, col_data = st.columns(3)
        with col_sesso:
            sesso = st.radio("Sesso", ["Maschio", "Femmina"])
        with col_eta:
            fascia_eta = st.radio(
                "Fascia di Età",
                ["<18 anni", "18-30 anni", "31-50 anni", "51-67 anni"],
            )
        with col_data:
            data_evento = st.date_input("Data (formato gg/mm/aaaa)")

        col_luogo, col_reparto = st.columns(2)
        with col_luogo:
            luogo = st.radio(
                "Luogo", ["In Azienda", "In itinere", "In missione"]
            )
        with col_reparto:
            reparto_aziendale = st.text_input("Reparto (se è In Azienda)")

        col_f1, col_f2 = st.columns(2)
        with col_f1:
            fascia_oraria = st.radio(
                "Fascia oraria di accadimento",
                ["0-6", "6-12", "12-18", "18-24"],
            )
        with col_f2:
            fascia_lavoratore = st.text_input(
                "Fascia oraria per il lavoratore (1, 2, 3 ora Max. 8 ore)"
            )

        descrizione = st.text_area(
            "Descrizione dell'evento o della criticità (campo a testo libero)"
        )

        st.markdown(
            "#### Possibili cause dell'evento / In caso di Non Conformità selezionare la tipologia"
        )

        col_c1, col_c2 = st.columns(2)
        with col_c1:
            c_err_proc = st.checkbox(
                "Errore procedurale (disattenzione, scarsa conoscenza procedure operative, ...)"
            )
            c_prob_comm = st.checkbox(
                "Problema di comunicazione (lingua, incertezza nei ruoli e/o compiti, ...)"
            )
            c_manc_proc = st.checkbox(
                "Mancanza/inadeguatezza di procedure operative"
            )
            c_manc_prot = st.checkbox(
                "Mancanza di protezioni sull'attrezzatura"
            )
            c_car_prot = st.checkbox(
                "Carenza (inadeguatezza) di protezioni sull'attrezzatura"
            )
            c_anom_guasto = st.checkbox(
                "Anomalia/guasto in avviamento/arresto/esercizio (funzionamento)"
            )
            c_unica_attrez = st.checkbox(
                "Unica attrezzatura disponibile ma non idonea alla lavorazione"
            )
            c_ass_attrez = st.checkbox(
                "Assenza di attrezzature idonee alla lavorazione"
            )
            c_stocc_err = st.checkbox(
                "Stoccaggio/etichettatura errato di materiali"
            )
            c_prob_mat = st.checkbox(
                "Problema legato alle caratteristiche/trasformazioni di materiali"
            )
            c_segnal_inad = st.checkbox(
                "Segnaletica di sicurezza/Cartellonistica inadeguata o assente"
            )
            c_ass_perc = st.checkbox(
                "Assenza o inadeguatezza di percorsi in sicurezza, vie di transito, uscite di emergenza"
            )

        with col_c2:
            c_illum_inad = st.checkbox("Illuminazione non idonea o assente")
            c_ass_barr = st.checkbox(
                "Assenza o inadeguatezza di barriere, protezioni, parapetti, armature"
            )
            c_spazi_inad = st.checkbox("Spazi inadeguati su postazioni di lavoro")
            c_ass_stocc = st.checkbox(
                "Assenza o inadeguatezza di aree di stoccaggio"
            )
            c_pres_liq = st.checkbox(
                "Presenza imprevista di liquidi (acqua, olio, ...)"
            )
            c_pres_gas = st.checkbox("Presenza imprevista di gas, vapori")
            c_crit_imp = st.checkbox(
                "Criticità su impianti generali a supporto dell'area di lavoro"
            )
            c_pres_elett = st.checkbox(
                "Presenza di elettricità/linea elettrica accessibile"
            )
            c_rumore = st.checkbox("Livelli di rumorosità inadeguati")
            c_manc_dpi = st.checkbox("Mancato uso o uso errato di DPI")
            c_dpi_non_forn = st.checkbox("DPI non fornito")
            c_dpi_inad = st.checkbox("DPI inadeguato")

        altro_specificare = st.text_input(
            "Altro (specificare campo a testo libero)"
        )
        storico_riscontro = st.radio(
            "In base alla tua esperienza lavorativa, la situazione rilevata o osservata si è già presentata in passato anche recente?",
            ["Sì frequentemente", "Sì raramente", "No"],
        )
        valutazioni_proposte = st.text_area(
            "Valutazioni / azioni / proposte di miglioramento (campo a testo libero)"
        )

        submit_modulo = st.form_submit_button(
            "Registra ed Invia Segnalazione", use_container_width=True
        )

        if submit_modulo:
            if not descrizione.strip():
                st.error(
                    "Errore: La descrizione dell'evento è obbligatoria per effettuare il salvataggio."
                )
            else:
                cause_selezionate = []
                mappa_cause = {
                    "Errore procedurale": c_err_proc,
                    "Problema comunicazione": c_prob_comm,
                    "Mancanza procedure": c_manc_proc,
                    "Mancanza protezioni": c_manc_prot,
                    "Carenza protezioni": c_car_prot,
                    "Anomalia guasto": c_anom_guasto,
                    "Unica attrezzatura non idonea": c_unica_attrez,
                    "Assenza attrezzature idonee": c_ass_attrez,
                    "Stoccaggio errato": c_stocc_err,
                    "Problema materiali": c_prob_mat,
                    "Segnaletica inadeguata": c_segnal_inad,
                    "Inadeguatezza percorsi": c_ass_perc,
                    "Illuminazione inadeguata": c_illum_inad,
                    "Assenza barriere": c_ass_barr,
                    "Spazi inadeguati": c_spazi_inad,
                    "Assenza aree stoccaggio": c_ass_stocc,
                    "Presenza liquidi": c_pres_liq,
                    "Presenza gas": c_pres_gas,
                    "Criticità impianti": c_crit_imp,
                    "Presenza elettricità": c_pres_elett,
                    "Rumorosità": c_rumore,
                    "Mancato uso DPI": c_manc_dpi,
                    "DPI non fornito": c_dpi_non_forn,
                    "DPI inadeguato": c_dpi_inad,
                }
                for nome_c, var_c in mappa_cause.items():
                    if var_c:
                        cause_selezionate.append(nome_c)
                if altro_specificare.strip():
                    cause_selezionate.append(
                        f"Altro: {altro_specificare.strip()}"
                    )

                now_str = datetime.now().strftime("%d-%m-%Y %H:%M:%S")

                nuovo_record = {
                    "Data Segnalazione": now_str,
                    "Tipo Evento": tipo_evento,
                    "Segnalatore": (
                        segnalatore.strip() if segnalatore.strip() else "Anonimo"
                    ),
                    "Sesso": sesso,
                    "Fascia Eta": fascia_eta,
                    "Data Evento Real": data_evento.strftime("%d/%m/%Y"),
                    "Luogo": luogo,
                    "Reparto": (
                        reparto_aziendale.strip()
                        if reparto_aziendale.strip()
                        else "N/D"
                    ),
                    "Fascia Oraria Accadimento": fascia_oraria,
                    "Fascia Oraria Lavoratore": (
                        fascia_lavoratore.strip()
                        if fascia_lavoratore.strip()
                        else "N/D"
                    ),
                    "Descrizione": descrizione.strip(),
                    "Immagine Allegata": immagine_salvata_nome,
                    "Cause Rilevate": ", ".join(cause_selezionate),
                    "Verificato in Passato": storico_riscontro,
                    "Proposte Miglioramento": valutazioni_proposte.strip(),
                    "Firma Presa in Carico": "Da firmare",
                }

                # ---------------------------------------------------------
                # SALVATAGGIO AUTOMATICO SU GITHUB (segnalazioni_near_miss.csv)
                # ---------------------------------------------------------
                if submit_button:
                    # 1. Unisci il nuovo record con i dati esistenti
                    df_totale = pd.concat([df_analisi, df_n], ignore_index=True)
                
                    # 2. Invia l'aggiornamento a GitHub
                    if salva_csv_su_github(
                        df_totale,
                        FILE_ANALISI_NM,
                        f"Aggiunta analisi del {datetime.now().strftime('%d/%m/%Y')}",
                    ):
                        st.success("Analisi salvata e sincronizzata con successo su GitHub!")
                        time.sleep(1)
                        st.rerun()

# ==================================================================
# --- SEZIONE 3: SCADENZARIO ADEMPIMENTI ---
# ==================================================================
if nav == "Scadenzario Adempimenti":
    st.header("Registro Scadenzario Adempimenti Aziendali")
    st.markdown("Sezione Protetta — Sincronizzata direttamente sul file Excel GitHub `scadenzario.xlsx`")
    
    if "autenticato_scadenze" not in st.session_state:
        st.session_state.autenticato_scadenze = False
    if "df_scadenzario_state" not in st.session_state:
        st.session_state.df_scadenzario_state = None
        
    correct_pwd_scad = st.secrets.get("PASSWORD_SEZIONE", "hse2026")

    def carica_da_github_scadenzario():
        file_path_repo = "scadenzario.xlsx"
        github_token = st.secrets.get("GITHUB_TOKEN", "")
        repo_name = st.secrets.get("REPO_NAME", "")
        
        df = None
        if github_token and repo_name:
            url = f"https://raw.githubusercontent.com/{repo_name}/main/{file_path_repo}"
            headers = {"Authorization": f"token {github_token}"}
            try:
                res = requests.get(url, headers=headers)
                if res.status_code == 200:
                    df = pd.read_excel(io.BytesIO(res.content))
            except Exception:
                pass
        
        if df is None:
            file_scad_target = FILE_SCADENZARIO if 'FILE_SCADENZARIO' in globals() else "scadenzario.xlsx"
            if os.path.exists(file_scad_target):
                try:
                    df = pd.read_excel(file_scad_target)
                except Exception:
                    df = pd.DataFrame(columns=COLONNE_SCADENZARIO)
            else:
                df = pd.DataFrame(columns=COLONNE_SCADENZARIO)

        for col in COLONNE_SCADENZARIO:
            if col not in df.columns:
                df[col] = ""
        df = df[COLONNE_SCADENZARIO].fillna("")
        
        # FORZA IL CALCOLO SU TUTTE LE RIGHE DEL FILE SCARICATO DA GITHUB
        df = applica_calcolo_scadenza_dataframe(df)
        return df

    if not st.session_state.autenticato_scadenze:
        pwd_scad = st.text_input("Inserisci la password di sblocco Scadenzario", type="password", key="pwd_scad_tab")
        if st.button("Convalida Password Scadenzario", use_container_width=True):
            if pwd_scad == correct_pwd_scad:
                st.session_state.autenticato_scadenze = True
                st.session_state.df_scadenzario_state = carica_da_github_scadenzario()
                st.rerun()
            else:
                st.error("Password non corretta. Accesso negato.")
                
    if st.session_state.autenticato_scadenze and st.session_state.df_scadenzario_state is not None:
        st.success("Accesso Concesso alle scadenze.")
        st.markdown("---")
        
        st.subheader("1. Modifica ed Inserimento Dati")
        
        # Garantiamo che i dati in ingresso nell'editor abbiano già le scadenze ricalcolate
        df_editor_input = applica_calcolo_scadenza_dataframe(st.session_state.df_scadenzario_state.copy())
        
        df_modificato = st.data_editor(
            df_editor_input,
            use_container_width=True,
            hide_index=True,
            num_rows="dynamic",
            key="editor_scadenzario_pure_data"
        )
        
        if st.button("Aggiorna Calcoli Automatici e Salva nel File Excel su GitHub", use_container_width=True):
            try:
                # 1. Copia dei dati dall'editor
                df_finale = df_modificato.copy().fillna("")
                
                # 2. RICALCOLO FORZATO E TASSATIVO SU OGNI RIGA PRIMA DI CREARE L'EXCEL
                df_finale = applica_calcolo_scadenza_dataframe(df_finale)
                
                # 3. Conversione numerica della colonna mesi per pulizia
                if "In vigore durante (mesi)" in df_finale.columns:
                    df_finale["In vigore durante (mesi)"] = pd.to_numeric(
                        df_finale["In vigore durante (mesi)"], errors='coerce'
                    ).fillna(0).astype(int)
                
                # 4. Generazione file Excel binario con le colonne e i calcoli aggiornati
                output_excel = io.BytesIO()
                with pd.ExcelWriter(output_excel, engine='openpyxl') as writer:
                    df_finale.to_excel(writer, index=False)
                excel_bytes = output_excel.getvalue()

                # 5. Salvataggio Locale Backup
                file_scad_target = FILE_SCADENZARIO if 'FILE_SCADENZARIO' in globals() else "scadenzario.xlsx"
                with open(file_scad_target, "wb") as f:
                    f.write(excel_bytes)

                # 6. Push diretto su GitHub
                salvato_gh = False
                if 'save_to_github' in globals():
                    salvato_gh = save_to_github("scadenzario.xlsx", excel_bytes, "Aggiornamento scadenzario.xlsx")
                
                # 7. Aggiornamento dello stato interno
                st.session_state.df_scadenzario_state = df_finale
                
                if salvato_gh:
                    st.success("Tutte le righe sono state ricalcolate e il file 'scadenzario.xlsx' è stato aggiornato su GitHub!")
                else:
                    st.success("Calcoli aggiornati e salvati in locale!")
                    
                time.sleep(0.5)
                st.rerun()
            except Exception as err:
                st.error(f"Errore durante il salvataggio: {err}")
                
        st.markdown("---")
        st.subheader("2. Registro Alert & Scadenze Effettive (Vista Finale)")
        
        # Ricalcolo di sicurezza anche per la vista finale
        df_vista_alert = applica_calcolo_scadenza_dataframe(st.session_state.df_scadenzario_state.copy())
        
        if not df_vista_alert.empty:
            if 'evidenzia_righe_scadenza' in globals():
                styler_colorato = df_vista_alert.style.apply(evidenzia_righe_scadenza, axis=1)
                st.dataframe(styler_colorato, use_container_width=True, hide_index=True)
            else:
                st.dataframe(df_vista_alert, use_container_width=True, hide_index=True)
        else:
            st.info("Nessun adempimento presente nel registro.")
# ==================================================================
# --- SEZIONE 4: ANALISI SEGNALAZIONI NEAR MISS ---
# ==================================================================
if nav == "Analisi Segnalazioni Near Miss":
    st.header("Analisi approfondita delle segnalazioni Near Miss")

    if "autenticato_rspp" not in st.session_state:
        st.session_state.autenticato_rspp = False

    if not st.session_state.autenticato_rspp:
        pwd_rspp = st.text_input(
            "Inserisci la Password di Accesso",
            type="password",
            key="pwd_rspp_tab",
        )
        if st.button("Convalida Accesso", use_container_width=True):
            if pwd_rspp == "hse2026":
                st.session_state.autenticato_rspp = True
                st.rerun()
            else:
                st.error("Credenziali errate.")

    if st.session_state.autenticato_rspp:
        st.success("Autenticato")

        # --- PERCORSI DEI FILE ---
        FILE_NEAR_MISS = "segnalazioni_near_miss.csv"
        FILE_MANUTENZIONE = os.path.join(
            "Segnalazione_NM_Manutenzione", "manutenzione.csv"
        )
        FILE_ANALISI_NM = "analisi_near_miss.csv"
        DIR_IMMAGINI_ANALISI = "immagini_analisi"

        if not os.path.exists(DIR_IMMAGINI_ANALISI):
            os.makedirs(DIR_IMMAGINI_ANALISI)

        # Funzione di supporto per salvare l'intero DataFrame aggiornato su GitHub
        def salva_df_analisi_su_github(
            df_target, message="Aggiornamento analisi_near_miss.csv"
        ):
            try:
                github_token = st.secrets["GITHUB_TOKEN"]
                repo_name = st.secrets["REPO_NAME"]
                g = Github(github_token)
                repo = g.get_repo(repo_name)

                csv_buffer = df_target.to_csv(index=False, sep=";")

                try:
                    file_content = repo.get_contents(FILE_ANALISI_NM)
                    repo.update_file(
                        path=FILE_ANALISI_NM,
                        message=message,
                        content=csv_buffer,
                        sha=file_content.sha,
                    )
                except GithubException as ge:
                    if ge.status == 404:
                        repo.create_file(
                            path=FILE_ANALISI_NM,
                            message=message,
                            content=csv_buffer,
                        )
                    else:
                        raise ge
                return True
            except Exception as ex:
                st.error(
                    f"Errore durante il salvataggio su GitHub: {ex}"
                )
                return False

        # --- LETTURA DELLE SEGNALAZIONI DAI DUE FILE ---
        lista_segnalazioni = []
        mappa_descrizioni = {}

        # 1. Lettura File "segnalazioni_near_miss.csv"
        if os.path.exists(FILE_NEAR_MISS):
            try:
                df_nm = pd.read_csv(
                    FILE_NEAR_MISS,
                    sep=";",
                    on_bad_lines="skip",
                    engine="python",
                )
                for idx, row in df_nm.iterrows():
                    data_ev = str(
                        row.get(
                            "Data Segnalazione", row.get("Data Evento", "N/D")
                        )
                    )
                    segnalatore = str(
                        row.get(
                            "Segnalatore", row.get("Nome Segnalatore", "N/D")
                        )
                    )
                    label = f"{data_ev} | Segnalazione NM | {segnalatore}"
                    lista_segnalazioni.append(label)
                    mappa_descrizioni[label] = str(row.get("Descrizione", ""))
            except Exception as e:
                st.warning(f"Impossibile leggere {FILE_NEAR_MISS}: {e}")

        # 2. Lettura File "manutenzione.csv"
        if os.path.exists(FILE_MANUTENZIONE):
            try:
                df_man = pd.read_csv(
                    FILE_MANUTENZIONE,
                    sep=";",
                    on_bad_lines="skip",
                    engine="python",
                )
                for idx, row in df_man.iterrows():
                    data_ev = str(
                        row.get(
                            "Data Segnalazione", row.get("Data Evento", "N/D")
                        )
                    )
                    segnalatore = str(
                        row.get(
                            "Segnalatore", row.get("Nome Segnalatore", "N/D")
                        )
                    )
                    label = f"{data_ev} | NM_Manutenzione | {segnalatore}"
                    lista_segnalazioni.append(label)
                    mappa_descrizioni[label] = str(row.get("Descrizione", ""))
            except Exception as e:
                st.warning(f"Impossibile leggere {FILE_MANUTENZIONE}: {e}")

        # Lettura file delle analisi
        df_analisi = (
            pd.read_csv(
                FILE_ANALISI_NM, sep=";", on_bad_lines="skip", engine="python"
            )
            if os.path.exists(FILE_ANALISI_NM)
            else pd.DataFrame()
        )

        if "sub_sezione_rspp" not in st.session_state:
            st.session_state.sub_sezione_rspp = "compilazione"

        col_m1, col_m2 = st.columns(2)
        with col_m1:
            if st.button("Apri Nuovo Modulo Analisi", use_container_width=True):
                st.session_state.sub_sezione_rspp = "compilazione"
        with col_m2:
            if st.button("Commento e firma RSPP", use_container_width=True):
                st.session_state.sub_sezione_rspp = "firma"

        st.markdown("---")

        # --- SUBSEZIONE COMPILAZIONE ---
        if st.session_state.sub_sezione_rspp == "compilazione":
            opzioni_tendina = [
                "Nessun collegamento (Crea analisi indipendente)"
            ] + lista_segnalazioni

            selezione_nm = st.selectbox(
                "Seleziona una segnalazione a cui allacciarti:", opzioni_tendina
            )
            desc_def = ""
            if (
                selezione_nm
                != "Nessun collegamento (Crea analisi indipendente)"
            ):
                desc_def = mappa_descrizioni.get(selezione_nm, "")
                st.info("Testo della segnalazione caricato.")

            st.markdown(
                "#### Inserimento immagine o allegato di supporto per l'Analisi (Facoltativo)"
            )
            opzione_media_analisi = st.radio(
                "Scegli la modalità di inserimento file/immagine per l'analisi:",
                ["Nessun file", "Carica file locale", "Scatta foto istantanea"],
                key="scelta_media_analisi_rspp",
            )

            allegato_analisi_nome = "Nessuna"
            if opzione_media_analisi == "Carica file locale":
                file_img_an = st.file_uploader(
                    "Scegli un file per l'analisi",
                    type=["png", "jpg", "jpeg", "pdf", "docx"],
                    key="uploader_analisi_rspp",
                )
                if file_img_an:
                    allegato_analisi_nome = file_img_an.name
                    with open(
                        os.path.join(
                            DIR_IMMAGINI_ANALISI, file_img_an.name
                        ),
                        "wb",
                    ) as f_local:
                        f_local.write(file_img_an.getbuffer())
                    st.caption(
                        f"File '{file_img_an.name}' salvato in {DIR_IMMAGINI_ANALISI}/"
                    )
            elif opzione_media_analisi == "Scatta foto istantanea":
                foto_scattata_an = st.camera_input(
                    "Scatta una foto della verifica tecnica",
                    key="camera_analisi_rspp",
                )
                if foto_scattata_an:
                    allegato_analisi_nome = f"analisi_scatto_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
                    with open(
                        os.path.join(
                            DIR_IMMAGINI_ANALISI, allegato_analisi_nome
                        ),
                        "wb",
                    ) as f_local:
                        f_local.write(foto_scattata_an.getbuffer())
                    st.caption(
                        f"Foto '{allegato_analisi_nome}' archiviata in {DIR_IMMAGINI_ANALISI}/"
                    )

            st.markdown("---")

            # FORM DI INSERIMENTO ANALISI
            with st.form("form_analisi_sup"):
                descrizione_finale = st.text_area(
                    "Integrazione dell'evento", value=desc_def
                )
                incidente_selezionato = st.multiselect(
                    "Incidente potenziale:",
                    [
                        "Caduta dall’alto o in profondità del lavoratore",
                        "Caduta in piano del lavoratore",
                        "Movimento incoordinato del lavoratore (che provoca urto contro, durante uso di attrezzatura manuale, …)",
                        "Caduta dall’alto di gravi",
                        "Proiezione di solidi",
                        "Avviamento inatteso/inopportuno di veicolo, macchina, attrezzatura, ecc.",
                        "Collisione/Urto alla guida di mezzo (contro elementi dell'ambiente di lavoro)",
                        "Investimento (anche mancato) da mezzi, veicoli, oggetti in movimento",
                        "Trascinato, impigliato, afferrato",
                        "Colpito, urtato da",
                        "Urtare contro, andare a sbattere",
                        "Tagliarsi, pungersi",
                        "Sollecitazioni fisiche (rumore, vibrazioni, radiazioni ecc.)",
                        "Disturbi causati da animali",
                        "Ribaltamento mezzo",
                        "Contatto elettrico diretto/indiretto",
                        "Esplosioni, Sviluppo di fiamme",
                        "Fuoriuscita di gas, fumi, aerosol e liquidi",
                        "Contatto con organi lavoratori in movimento",
                        "Contatto con oggetti o materiali caldi, fiamme libere, etc. (nella loro abituale sede)",
                        "Contatto con gas, fumi, aerosol e liquidi (nella loro abituale sede",
                        "Contatto con oggetti o materiali a bassissima temperatura (nella loro abituale sed",
                        "Stretto, schiacciato da",
                        "Travolta, sommerso da",
                        "Travolto, investito da",
                        "Danni alla salute a causa di una postura sbagliata",
                        "Sforzo eccessivo per trasporto a mano di carichi",
                        "Contatto con sostanze pericolose",
                        "Annegamento",
                        "Altro",
                    ],
                )
                attivita_selezionata = st.multiselect(
                    "Attività svolta:",
                    [
                        "Lavori manuali senza utensili",
                        "Lavori manuali con utensili",
                        "Azionamento macchine",
                        "Preparazione dei lavori",
                        "Eliminazione guasti, riparazioni",
                        "Lavori di manutenzione",
                        "Lavori di ampliamento, prove/test",
                        "Imballaggio, disimballaggio",
                        "Trasporto a mano, compresi il carico e lo scarico",
                        "Movimentazione e trasporto con apparecchi",
                        "Movimentazione e trasporto con apparecchi mobili",
                        "Smistamento, manovra",
                        "Andare avanti e indietro",
                        "Pulizia, riordino locali",
                        "Lavaggio, cambio abiti",
                        "Pausa, ristoro",
                        "Attività sportive, giochi, liti",
                        "Contatto con animali vivi",
                        "Altro",
                    ],
                )
                cause_selezionate = st.multiselect(
                    "Cause radice:",
                    [
                        "Errore procedurale (disattenzione, scarsa conoscenza procedure operative, …)",
                        "Illuminazione non idonea o assente",
                        "Problema di comunicazione (lingua, incertezza nei ruoli e/o compiti)",
                        "Assenza o inadeguatezza di barriere, protezioni, parapetti, armatur",
                        "Mancanza/inadeguatezza di procedure operative",
                        "Spazi inadeguati su postazioni di lavoro",
                        "Mancanza di protezioni sull'attrezzatura",
                        "Assenza o inadeguatezza di aree di stoccaggio",
                        "Carenza (inadeguatezza) di protezioni sull'attrezzatura",
                        "Presenza imprevista di liquidi (acqua, olio, …)",
                        "Presenza imprevista di gas, vapori",
                        "Anomalia/guasto in avviamento/arresto/esercizio (funzionamento)",
                        "Unica attrezzatura disponibile ma non idonea alla lavorazione",
                        "Criticità su impianti generali a supporto dell'area di lavoro (sistemi di ventilazione, aerazione)",
                        "Assenza di attrezzature idonee alla lavorazione",
                        "Presenza di elettricità/linea elettrica accessibile",
                        "Stoccaggio/etichettatura errato di materiali",
                        "Livelli di rumorosità inadeguati",
                        "Problema legato alle caratteristiche/trasformazioni di materiali",
                        "Mancanza o Uso errato DPI",
                        "DPI non fornito",
                        "DPI inadeguato",
                        "Segnaletica di sicurezza/Cartellonistica inadeguata o assente",
                        "Assenza o inadeguatezza di percorsi in sicurezza, vie di transito, uscite di emergenza (ingombro di materiali, irregolarità su pavimentazioni)"
                        "Altro",
                    ],
                )
                storico_eventi = st.radio(
                    "Già verificato in passato?",
                    ["Sì frequentemente", "Sì raramente", "No"],
                )
                criticita_selezionate = st.multiselect(
                    "Criticità:",
                    [
                        "Vigilanza/Coordinamento",
                        "Dvr/duvri/psc/pos",
                        "Emergenze e Antincendio",
                        "Piani di manutenzione e pulizia",
                        "Informazione"
                        "Formazione carente",
                        "Sorveglianza sanitaria",
                        "Verifiche periodiche e certificazione conformità impianti",
                        "Primo soccorso",
                        "Verifica idoneità tecnico-professionale",
                        "Nomine e designazioni",
                        "Nessuna",
                    ],
                )

                colX, colY = st.columns(2)
                with colX:
                    danno_strutture = st.radio(
                        "Danno a strutture",
                        ["nessuno", "lieve", "medio", "notevole"],
                    )
                    danno_produttivita = st.radio(
                        "Danno produttivo",
                        ["nessuna", "breve", "media", "rilevante"],
                    )
                with colY:
                    danno_persone = st.radio(
                        "Danno potenziale persone",
                        ["nessuno", "lieve", "grave", "gravissimo"],
                    )
                    frequenza = st.radio(
                        "Frequenza stimata",
                        ["rara", "frequente", "molto frequente"],
                    )

                submit_button = st.form_submit_button("Salva Modulo Direzione")

            # PULSANTE DI INVIO
            if submit_button:
                now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

                # Costruzione riga dati
                nuova_risposta = {
                    "Data Analisi": now_str,
                    "Segnalazione Collegata": selezione_nm,
                    "Descrizione": descrizione_finale.replace("\n", " ")
                    .replace("\r", " ")
                    .replace(";", ","),
                    "Incidente": ", ".join(incidente_selezionato),
                    "Attività": ", ".join(attivita_selezionata),
                    "Cause": ", ".join(cause_selezionate),
                    "Storico": storico_eventi,
                    "Criticità": ", ".join(criticita_selezionate),
                    "Danno Strutture": danno_strutture,
                    "Danno Produttività": danno_produttivita,
                    "Danno Persone": danno_persone,
                    "Frequenza": frequenza,
                    "Commento RSPP": "",
                    "Firma RSPP (Stato)": "Non Firmato",
                    "Allegato Analisi": allegato_analisi_nome,
                }

                # Unisci il nuovo record al DataFrame o crealo se non esiste
                df_nuovo_rec = pd.DataFrame([nuova_risposta])
                df_totale = pd.concat(
                    [df_analisi, df_nuovo_rec], ignore_index=True
                )

                if salva_df_analisi_su_github(
                    df_totale, f"Nuova analisi near miss ({now_str})"
                ):
                    st.session_state[
                        "ultima_segnalazione_analisi_near_miss"
                    ] = nuova_risposta
                    st.success(
                        "Analisi acquisita e salvata con successo su GitHub!"
                    )
                    time.sleep(1)
                    st.rerun()

        # --- SUBSEZIONE COMMENTO E FIRMA RSPP ---
        elif st.session_state.sub_sezione_rspp == "firma":
            if df_analisi.empty:
                st.warning(
                    "Nessuna analisi presente nel file 'analisi_near_miss.csv'."
                )
            else:
                opzioni_r = []
                mappatura = {}
                for idx, r in df_analisi.iterrows():
                    testo_o = f"Analisi del {r.get('Data Analisi', 'N/D')} | Collegamento: {r.get('Segnalazione Collegata', 'Nessuno')}"
                    opzioni_r.append(testo_o)
                    mappatura[testo_o] = idx
        
                scelta_rec = st.selectbox(
                    "Scegli l'analisi da integrare con commento e firma:",
                    opzioni_r,
                )
                idx_sel = mappatura[scelta_rec]
        
                if "Commento RSPP" not in df_analisi.columns:
                    df_analisi["Commento RSPP"] = ""
                if "Firma RSPP (Stato)" not in df_analisi.columns:
                    df_analisi["Firma RSPP (Stato)"] = "Non Firmato"
        
                comm_pre = (
                    str(df_analisi.at[idx_sel, "Commento RSPP"])
                    if pd.notna(df_analisi.at[idx_sel, "Commento RSPP"])
                    else ""
                )
                comm_in = st.text_area(
                    "Note / Commenti del Professionista (RSPP):",
                    value=comm_pre,
                )
                file_f = st.file_uploader(
                    "Carica Firma Grafica",
                    type=["png", "jpg", "jpeg"],
                    key="uploader_firma_rspp",
                )
        
                if file_f:
                    st.image(file_f, width=150)
        
                if st.button("Salva ed Applica Modifiche in Riga"):
                    if not comm_in.strip():
                        st.error("Inserire un commento prima di salvare.")
                    else:
                        commento_pulito = (
                            comm_in.strip()
                            .replace("\n", " ")
                            .replace("\r", " ")
                            .replace(";", ",")
                        )
        
                        # 1. Assegna i tipi di dato corretti per evitare TypeError
                        df_analisi["Commento RSPP"] = df_analisi["Commento RSPP"].astype(object)
                        df_analisi["Firma RSPP (Stato)"] = df_analisi["Firma RSPP (Stato)"].astype(object)
        
                        # 2. Aggiorna i valori nel DataFrame PRIMA di salvare
                        df_analisi.at[idx_sel, "Commento RSPP"] = commento_pulito
                        if file_f:
                            df_analisi.at[idx_sel, "Firma RSPP (Stato)"] = "Firmato"
        
                        # 3. Invia il DataFrame aggiornato a GitHub
                        if salva_df_analisi_su_github(df_analisi):
                            st.success(
                                "Commento e firma RSPP salvati e sincronizzati con successo su GitHub!"
                            )
                            time.sleep(1)
                            st.rerun()



# ==================================================================
# --- SEZIONE 5: Consultazione CONFORMITÀ LEGISLATIVA (RIFERIMENTI REALI) ---
# ==================================================================
if nav == "Consultazione":
    st.header("Consultazione Documenti e Procedure")
    
    if "autenticato_upload_legale" not in st.session_state:
        st.session_state.autenticato_upload_legale = False
        
    correct_pwd_legale = st.secrets.get("PASSWORD_SEZIONE", "")
    github_token = st.secrets.get("GITHUB_TOKEN", "")
    repo_name = st.secrets.get("REPO_NAME", "")
    folder_repo_conformita = "documenti_conformita"

    # Helper function per recuperare la lista dei file su GitHub
    def elenca_file_github_conformita():
        if github_token and repo_name:
            url = f"https://api.github.com/repos/{repo_name}/contents/{folder_repo_conformita}"
            headers = {"Authorization": f"token {github_token}"}
            try:
                res = requests.get(url, headers=headers)
                if res.status_code == 200:
                    return res.json()  # Restituisce la lista di oggetti dei file
            except Exception:
                pass
        return []

    # Helper function per scaricare un singolo file da GitHub
    def scarica_file_github_conformita(nome_file):
        if github_token and repo_name:
            url = f"https://raw.githubusercontent.com/{repo_name}/main/{folder_repo_conformita}/{nome_file}"
            headers = {"Authorization": f"token {github_token}"}
            try:
                res = requests.get(url, headers=headers)
                if res.status_code == 200:
                    return res.content
            except Exception:
                pass
        
        # Fallback locale
        local_path = os.path.join(DIR_CONFORMITA if 'DIR_CONFORMITA' in globals() else folder_repo_conformita, nome_file)
        if os.path.exists(local_path):
            with open(local_path, "rb") as lf:
                return lf.read()
        return None

    col_bot, col_upload = st.columns([1, 1])
    
    # -----------------------------------------------------------
    # AREA PUBBLICA (col_bot)
    # -----------------------------------------------------------
    with col_bot:
        st.subheader("Assistente di Scansione Istantanea")
        st.caption("Accesso Pubblico - Ricerca libera e lettura QR Code.")
        # 1. Ricerca Documentale
        prompt_utente = st.text_input("Scrivi l'oggetto della richiesta (es: requisiti antincendio, obblighi, nomine):", key="txt_prompt_sup")
        if st.button("Interroga Archivio Normativo", use_container_width=True):
            if prompt_utente.strip():
                with st.spinner("Scansione testi in corso..."):
                    dir_target = DIR_CONFORMITA if 'DIR_CONFORMITA' in globals() else folder_repo_conformita
                    risultati_ricerca = cerca_nei_documenti(prompt_utente, dir_target)
                    if risultati_ricerca == "specifica":
                        st.warning("Inserisci una query di ricerca più dettagliata.")
                    elif not risultati_ricerca:
                        st.info("Nessun paragrafo corrisponde alla ricerca basata su parole chiave.")
                    else:
                        st.markdown("---")
                        st.markdown("### Riscontri della Ricerca:")
                        for r in risultati_ricerca:
                            st.markdown(f"**Fonte:** {r['fonte']} (Rilevanza: {r['score']}%)")
                            st.info(f"> {r['testo'][:1000]}")
                            try:
                                # Recupero del file dal repository GitHub o locale
                                dati_file = scarica_file_github_conformita(r['nome_file'])
                                if dati_file:
                                    st.download_button(
                                        label=f"Apri / Scarica File Originale ({r['nome_file']})",
                                        data=dati_file,
                                        file_name=r['nome_file'],
                                        mime="application/octet-stream",
                                        key=f"dl_{r['fonte'].replace(' ','_')}_{time.time()}"
                                    )
                                else:
                                    st.error("Impossibile recuperare il file dal repository.")
                            except Exception as e_file:
                                st.error(f"Impossibile agganciare l'anteprima del file: {e_file}")
            else:
                st.warning("Inserire un testo di ricerca prima di premere il pulsante.")
                
        # -----------------------------------------------------------
        # POSIZIONAMENTO PUBBLICO E STRUTTURAZIONE NATIVA QR CODE
        # -----------------------------------------------------------
        st.markdown("---")
        st.subheader("Lettura Link tramite QR Code")
        st.caption("Strumento pubblico di codifica rapida per la documentazione HSE aziendale.")
        
        # Scelta della modalità di input
        metodo_input = st.radio(
            "Scegli come inserire il codice:", 
            ["Carica file dal dispositivo", "Fotocamera"], 
            horizontal=True,
            key="qr_pub_metodo"
        )
        immagine_input = None
        # Gestione logica in base alla scelta
        if metodo_input == "Carica file dal dispositivo":
            immagine_input = st.file_uploader(
                "Seleziona un'immagine (png, jpg, jpeg):", 
                type=["png", "jpg", "jpeg"], 
                key="up_qr_pub"
            )
        else:
            immagine_input = st.camera_input("Scatta una foto del QR Code:", key="cam_qr_pub")
        
        risultato_pub = None
        # Elaborazione dell'immagine (funziona con entrambi gli input)
        if immagine_input:
            # La funzione decodifica_qr_opencv gestisce l'oggetto file_like (file o stream camera)
            risultato_pub = decodifica_qr_opencv(immagine_input)
    
        if risultato_pub:
            st.success("Codice QR decodificato con successo!")
            st.code(risultato_pub, language="text")
        
            if risultato_pub.startswith("http://") or risultato_pub.startswith("https://"):
                st.markdown(f"[Clicca qui per aprire il link rilevato]({risultato_pub})")
        else:
            st.warning("Nessun codice QR valido rilevato nell'immagine.")

    # -----------------------------------------------------------
    # AREA PRIVATA / AMMINISTRATIVA (col_upload)
    # -----------------------------------------------------------        
    with col_upload:
        st.subheader("Area Amministrazione Archivio")
        
        if not st.session_state.autenticato_upload_legale:
            st.caption("Questa funzione richiede i privileges di Amministrazione per inserire nuovi testi ed estrarre la check-list.")
            pwd_leg = st.text_input("Inserisci la password amministrativa", type="password", key="pwd_leg_tab")
            if st.button("Sblocca Funzioni Avanzate", use_container_width=True):
                if pwd_leg == correct_pwd_legale:
                    st.session_state.autenticato_upload_legale = True
                    st.rerun()
                else:
                    st.error("Password di conformità errata. Accesso negato.")
        else:
            st.success("Modalità di amministrazione attiva.")
            if st.button("Chiudi sessione Amministratore (Blocca)", use_container_width=True):
                st.session_state.autenticato_upload_legale = False
                st.rerun()
                
            # --- Funzioni Admin ---
            tab_admin1, tab_admin2 = st.tabs(["Documenti & Check-list", "Gestione QR & Normativa"])
            
            with tab_admin1:
                st.markdown("Conformità legislativa e documentale")
                st.subheader("Caricamento Documenti")
                file_caricati = st.file_uploader(
                    "Trascina qui Decreti, Testi Unici o Tabelle Word/Excel da archiviare nel repository online GitHub (cartella 'documenti_conformità'):",
                    type=["pdf", "csv", "xlsx", "xls", "docx", "doc", "jpg", "jpeg", "png"],
                    accept_multiple_files=True,
                    key="uploader_leg_sup"
                )
                if file_caricati:
                    for f in file_caricati:
                        file_bytes = f.getbuffer().tobytes()
                        path_in_repo = f"{folder_repo_conformita}/{f.name}"
                        
                        # Backup locale
                        dir_target = DIR_CONFORMITA if 'DIR_CONFORMITA' in globals() else folder_repo_conformita
                        os.makedirs(dir_target, exist_ok=True)
                        save_path = os.path.join(dir_target, f.name)
                        with open(save_path, "wb") as f_out:
                            f_out.write(file_bytes)
                        
                        # Salvataggio su repository GitHub
                        if 'save_to_github' in globals():
                            save_to_github(path_in_repo, file_bytes, f"Caricamento documento conformità {f.name}")
                            
                    st.success("File salvati con successo in locale e sul repository GitHub!")
                
                # Check-list
                st.markdown("---")
                st.subheader("Estrazione Check-list Documentale")
            
                if st.button("Genera ed Elabora Check-list CSV", use_container_width=True, key="btn_genera_checklist_doc"):
                    file_github = elenca_file_github_conformita()
                    dir_target = DIR_CONFORMITA if 'DIR_CONFORMITA' in globals() else folder_repo_conformita
                    
                    file_presenti_nomi = []
                    if file_github:
                        file_presenti_nomi = [item["name"] for item in file_github if item["type"] == "file"]
                    elif os.path.exists(dir_target):
                        file_presenti_nomi = [f for f in os.listdir(dir_target) if os.path.isfile(os.path.join(dir_target, f))]

                    if not file_presenti_nomi:
                        st.warning("L'archivio è vuoto. Carica almeno un documento.")
                    else:
                        with st.spinner("Analisi ed estrazione strutturata..."):
                            righe_checklist = []
                            for fn in file_presenti_nomi:
                                fp = os.path.join(dir_target, fn)
                                data_caricamento = datetime.now().strftime("%d-%m-%Y")
                                if os.path.exists(fp):
                                    timestamp_creazione = os.path.getmtime(fp)
                                    data_caricamento = datetime.fromtimestamp(timestamp_creazione).strftime("%d-%m-%Y")
                                    
                                titolo, regolamento, doc_produrre, settore_aziendale = estrai_info_checklist(fp if os.path.exists(fp) else fn)
                            
                                righe_checklist.append({
                                    "Quando fu caricata": data_caricamento,
                                    "Titolo": titolo,
                                    "Settore aziendale": settore_aziendale,
                                    "Regolamento o legge di riferimento": regolamento,
                                    "Documenti da produrre o allegati": doc_produrre
                                })
                            df_checklist = pd.DataFrame(righe_checklist)
                            st.session_state["cached_df_checklist"] = df_checklist
                            st.success("Check-list aggiornata in memoria.")
                            
                    if "cached_df_checklist" in st.session_state:
                        df_visualizza = st.session_state["cached_df_checklist"]
                        st.dataframe(df_visualizza, use_container_width=True, hide_index=True)
                        csv_buffer = df_visualizza.to_csv(index=False, sep=";").encode('utf-8')
                        st.download_button(
                            label="Scarica Check-list in formato CSV",
                            data=csv_buffer,
                            file_name=f"checklist_conformita_{datetime.now().strftime('%Y%m%d')}.csv",
                            mime="text/csv",
                            use_container_width=True
                        )
                        
            # -----------------------------------------------------------
            # GENERA QRCode e Gestione Normativa
            # -----------------------------------------------------------    
            with tab_admin2:
                # Generazione QR
                st.markdown("#### Genera QR Code")
                opzione_qr = st.radio("Seleziona operazione QR:", ["Genera QR Code", "Leggi/Decodifica QR Code"], key="radio_opzioni_qr_pubblico")
                img_ottenuta = None
        
                if opzione_qr == "Genera QR Code":
                    testo_da_convertire = st.text_input("Inserisci l'URL o il testo da inserire nel QR Code:", placeholder="https://www.gazzettaufficiale.it/", key="txt_input_qr_gen_unq")
                    if testo_da_convertire.strip():
                        img_ottenuta = genera_qr_nativo(testo_da_convertire)
        
                    if img_ottenuta:
                        percorso_temp_qr = "temp_generated_qr.png"
                        img_ottenuta.save(percorso_temp_qr)
                
                        st.image(percorso_temp_qr, width=220, caption="Codice QR generato con successo!")
                
                        with open(percorso_temp_qr, "rb") as f_qr:
                            st.download_button(
                                label="Scarica Immagine QR Code (.png)",
                                data=f_qr.read(),
                                file_name="qr_code_hse.png",
                                mime="image/png",
                                use_container_width=True,
                                key="btn_dl_qr_unq"
                            )
                    
                elif opzione_qr == "Leggi/Decodifica QR Code":
                    file_qr_caricato = st.file_uploader("Carica un'immagine contenente un codice QR:", type=["png", "jpg", "jpeg"], key="uploader_file_qr_nativo")
                    if file_qr_caricato:
                        risultato_priv = decodifica_qr_opencv(file_qr_caricato)
                        if risultato_priv:
                            st.success("Codice QR decodificato con successo!")
                            st.code(risultato_priv, language="text")
                    
                            if risultato_priv.startswith("http://") or risultato_priv.startswith("https://"):
                                st.markdown(f"[Clicca qui per aprire il link rilevato]({risultato_priv})")
                        else:
                            st.warning("Nessun codice QR valido rilevato all'interno dell'immagine caricata.")

                # -----------------------------------------------------------
                # UNICA AREA: GESTIONE DATABASE NORMATIVO & CONFRONTO LEGISLATIVO
                # -----------------------------------------------------------
                st.markdown("---")
                st.subheader("Confronto Legislativo & Stato Conformità")
            
                with st.expander("Aggiungi nuova norma al database interno (normative.json su GitHub)"):
                    k = st.text_input("Parola chiave (es: rifiuti):", key="txt_norma_key_unique")
                    t = st.text_input("Titolo:", key="txt_norma_titolo_unique")
                    u = st.text_input("URL ufficiale:", key="txt_norma_url_unique")
                    d = st.text_area("Descrizione della norma:", key="txt_norma_desc_unique")
                    
                    if st.button("Salva nel database", key="btn_salva_norma_unique", use_container_width=True): 
                        if k.strip():
                            esito = salva_normativa(k, t, u, d)
                            if esito:
                                st.success(f"Norma '{k}' salvata e sincronizzata con successo nel database su GitHub!")
                                time.sleep(1)
                                st.rerun()
                        else:
                            st.warning("Inserisci obbligatoriamente una parola chiave prima di salvare.")

                note_libere = st.text_area("Analisi testo libero / criticità aziendali:", placeholder="es: gestione rifiuti pericolosi, amianto, dispositivi DPI, ecc.", key="txt_note_libere_conformita_unq")
            
                if st.button("Avvia Analisi di Conformità", use_container_width=True, key="btn_avvia_analisi_conformita_avanzata"):
                    with st.spinner("Scansione in corso di tutti i documenti di GitHub e verifica dei link nel database normativo..."):
                        
                        # 1. Caricamento Database Normativo (da GitHub o locale)
                        db_norme = {}
                        if github_token and repo_name:
                            url_github_api = f"https://api.github.com/repos/{repo_name}/contents/normative.json"
                            headers = {"Authorization": f"token {github_token}"}
                            try:
                                res = requests.get(url_github_api, headers=headers)
                                if res.status_code == 200:
                                    import base64
                                    content_decoded = base64.b64decode(res.json().get("content", "")).decode("utf-8")
                                    db_norme = json.loads(content_decoded)
                            except Exception:
                                pass

                        if not db_norme and os.path.exists("normative.json"):
                            try:
                                with open("normative.json", "r", encoding="utf-8") as f:
                                    db_norme = json.load(f)
                            except Exception as e:
                                st.error(f"Errore lettura database: {e}")

                        if not db_norme:
                            st.warning("Il database normativo (normative.json) è vuoto o assente. Aggiungi una norma nel pannello sopra.")
                        
                        # 2. Scansione Documenti da GitHub
                        file_github = elenca_file_github_conformita() if 'elenca_file_github_conformita' in globals() else []
                        dir_target = DIR_CONFORMITA if 'DIR_CONFORMITA' in globals() else folder_repo_conformita
                        
                        contenuto_documenti = {}
                        
                        if file_github:
                            for item in file_github:
                                if item.get("type") == "file":
                                    n_file = item["name"]
                                    bytes_file = scarica_file_github_conformita(n_file) if 'scarica_file_github_conformita' in globals() else None
                                    txt_ext = estrai_testo_da_bytes_robusto(n_file, bytes_file)
                                    contenuto_documenti[n_file] = txt_ext
                        elif os.path.exists(dir_target):
                            for n_file in os.listdir(dir_target):
                                p_file = os.path.join(dir_target, n_file)
                                if os.path.isfile(p_file):
                                    with open(p_file, "rb") as lf:
                                        txt_ext = estrai_testo_da_bytes_robusto(n_file, lf.read())
                                        contenuto_documenti[n_file] = txt_ext

                        testo_globale_archivio = " ".join(contenuto_documenti.values()).lower() + " " + note_libere.lower()
                        
                        st.markdown("### Risultati Analisi di Conformità Documentale")
                        st.caption(f"Documenti analizzati nell'archivio 'documenti_conformità': **{len(contenuto_documenti)}**")
                        
                        if not db_norme:
                            st.stop()

                        trovati_totali = 0
                        conformi_totali = 0
                        non_conformi_totali = 0

                        # 3. Analisi per ogni Norma
                        for chiave, info in db_norme.items():
                            chiave_clean = chiave.strip().lower()
                            url_norma = info.get("url", "")
                            
                            testo_link = estrai_contenuto_link(url_norma) if url_norma else ""
                            
                            doc_corrispondenti = [
                                nome_doc for nome_doc, testo_doc in contenuto_documenti.items()
                                if chiave_clean in testo_doc.lower() or chiave_clean in nome_doc.lower()
                            ]
                            
                            is_conforme = len(doc_corrispondenti) > 0 or (chiave_clean in note_libere.lower())

                            trovati_totali += 1
                            if is_conforme:
                                conformi_totali += 1
                            else:
                                non_conformi_totali += 1

                            st.markdown("---")
                            col_stato, col_info = st.columns([1, 3])
                            
                            with col_stato:
                                if is_conforme:
                                    st.success(" CONFORME")
                                    st.caption(f"Trovati {len(doc_corrispondenti)} documenti relativi")
                                else:
                                    st.error(" NON CONFORME")
                                    st.caption("Nessun documento idoneo rilevato")

                            with col_info:
                                st.subheader(f"Tema: {info.get('titolo', chiave)}")
                                st.markdown(f"**Parola chiave monitorata:** `{chiave}`")
                                st.write(f"**Descrizione:** {info.get('desc', 'N.D.')}")
                                
                                if url_norma:
                                    st.link_button("Vai alla Fonte/Link Ufficiale", url_norma)

                                if is_conforme:
                                    st.markdown("** Documenti presenti che attestano la conformità:**")
                                    for doc_ok in doc_corrispondenti:
                                        st.markdown(f"- `{doc_ok}`")
                                else:
                                    doc_mancanti, azioni_suggerite = genera_piano_adeguamento_e_gap(chiave, info, testo_link)
                                    
                                    st.warning("⚠️ **AZIONI CORRETTIVE RICHIESTE PER L'ADEGUAMENTO:**")
                                    
                                    st.markdown("####  1. Documenti specifici da produrre / allegare:")
                                    for doc_m in doc_mancanti:
                                        st.markdown(f" * {doc_m}")

                                    st.markdown("####  2. Azioni operative da intraprendere:")
                                    for az_m in azioni_suggerite:
                                        st.markdown(f" * {az_m}")

                        st.markdown("---")
                        st.markdown("###  Sintesi finale dell'Analisi di Conformità")
                        col_m1, col_m2, col_m3 = st.columns(3)
                        col_m1.metric("Totale Temi Controllati", trovati_totali)
                        col_m2.metric("Conformi", conformi_totali)
                        col_m3.metric("Non Conformi (Gap)", non_conformi_totali, delta_color="inverse")
# ==================================================================
# --- SEZIONE 6: Analisi - Fase 2 ---
# ==================================================================
# ==================================================================
# --- RECUPERO SICURO DELLE VARIABILI DI GITHUB ---
# ==================================================================
# Cerca prima nei Secrets di Streamlit Cloud, poi nelle variabili di ambiente del sistema.
GITHUB_TOKEN = st.secrets.get("GITHUB_TOKEN", os.environ.get("GITHUB_TOKEN", ""))
REPO_NAME = st.secrets.get("REPO_NAME", os.environ.get("REPO_NAME", ""))

if nav == "Analisi - Fase 2":
    st.header("Analisi - Fase 2 delle segnalazioni near miss")
    # --- SEZIONE 2: ANALISI ISHIKAWA ---
    if "autenticato_fase2" not in st.session_state:
        st.session_state.autenticato_fase2 = False
        
    if not st.session_state.autenticato_fase2:
        pwd_fase2 = st.text_input("Inserisci la Password di Accesso", type="password", key="pwd_fase2_tab")
        if st.button("Convalida Accesso", use_container_width=True):
            if pwd_fase2 == "hse2026":
                st.session_state.autenticato_fase2 = True
                st.rerun()
            else:
                st.error("Credenziali errate.")
                
    if st.session_state.autenticato_fase2:

        # Caricamento e unione dati per dropdown
        opzioni = ["Nessuna (Nuova analisi)"]
            
        # Leggi Near Miss
        if os.path.exists(FILE_NEAR_MISS):
            df_nm = pd.read_csv(FILE_NEAR_MISS, sep=";")
            for idx, r in df_nm.iterrows():
                opzioni.append(f"NM | {r.get('Data Segnalazione', 'N/D')} | {r.get('Tipo Evento', 'Evento')}")
            
        # Leggi Analisi già fatte
        if os.path.exists(FILE_ANALISI_NM):
            df_an = pd.read_csv(FILE_ANALISI_NM, sep=";")
            for idx, r in df_an.iterrows():
                opzioni.append(f"AN | {r.get('Data Analisi', 'N/D')} | Collegamento: {r.get('Segnalazione Collegata', 'Analisi')}")
            
        scelta_rif = st.selectbox("Seleziona evento/analisi collegata:", opzioni)

        # Campi Input (Ishikawa)
        c_ma, c_mb = st.columns(2)
        with c_ma:
            macchina = st.text_area("Macchina / Infrastruttura", key="inp_m")
            materiale = st.text_area("Materiale", key="inp_mat")
        with c_mb:
            metodo = st.text_area("Metodo / Procedura", key="inp_met")
            manodopera = st.text_area("Manodopera / Comportamento", key="inp_man")
            
        w1 = st.text_input("Perché 1:", key="w1")
        w2 = st.text_input("Perché 2:", key="w2")
        w3 = st.text_input("Perché 3:", key="w3")
        w4 = st.text_input("Perché 4:", key="w4")
        w5 = st.text_input("Perché 5:", key="w5")
        conc = st.text_area("Conclusioni:", key="conc")

        if st.button("Genera Report e File"):
            temp_img = os.path.join(DIR_ANALISI, "temp_ishikawa.png")
                
            # Matplotlib setup
            fig, ax = plt.subplots(figsize=(10, 6))
            ax.axis('off')
            ax.plot([1, 9], [0, 0], color='black', lw=3)
                
            # Aggiunta branchie (Ishikawa)
            def wrap(text): return "\n".join(textwrap.wrap(str(text), width=20))
            branches = [
                ((2, 0), (1.5, 1.5), f"Macchina:\n{wrap(macchina)}"),
                ((4, 0), (3.5, -1.5), f"Metodo:\n{wrap(metodo)}"),
                ((6, 0), (6.5, 1.5), f"Materiale:\n{wrap(materiale)}"),
                ((8, 0), (8.5, -1.5), f"Manodopera:\n{wrap(manodopera)}")
            ]
            for start, end, label in branches:
                ax.plot([start[0], end[0]], [start[1], end[1]], 'k-', lw=2)
                ax.text(end[0], end[1], label, fontsize=9, ha='center', va='center', bbox=dict(facecolor='white', alpha=0.8))
                
            plt.savefig(temp_img, bbox_inches='tight', dpi=300)
            plt.close()
            st.pyplot(fig)
            
            # Lettura dell'immagine PNG creata in binario per il salvataggio in sessione
            with open(temp_img, "rb") as img_file:
                png_bytes = img_file.read()

            # Raccolta metadati richiesti (Data generazione e file associati)
            data_generazione = datetime.now().strftime("%d-%m-%Y %H:%M:%S")
            file_obbligatorio = FILE_ANALISI_NM
            file_facoltativo = FILE_NEAR_MISS if os.path.exists(FILE_NEAR_MISS) else "Non utilizzato / Assente"

            # Conversione dell'immagine in stringa Base64 per l'inclusione nel JSON
            png_base64 = base64.b64encode(png_bytes).decode('utf-8')

            # Dati strutturati per Export
            dati_report = {
                "Data Generazione Report": data_generazione,
                "File Obbligatorio Associato": file_obbligatorio,
                "File Facoltativo Associato": file_facoltativo,
                "Riferimento Selezionato": scelta_rif,
                "4M": {
                    "Macchina": macchina, 
                    "Materiale": materiale, 
                    "Metodo": metodo, 
                    "Manodopera": manodopera
                },
                "5Whys": [w1, w2, w3, w4, w5],
                "Conclusioni": conc
            }

            # 1. PDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", 'B', 16)
            pdf.cell(0, 10, "Report Analisi Near Miss", ln=True, align='C')
            pdf.ln(10) # Spazio
        
            pdf.set_font("Arial", size=10)
            pdf.cell(0, 6, f"Data Generazione: {data_generazione}", ln=True)
            pdf.cell(0, 6, f"File Obbligatorio: {file_obbligatorio}", ln=True)
            pdf.cell(0, 6, f"File Facoltativo: {file_facoltativo}", ln=True)
            pdf.cell(0, 6, f"Riferimento: {scelta_rif}", ln=True)
            pdf.ln(5)

            # Scrittura Analisi 4M
            pdf.set_font("Arial", 'B', 14)
            pdf.cell(0, 10, "Analisi 4M:", ln=True)
            pdf.set_font("Arial", size=12)
            for key, val in dati_report["4M"].items():
                pdf.cell(0, 10, f"- {key}: {val}", ln=True)
            pdf.ln(5)

            # Scrittura 5Whys (Inclusione di tutti i 5 Perché)
            pdf.set_font("Arial", 'B', 14)
            pdf.cell(0, 10, "5Whys:", ln=True)
            pdf.set_font("Arial", size=12)
            for i, why in enumerate(dati_report['5Whys'], 1):
                pdf.cell(0, 10, f"Perché {i}: {why}", ln=True)
            pdf.ln(5)

            # Scrittura Conclusioni
            pdf.set_font("Arial", 'B', 14)
            pdf.cell(0, 10, "Conclusioni:", ln=True)
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, dati_report['Conclusioni'])
            pdf.ln(5)

            # Inserimento visivo del grafico PNG nel PDF (Diagramma di Ishikawa)
            pdf.set_font("Arial", 'B', 10)
            pdf.cell(0, 10, "Diagramma di Ishikawa:", ln=True)
            pdf.image(temp_img, w=180)

            # Output del file in bytes
            pdf_output = bytes(pdf.output(dest='S')) 
            st.session_state.pdf_bytes = pdf_output

            # 2. XLSX
            import io
            buffer_xlsx = io.BytesIO()
            pd.DataFrame([dati_report["4M"]]).to_excel(buffer_xlsx, index=False)
            st.session_state.xlsx_bytes = buffer_xlsx.getvalue()

            # 3. JSON (Include i metadati, i testi e l'immagine codificata in Base64)
            st.session_state.json_bytes = json.dumps(dati_report, indent=4, ensure_ascii=False).encode('utf-8')
                
            # 4. PNG (File immagine separato)
            st.session_state.png_bytes = png_bytes

            # ==================================================================
            # --- SALVATAGGIO AUTOMATICO SU GITHUB ONLINE ---
            # ==================================================================
            if GITHUB_TOKEN and REPO_NAME:
                # Sanificazione stringa del riferimento evento per uso nei file
                rif_sanitizzato = re.sub(r'[\\/*?:"<>|]', '_', scelta_rif).replace(' ', '_')
                giorno_str = datetime.now().strftime("%Y-%m-%d")

                # Nome base: "Giorno della generazione_riferimento dell'evento_Analisi Fase 2"
                base_filename = f"{giorno_str}_{rif_sanitizzato}_Analisi Fase 2"

                # Dizionario dei file da caricare nella cartella "Analisi_Fase2"
                files_to_upload = {
                    f"Analisi_Fase2/{base_filename}.pdf": pdf_output,
                    f"Analisi_Fase2/{base_filename}_Ishikawa.png": png_bytes,
                    f"Analisi_Fase2/{base_filename}.xlsx": st.session_state.xlsx_bytes,
                    f"Analisi_Fase2/{base_filename}.json": st.session_state.json_bytes
                }

                headers = {
                    "Authorization": f"token {GITHUB_TOKEN}",
                    "Accept": "application/vnd.github.v3+json"
                }

                caricamento_ok = True
                for path_in_repo, content_bytes in files_to_upload.items():
                    url = f"https://api.github.com/repos/{REPO_NAME}/contents/{path_in_repo}"
                    
                    # Controllo se il file esiste già per recuperare lo SHA (per sovrascrittura)
                    res_get = requests.get(url, headers=headers)
                    sha = res_get.json().get("sha") if res_get.status_code == 200 else None

                    # Encoding del contenuto binario in Base64
                    content_b64 = base64.b64encode(content_bytes).decode('utf-8')

                    payload = {
                        "message": f"Auto-save report: {base_filename}",
                        "content": content_b64
                    }
                    if sha:
                        payload["sha"] = sha

                    # Push su GitHub
                    res_put = requests.put(url, json=payload, headers=headers)
                    if res_put.status_code not in [200, 201]:
                        caricamento_ok = False
                        st.error(f"Errore nel salvataggio su GitHub per {path_in_repo}: {res_put.json().get('message')}")

                if caricamento_ok:
                    st.success(f"Report salvato automaticamente in 'Analisi_Fase2/' su GitHub!")
            else:
                st.warning("Variabili GITHUB_TOKEN o REPO_NAME non trovate. Impossibile salvare online.")

            st.session_state.report_ready = True
        # Bottoni Download
        if st.session_state.get("report_ready"):
            col_d1, col_d2, col_d3 = st.columns(3)
            col_d1.download_button("Scarica PDF", st.session_state.pdf_bytes, "Report.pdf")
            col_d2.download_button("Scarica XLSX", st.session_state.xlsx_bytes, "Dati.xlsx")
            col_d3.download_button("Scarica JSON", st.session_state.json_bytes, "Dati.json")

# ==================================================================
# --- SEZIONE 7: KPI ---
# ==================================================================
if nav == "KPI":
    st.header("KPI & Indicatori di Performance HSE")
    
    # ---------------------------------------------------------
    # Autenticazione Password Sezione KPI tramite Secrets
    # ---------------------------------------------------------
    if "autenticato_kpi" not in st.session_state:
        st.session_state.autenticato_kpi = False
        
    correct_pwd_kpi = st.secrets.get("PASSWORD_SEZIONE", "hse2026")
        
    if not st.session_state.autenticato_kpi:
        pwd_kpi = st.text_input("Inserisci la Password di Accesso per la sezione KPI", type="password", key="pwd_kpi_sec")
        if st.button("Convalida Accesso KPI", use_container_width=True):
            if pwd_kpi == correct_pwd_kpi:
                st.session_state.autenticato_kpi = True
                st.rerun()
            else:
                st.error("Password errata.")
        
    if st.session_state.autenticato_kpi:
        def render_sezione_kpi():
            st.title("Sezione KPI Aziendali")
            st.markdown(
                "Gestione centralizzata, monitoraggio e storico delle performance di"
                " sicurezza sul lavoro."
            )
            st.markdown("---")
            oggi = datetime.today().date()
            data_str = oggi.strftime('%Y-%m-%d')
            
            # Creazione della cartella KPI locale se non esiste
            os.makedirs("KPI", exist_ok=True)
            
            # Percorsi per la persistenza automatica dei dati tra i refresh
            path_master_pers = os.path.join("KPI", "master_kpi_definitions.csv")
            path_storico_pers = os.path.join("KPI", "storico_misure.csv")
            
            # Helper function per il salvataggio su GitHub sicuro
            def salva_su_github_kpi(rel_path, df_to_save, commit_msg):
                try:
                    os.makedirs(os.path.dirname(rel_path), exist_ok=True)
                    df_to_save.to_csv(rel_path, index=False, encoding='utf-8')
                    csv_bytes = df_to_save.to_csv(index=False, encoding='utf-8').encode('utf-8')
                    if 'save_to_github' in globals():
                        return save_to_github(rel_path, csv_bytes, commit_msg)
                    return True
                except Exception as e:
                    st.error(f"Errore durante il salvataggio: {e}")
                    return False

            # ==========================================
            # ARCHIVIO CONDIVISO UNICO (PRESENTE IN ENTRAMBE LE TAB)
            # ==========================================
            if "t1_kpi_definitions" not in st.session_state:
                if os.path.exists(path_master_pers):
                    try:
                        st.session_state.t1_kpi_definitions = pd.read_csv(path_master_pers).fillna("")
                    except Exception:
                        st.session_state.t1_kpi_definitions = pd.DataFrame()
                else:
                    st.session_state.t1_kpi_definitions = pd.DataFrame([
                        {
                            "ID_KPI": "NM-01",
                            "Nome_KPI": "Numero Totale Near Miss Segnalati",
                            "Categoria": "Volume",
                            "Unita_Misura": "Numero",
                            "Cadenza": "Settimanale",
                            "Formula": "Conteggio assoluto segnalazioni validate",
                            "Parametro_1": "Segnalazioni valide",
                            "Parametro_2": "",
                            "Parametro_3": "",
                            "Baseline": "Media storica di riferimento: 10",
                        },
                        {
                            "ID_KPI": "NM-02",
                            "Nome_KPI": "Tasso di Risoluzione Near Miss",
                            "Categoria": "Efficacia",
                            "Unita_Misura": "%",
                            "Cadenza": "Settimanale",
                            "Formula": "(Parametro_1 / Parametro_2) * 100",
                            "Parametro_1": "Near Miss Chiusi",
                            "Parametro_2": "Near Miss Totali",
                            "Parametro_3": "",
                            "Baseline": "Target minimo aziendale: 75%",
                        },
                        {
                            "ID_KPI": "NM-03",
                            "Nome_KPI": "Tempo Medio di Presa in Carico",
                            "Categoria": "Reattività",
                            "Unita_Misura": "Giorni",
                            "Cadenza": "Bisettimanale",
                            "Formula": "Parametro_1 / Parametro_2",
                            "Parametro_1": "Somma giorni attesa",
                            "Parametro_2": "Numero totale pratiche",
                            "Parametro_3": "",
                            "Baseline": "Target massimo accettabile: 2 giorni",
                        },
                    ])
                    salva_su_github_kpi(path_master_pers, st.session_state.t1_kpi_definitions, "Inizializzazione master_kpi_definitions.csv")

            # Storico delle misurazioni periodiche nel tempo
            if "t2_storico_misure" not in st.session_state:
                if os.path.exists(path_storico_pers):
                    try:
                        df_loaded = pd.read_csv(path_storico_pers)
                        df_loaded["Data_Monitoraggio"] = pd.to_datetime(df_loaded["Data_Monitoraggio"]).dt.date
                        st.session_state.t2_storico_misure = df_loaded.fillna("")
                    except Exception:
                        st.session_state.t2_storico_misure = pd.DataFrame()
                else:
                    st.session_state.t2_storico_misure = pd.DataFrame([
                        {
                            "Data_Monitoraggio": oggi - timedelta(days=7),
                            "ID_KPI": "NM-01",
                            "Nome_KPI": "Numero Totale Near Miss Segnalati",
                            "Valore_Registrato": 12.0,
                            "Dettaglio_Fattori": "Segnalazioni valide: 12",
                        },
                        {
                            "Data_Monitoraggio": oggi,
                            "ID_KPI": "NM-01",
                            "Nome_KPI": "Numero Totale Near Miss Segnalati",
                            "Valore_Registrato": 14.0,
                            "Dettaglio_Fattori": "Segnalazioni valide: 14",
                        },
                        {
                            "Data_Monitoraggio": oggi - timedelta(days=7),
                            "ID_KPI": "NM-02",
                            "Nome_KPI": "Tasso di Risoluzione Near Miss",
                            "Valore_Registrato": 70.0,
                            "Dettaglio_Fattori": "Near Miss Chiusi: 35.0 / Near Miss Totali: 50.0",
                        },
                        {
                            "Data_Monitoraggio": oggi,
                            "ID_KPI": "NM-02",
                            "Nome_KPI": "Tasso di Risoluzione Near Miss",
                            "Valore_Registrato": 78.5,
                            "Dettaglio_Fattori": "Near Miss Chiusi: 47.0 / Near Miss Totali: 60.0",
                        },
                    ])
                    salva_su_github_kpi(path_storico_pers, st.session_state.t2_storico_misure, "Inizializzazione storico_misure.csv")

            # --- NAVIGAZIONE PRINCIPALE STABILE (PERSISTENTE) ---
            if "kpi_main_nav" not in st.session_state:
                st.session_state.kpi_main_nav = "📊 Definizione Master KPI"
            st.session_state.kpi_main_nav = st.radio(
                "Seleziona Sezione Principale KPI",
                ["📊 Definizione Master KPI", "⚙️ Strumento di Monitoraggio & Storico"],
                horizontal=True,
                key="radio_kpi_main_nav_selector"
            )
            st.markdown("---")

            # ==========================================
            # TAB 1: DEFINIZIONE MASTER KPI
            # ==========================================
            if st.session_state.kpi_main_nav == "📊 Definizione Master KPI":
                st.subheader("Gestione Master dei KPI, Formule e Parametri")
                st.markdown(
                    "I KPI inseriti o modificati qui vengono riconosciuti e sincronizzati"
                    " automaticamente nella `tab_kpi_2`."
                )
                st.markdown("---")
                edited_t1_kpi = st.data_editor(
                    st.session_state.t1_kpi_definitions,
                    num_rows="dynamic",
                    key="t1_kpi_editor",
                    use_container_width=True,
                )
                if st.button("💾 Salva Modifiche Master (Tab 1)", key="t1_save_btn"):
                    st.session_state.t1_kpi_definitions = edited_t1_kpi.copy()
                    salva_su_github_kpi(path_master_pers, st.session_state.t1_kpi_definitions, "Aggiornamento Master KPI Definizioni")
                    st.success("Definizioni e parametri sincronizzati con successo sia in locale che su GitHub!")
                    st.rerun()
                st.markdown("---")
                
                # Sezione Download e Salvataggio Tab 1
                st.markdown("##### 📥 Esportazione Tabella Master KPI")
                nome_file_t1 = f"{data_str}_Tab1-MasterKPI.csv"
                path_t1 = os.path.join("KPI", nome_file_t1)
                csv_t1 = st.session_state.t1_kpi_definitions.to_csv(index=False, encoding='utf-8').encode('utf-8')
                col_exp_1, col_exp_2 = st.columns(2)
                with col_exp_1:
                    if st.button("Salva in cartella 'KPI' e GitHub (Tab 1)", key="btn_save_folder_t1", use_container_width=True):
                        salva_su_github_kpi(path_t1, st.session_state.t1_kpi_definitions, f"Salvataggio {nome_file_t1}")
                        st.success(f"Salvato con successo in: {path_t1} e nel repository GitHub!")
                with col_exp_2:
                    st.download_button(
                        label="Scarica CSV Master KPI",
                        data=csv_t1,
                        file_name=nome_file_t1,
                        mime="text/csv",
                        key="dl_btn_t1",
                        use_container_width=True
                    )
                st.markdown("---")
                
                # --- ELIMINAZIONE PROTETTA MASTER KPI ---
                st.markdown("##### Eliminazione Master KPI (Richiede Password)")
                with st.form("form_elimina_master_kpi"):
                    kpi_list_names = st.session_state.t1_kpi_definitions["Nome_KPI"].tolist() if not st.session_state.t1_kpi_definitions.empty else []
                    kpi_da_eliminare = st.selectbox("Seleziona KPI da eliminare permanentemente", kpi_list_names, key="sel_kpi_del")
                    pwd_del_kpi = st.text_input("Inserisci la password di sezione per confermare l'eliminazione", type="password", key="pwd_del_kpi_field")
                    
                    if st.form_submit_button("Conferma ed Elimina Master KPI", use_container_width=True):
                        if pwd_del_kpi == correct_pwd_kpi:
                            if kpi_da_eliminare:
                                # Rimuovi da Master KPI
                                st.session_state.t1_kpi_definitions = st.session_state.t1_kpi_definitions[
                                    st.session_state.t1_kpi_definitions["Nome_KPI"] != kpi_da_eliminare
                                ]
                                # Rimuovi anche dallo storico collegato
                                st.session_state.t2_storico_misure = st.session_state.t2_storico_misure[
                                    st.session_state.t2_storico_misure["Nome_KPI"] != kpi_da_eliminare
                                ]
                                # Aggiorna i file sia in locale che su GitHub
                                salva_su_github_kpi(path_master_pers, st.session_state.t1_kpi_definitions, f"Eliminazione Master KPI: {kpi_da_eliminare}")
                                salva_su_github_kpi(path_storico_pers, st.session_state.t2_storico_misure, f"Pulizia storico per eliminazione KPI: {kpi_da_eliminare}")
                                st.success(f"KPI '{kpi_da_eliminare}' eliminato con successo!")
                                st.rerun()
                        else:
                            st.error("Password errata. Impossibile procedere con l'eliminazione.")
                st.markdown("---")
                st.metric(
                    "KPI Attivi nel Sistema", len(st.session_state.t1_kpi_definitions)
                )

            # ==========================================
            # TAB 2: STRUMENTO DI MONITORAGGIO E STORICO
            # ==========================================
            else:
                st.subheader(
                    "Strumento di Monitoraggio (Riconoscimento Automatico KPI da Tab 1)"
                )
                st.markdown(
                    "Qui vengono letti gli stessi KPI definiti nella Tab 1, consentendo di"
                    " inserire i fattori di calcolo e registrare nuove misurazioni nel"
                    " tempo."
                )
                st.markdown("---")
                # Acquisizione diretta degli stessi KPI della Tab 1
                df_kpi_master = st.session_state.t1_kpi_definitions.copy()
                def calcola_prossima_data_approssimata(row):
                    storico_kpi = st.session_state.t2_storico_misure[
                        st.session_state.t2_storico_misure["ID_KPI"] == row["ID_KPI"]
                    ]
                    if not storico_kpi.empty:
                        ultima_data = pd.to_datetime(
                            storico_kpi["Data_Monitoraggio"].max()
                        ).date()
                        cad = str(row["Cadenza"])
                        if "Settimanale" in cad:
                            delta_gg = 7
                        elif "Bisettimanale" in cad:
                            delta_gg = 14
                        elif "Mensile" in cad:
                            delta_gg = 30
                        elif "Trimestrale" in cad:
                            delta_gg = 90
                        elif "Semestrale" in cad:
                            delta_gg = 180
                        else:
                            delta_gg = 365
                        return ultima_data + timedelta(days=delta_gg)
                    else:
                        return oggi + timedelta(days=7)
                    
                if not df_kpi_master.empty:
                    df_kpi_master["Prossima_Data"] = df_kpi_master.apply(
                        calcola_prossima_data_approssimata, axis=1
                    )
                    def calcola_stato_t2(row):
                        delta = (row["Prossima_Data"] - oggi).days
                        if delta < 0:
                            return "Scaduto", "🔴"
                        elif delta <= 3:
                            return "In Scadenza", "🟡"
                        else:
                            return "Regolare", "🟢"
                    df_kpi_master["Stato_Testo"], df_kpi_master["Stato_Colore"] = zip(
                        *df_kpi_master.apply(calcola_stato_t2, axis=1)
                    )
                    df_kpi_master["Giorni_Rimasti"] = df_kpi_master["Prossima_Data"].apply(
                        lambda x: (x - oggi).days
                    )
                # --- SOTTO-NAVIGAZIONE STABILE PER IL MONITORAGGIO ---
                if "kpi_sub_nav" not in st.session_state:
                    st.session_state.kpi_sub_nav = "Dashboard & Grafici"
                st.session_state.kpi_sub_nav = st.radio(
                    "Seleziona Sottosezione Monitoraggio",
                    [
                        "Dashboard & Grafici",
                        "Registra Nuovo Monitoraggio",
                        "Aggiungi Nuovo KPI & Configurazione",
                    ],
                    horizontal=True,
                    key="radio_kpi_sub_nav_selector"
                )
                st.markdown("---")
                
                # --- SOTTO-TAB 1: MONITORAGGIO E GRAFICI ---
                if st.session_state.kpi_sub_nav == "Dashboard & Grafici":
                    if not df_kpi_master.empty:
                        scadenze_critiche_t2 = df_kpi_master[
                            df_kpi_master["Giorni_Rimasti"] <= 3
                        ]
                        if not scadenze_critiche_t2.empty:
                            st.warning(
                                f"⚠️ Attenzione: {len(scadenze_critiche_t2)} KPI richiedono un"
                                " monitoraggio imminente!"
                            )
                        else:
                            st.success("Tutti i KPI condivisi sono regolari.")
                        st.markdown("#### Tabella di Controllo KPI Sincronizzati")
                        def color_status_t2(val):
                            if val == "Scaduto":
                                return (
                                    "background-color: #ffcccc; color: #900c3f; font-weight: bold;"
                                )
                            elif val == "In Scadenza":
                                return (
                                    "background-color: #fff3cd; color: #856404; font-weight: bold;"
                                )
                            else:
                                return "background-color: #d4edda; color: #155724;"
                        df_controllo_display = df_kpi_master[[
                                "Stato_Colore",
                                "ID_KPI",
                                "Nome_KPI",
                                "Categoria",
                                "Unita_Misura",
                                "Cadenza",
                                "Prossima_Data",
                                "Formula",
                                "Baseline",
                                "Stato_Testo",
                            ]]
                        st.dataframe(
                            df_controllo_display
                            .style.map(color_status_t2, subset=["Stato_Testo"])
                            .format({"Prossima_Data": lambda t: t.strftime("%d/%m/%Y")}),
                            use_container_width=True,
                        )
                        # Esportazione Tabella di Controllo KPI (Sottosezione Tab2 - Controllo)
                        st.markdown("##### 📥 Esportazione Tabella di Controllo KPI")
                        nome_file_controllo = f"{data_str}_Tab2-ControlloKPI.csv"
                        path_controllo = os.path.join("KPI", nome_file_controllo)
                        csv_controllo = df_controllo_display.to_csv(index=False, encoding='utf-8').encode('utf-8')
                        col_c1, col_c2 = st.columns(2)
                        with col_c1:
                            if st.button("Salva in cartella 'KPI' e GitHub (Controllo KPI)", key="btn_save_folder_ctrl", use_container_width=True):
                                salva_su_github_kpi(path_controllo, df_controllo_display, f"Salvataggio {nome_file_controllo}")
                                st.success(f"Salvato con successo in: {path_controllo} e su GitHub!")
                        with col_c2:
                            st.download_button(
                                label="Scarica CSV Controllo KPI",
                                data=csv_controllo,
                                file_name=nome_file_controllo,
                                mime="text/csv",
                                key="dl_btn_ctrl",
                                use_container_width=True
                            )
                        st.markdown("---")
                        st.markdown("#### Evoluzione Storica dei KPI")
                        kpi_selezionato_storico = st.selectbox(
                            "Seleziona KPI per visualizzare il grafico temporale:",
                            df_kpi_master["Nome_KPI"],
                            key="t2_sel_storico_chart",
                        )
                        df_storico_filtrato = st.session_state.t2_storico_misure[
                            st.session_state.t2_storico_misure["Nome_KPI"]
                            == kpi_selezionato_storico
                        ]
                        if not df_storico_filtrato.empty:
                            df_chart = df_storico_filtrato.set_index("Data_Monitoraggio")[
                                ["Valore_Registrato"]
                            ]
                            st.line_chart(df_chart)
                        else:
                            st.info("Nessuno storico registrato per questo KPI.")
                    else:
                        st.info("Nessun KPI presente.")
                        
                # --- SOTTO-TAB 2: REGISTRA NUOVO MONITORAGGIO PERIODICO ---
                elif st.session_state.kpi_sub_nav == "Registra Nuovo Monitoraggio":
                    st.markdown(
                        "#### Aggiungi Nuovo Monitoraggio (Nuova Riga nello Storico)"
                    )
                    st.markdown(
                        "I KPI sottostanti sono riconosciuti direttamente da quelli definiti"
                        " nella `tab_kpi_1`."
                    )
                    if not df_kpi_master.empty:
                        kpi_sel_upd = st.selectbox(
                            "Seleziona KPI Riconosciuto",
                            df_kpi_master["Nome_KPI"],
                            key="t2_sel_kpi_storico",
                        )
                        kpi_row = df_kpi_master[
                            df_kpi_master["Nome_KPI"] == kpi_sel_upd
                        ].iloc[0]
                        formula_corrente = str(kpi_row["Formula"])
                        cadenza_corrente = str(kpi_row["Cadenza"])
                        id_kpi_corrente = str(kpi_row["ID_KPI"])
                        p1_nome = str(kpi_row["Parametro_1"])
                        p2_nome = str(kpi_row["Parametro_2"])
                        p3_nome = str(kpi_row["Parametro_3"])
                        st.info(
                            f"**Formula:** `{formula_corrente}` | **Baseline:**"
                            f" `{kpi_row['Baseline']}`"
                        )
                        col_f1, col_f2 = st.columns(2)
                        with col_f1:
                            valori_fattori = []
                            dettagli_lista = []
                            if p1_nome and p1_nome.lower() != "nan" and p1_nome.strip() != "":
                                v1 = st.number_input(
                                    f"Parametro 1: {p1_nome}",
                                    value=10.0,
                                    format="%.2f",
                                    key="t2_p1_val",
                                )
                                valori_fattori.append(v1)
                                dettagli_lista.append(f"{p1_nome}: {v1}")
                            if p2_nome and p2_nome.lower() != "nan" and p2_nome.strip() != "":
                                v2 = st.number_input(
                                    f"Parametro 2: {p2_nome}",
                                    value=15.0,
                                    format="%.2f",
                                    key="t2_p2_val",
                                )
                                valori_fattori.append(v2)
                                dettagli_lista.append(f"{p2_nome}: {v2}")
                            if p3_nome and p3_nome.lower() != "nan" and p3_nome.strip() != "":
                                v3 = st.number_input(
                                    f"Parametro 3: {p3_nome}",
                                    value=0.0,
                                    format="%.2f",
                                    key="t2_p3_val",
                                )
                                valori_fattori.append(v3)
                                dettagli_lista.append(f"{p3_nome}: {v3}")
                            if len(valori_fattori) >= 2 and valori_fattori[1] > 0:
                                if "%" in str(kpi_row["Unita_Misura"]):
                                    valore_calcolato = round(
                                        (valori_fattori[0] / valori_fattori[1]) * 100, 2
                                    )
                                else:
                                    valore_calcolato = round(
                                        valori_fattori[0] / valori_fattori[1], 2
                                    )
                            elif len(valori_fattori) == 1:
                                valore_calcolato = valori_fattori[0]
                            else:
                                valore_calcolato = st.number_input(
                                    "Valore Misurato", value=0.0, format="%.2f", key="t2_fallback_v"
                                )
                                dettagli_lista.append(f"Valore Assoluto: {valore_calcolato}")
                        dettaglio_fattori_str = " / ".join(dettagli_lista)
                        st.metric(
                            "Valore Risultante Calcolato",
                            f"{valore_calcolato} {kpi_row['Unita_Misura']}",
                        )
                        with col_f2:
                            data_monitoraggio = st.date_input(
                                "Data monitoraggio", value=oggi, key="t2_data_mon_input"
                            )
                            if "Settimanale" in cadenza_corrente:
                                gg_agg = 7
                            elif "Bisettimanale" in cadenza_corrente:
                                gg_agg = 14
                            elif "Mensile" in cadenza_corrente:
                                gg_agg = 30
                            elif "Trimestrale" in cadenza_corrente:
                                gg_agg = 90
                            elif "Semestre" in cadenza_corrente:
                                gg_agg = 180
                            else:
                                gg_agg = 365
                            prossima_scadenza_calc = data_monitoraggio + timedelta(days=gg_agg)
                            st.write(
                                f"📅 **Cadenza:** {cadenza_corrente} -> Nuova scadenza stimata:"
                                f" **{prossima_scadenza_calc.strftime('%d/%m/%Y')}**"
                            )
                        if st.button(
                            "Registra Nuovo Monitoraggio nello Storico",
                            key="t2_btn_registra_storico",
                        ):
                            nuova_riga_storico = pd.DataFrame([{
                                "Data_Monitoraggio": data_monitoraggio,
                                "ID_KPI": id_kpi_corrente,
                                "Nome_KPI": kpi_sel_upd,
                                "Valore_Registrato": valore_calcolato,
                                "Dettaglio_Fattori": dettaglio_fattori_str,
                            }])
                            st.session_state.t2_storico_misure = pd.concat(
                                [st.session_state.t2_storico_misure, nuova_riga_storico],
                                ignore_index=True,
                            )
                            # Salva permanentemente sia in locale che su GitHub
                            salva_su_github_kpi(path_storico_pers, st.session_state.t2_storico_misure, f"Nuovo monitoraggio per KPI {id_kpi_corrente}")
                            st.success(
                                "Nuovo monitoraggio aggiunto con successo nello storico e salvato su GitHub!"
                            )
                            st.rerun()
                        st.markdown("---")
                        st.markdown("#### Tabella Storico Monitoraggi")
                        df_storico_display = st.session_state.t2_storico_misure.sort_values(
                            by="Data_Monitoraggio", ascending=False
                        )
                        st.dataframe(df_storico_display, use_container_width=True)
                        
                        # Esportazione Tabella Storico Monitoraggi (Sottosezione Tab2 - Storico)
                        st.markdown("##### Esportazione Tabella Storico Monitoraggi")
                        nome_file_storico = f"{data_str}_Tab2-StoricoMisure.csv"
                        path_storico = os.path.join("KPI", nome_file_storico)
                        csv_storico = df_storico_display.to_csv(index=False, encoding='utf-8').encode('utf-8')
                        col_s1, col_s2 = st.columns(2)
                        with col_s1:
                            if st.button("Salva in cartella 'KPI' e GitHub (Storico Misure)", key="btn_save_folder_storico", use_container_width=True):
                                salva_su_github_kpi(path_storico, df_storico_display, f"Salvataggio {nome_file_storico}")
                                st.success(f"Salvato con successo in: {path_storico} e su GitHub!")
                        with col_s2:
                            st.download_button(
                                label="Scarica CSV Storico Misure",
                                data=csv_storico,
                                file_name=nome_file_storico,
                                mime="text/csv",
                                key="dl_btn_storico",
                                use_container_width=True
                            )
                            
                        st.markdown("---")
                        
                        # --- ELIMINAZIONE PROTETTA REGISTRO STORICO ---
                        st.markdown("##### Eliminazione Registro di Monitoraggio (Richiede Password)")
                        if not st.session_state.t2_storico_misure.empty:
                            storico_temp_del = st.session_state.t2_storico_misure.copy()
                            storico_temp_del["Etichetta_Riga"] = (
                                storico_temp_del["Data_Monitoraggio"].astype(str)
                                + " | "
                                + storico_temp_del["Nome_KPI"]
                                + " | Valore: "
                                + storico_temp_del["Valore_Registrato"].astype(str)
                            )
                            
                            with st.form("form_elimina_registro_storico"):
                                riga_selezionata_del = st.selectbox(
                                    "Seleziona il registro di monitoraggio da eliminare",
                                    storico_temp_del["Etichetta_Riga"].tolist(),
                                    key="sel_storico_del"
                                )
                                pwd_del_storico = st.text_input("Inserisci la password di sezione per confermare l'eliminazione", type="password", key="pwd_del_storico_field")
                                
                                if st.form_submit_button("Conferma ed Elimina Registro", use_container_width=True):
                                    if pwd_del_storico == correct_pwd_kpi:
                                        idx_da_rimuovere = storico_temp_del[
                                            storico_temp_del["Etichetta_Riga"] == riga_selezionata_del
                                        ].index
                                        st.session_state.t2_storico_misure = st.session_state.t2_storico_misure.drop(idx_da_rimuovere).reset_index(drop=True)
                                        # Aggiorna il file sia in locale che su GitHub
                                        salva_su_github_kpi(path_storico_pers, st.session_state.t2_storico_misure, "Eliminazione registro dallo storico misure")
                                        st.success("Registro di monitoraggio eliminato con successo!")
                                        st.rerun()
                                    else:
                                        st.error("Password errata. Impossibile procedere con l'eliminazione.")
                        else:
                            st.info("Nessun registro di monitoraggio disponibile per l'eliminazione.")
                    else:
                        st.info("Nessun KPI disponibile.")
                        
                # --- SOTTO-TAB 3: AGGIUNGI NUOVO KPI & CONFIGURAZIONE ---
                else:
                    st.markdown("#### Aggiungi un Nuovo KPI (Sincronizzato con Tab 1)")
                    st.markdown(
                        "I KPI inseriti qui saranno aggiunti all'archivio comune e visibili"
                        " in entrambe le tab."
                    )
                    with st.form("t2_form_nuovo_kpi"):
                        c_a1, c_a2 = st.columns(2)
                        with c_a1:
                            nid = st.text_input("ID KPI (es. NM-04)", value="NM-04", key="t2_nid")
                            nnome = st.text_input(
                                "Nome KPI", value="Indice Chiusura Prescrizioni", key="t2_nnome"
                            )
                            ncat = st.selectbox(
                                "Categoria",
                                [
                                    "Volume",
                                    "Efficacia",
                                    "Reattività",
                                    "Cultura HSE",
                                    "Prevenzione",
                                    "Ambiente",
                                ],
                                key="t2_ncat",
                            )
                            nmis = st.selectbox(
                                "Unità di Misura", ["%", "Numero", "Giorni"], key="t2_nmis"
                            )
                            ncad = st.selectbox(
                                "Cadenza",
                                ["Settimanale", "Bisettimanale", "Mensile", "Trimestrale", "Semestrale", "Annuale"],
                                key="t2_ncad",
                            )
                        with c_a2:
                            nform = st.text_input(
                                "Formula (es. (Parametro_1 / Parametro_2) * 100)",
                                value="(Parametro_1 / Parametro_2) * 100",
                                key="t2_nform",
                            )
                            np1 = st.text_input(
                                "Parametro 1 (Fattore)",
                                value="Valore Numeratore",
                                key="t2_np1",
                            )
                            np2 = st.text_input(
                                "Parametro 2 (Fattore)",
                                value="Valore Denominatore",
                                key="t2_np2",
                            )
                            np3 = st.text_input(
                                "Parametro 3 (Opzionale)", value="", key="t2_np3"
                            )
                        # Gestione Baseline (Sì o No con Text Area dedicata)
                        nbase_choice = st.selectbox(
                            "Baseline Esistente", ["No", "Sì"], key="t2_nbase_choice"
                        )
                        nbaseline_text = st.text_area(
                            "Valore o Descrizione della Baseline",
                            value=(
                                "Inserisci qui i dettagli della baseline o target iniziale..."
                                if nbase_choice == "Sì"
                                else "Nessuna baseline storica precedente."
                            ),
                            key="t2_nbaseline_textarea",
                        )
                        if st.form_submit_button("Crea e Sincronizza Nuovo KPI"):
                            if nid in st.session_state.t1_kpi_definitions["ID_KPI"].values:
                                st.error(f"L'ID KPI '{nid}' esiste già.")
                            else:
                                nuova_riga_kpi = pd.DataFrame([{
                                    "ID_KPI": nid,
                                    "Nome_KPI": nnome,
                                    "Categoria": ncat,
                                    "Unita_Misura": nmis,
                                    "Cadenza": ncad,
                                    "Formula": nform,
                                    "Parametro_1": np1,
                                    "Parametro_2": np2,
                                    "Parametro_3": np3,
                                    "Baseline": nbaseline_text,
                                }])
                                st.session_state.t1_kpi_definitions = pd.concat(
                                    [st.session_state.t1_kpi_definitions, nuova_riga_kpi],
                                    ignore_index=True,
                                )
                                # Salva permanentemente sia in locale che su GitHub
                                salva_su_github_kpi(path_master_pers, st.session_state.t1_kpi_definitions, f"Creazione nuovo KPI: {nid}")
                                st.success(
                                    "Nuovo KPI creato e riconosciuto correttamente sia in locale che su GitHub!"
                                )
                                st.rerun()
        render_sezione_kpi()
# ==========================================
# SEZIONE 8: PIANO DI MIGLIORAMENTO
# ==========================================
if nav == "Piano Miglioramento":
    st.title("Piano di Miglioramento")

    # Inizializzazione Session State Autenticazione
    if "authenticated_sec8" not in st.session_state:
        st.session_state["authenticated_sec8"] = False

    # Definizione tab
    tab_list = ["Documentazione di riferimento", "Valutazione del Rischio", "Azioni Piano di Miglioramento"]
    tabs = st.tabs(tab_list)

    # ------------------------------------------
    # SOTTOSEZIONE 1: Documentazione di riferimento (PUBBLICA)
    # ------------------------------------------
    with tabs[0]:
        st.subheader("Documentazione di Riferimento")
        pdf_path = os.path.join("Piano_Miglioramento", "Istruzione_Piano_Miglioramento.pdf")
        
        if os.path.exists(pdf_path):
            with open(pdf_path, "rb") as f:
                pdf_data = f.read()
            
            st.download_button(
                label="Scarica Istruzione Piano Miglioramento (PDF)",
                data=pdf_data,
                file_name="Istruzione_Piano_Miglioramento.pdf",
                mime="application/pdf",
                key="btn_download_pdf_sec8"
            )
            
            # Rendering PDF via iframe HTML
            import base64
            base64_pdf = base64.b64encode(pdf_data).decode('utf-8')
            pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="600" type="application/pdf"></iframe>'
            st.markdown(pdf_display, unsafe_allow_html=True)
        else:
            st.warning(f"Il file '{pdf_path}' non è stato trovato sul server.")

    # ------------------------------------------
    # GESTIONE AUTENTICAZIONE DINAMICA PER SEZIONI PRIVATE
    # ------------------------------------------
    def render_auth_form(section_prefix="sec8"):
        if not st.session_state["authenticated_sec8"]:
            st.warning("🔒 Quest'area è riservata. Inserisci la password per accedere.")
            pwd_input = st.text_input("Password di accesso", type="password", key=f"pwd_{section_prefix}")
            if st.button("Sblocca Sezioni Private", key=f"btn_auth_{section_prefix}"):
                correct_pwd = st.secrets.get("PASSWORD_SEZIONE", "admin")
                if pwd_input == correct_pwd:
                    st.session_state["authenticated_sec8"] = True
                    st.rerun()
                else:
                    st.error("Password errata.")
            return False
        return True

    # UTILITY PER CARICAMENTO EVENTI (Inclusione sicura di manutenzione.csv)
    def load_events():
        events = []
        file_sources = [
            ("segnalazioni_near_miss.csv", "segnalazione"),
            ("analisi_near_miss.csv", "analisi"),
            (os.path.join("Segnalazione_NM_Manutenzione", "manutenzione.csv"), "manutenzione"),
            (os.path.join("Segnalazioni_NM_Manutenzione", "manutenzione.csv"), "manutenzione")
        ]
        
        for filepath, tipo in file_sources:
            if os.path.exists(filepath):
                try:
                    df = pd.read_csv(filepath)
                    if not df.empty:
                        for val in df.iloc[:, 0].dropna().astype(str).tolist():
                            events.append(f"{val} ({tipo})")
                except Exception:
                    pass
        return sorted(list(set(events))) if events else ["Nessun evento disponibile"]

    # UTILITY PER ACCORCIARE IL NOME DELL'EVENTO
    def format_event_name_short(evento_str):
        import re
        tipo = "evento"
        if "(segnalazione)" in evento_str.lower():
            tipo = "segnalazione"
        elif "(analisi)" in evento_str.lower():
            tipo = "analisi"
        elif "(manutenzione)" in evento_str.lower():
            tipo = "manutenzione"
        
        match_date = re.search(r'(\d{4}[-/.]?\d{2}[-/.]?\d{2}|\d{2}[-/.]?\d{2}[-/.]?\d{4})', evento_str)
        if match_date:
            clean_date = re.sub(r'[-/.]', '', match_date.group(1))
        else:
            clean_date = datetime.now().strftime("%Y%m%d")
            
        return f"{clean_date}_{tipo}"

    # UTILITY PER LEGGERE DATI VR DALLA CARTELLA SE NON IN SESSIONE
    def get_existing_vr_data(short_event_prefix):
        vr_folder = os.path.join("Piano_Miglioramento", "Valutazione_Rischio")
        if os.path.exists(vr_folder):
            try:
                files = [f for f in os.listdir(vr_folder) if f.startswith(short_event_prefix) and f.endswith(".xlsx")]
                if files:
                    files.sort(reverse=True)  # Prende il file più recente
                    latest_file = os.path.join(vr_folder, files[0])
                    return pd.read_excel(latest_file, sheet_name="Valutazione Rischio")
            except Exception:
                pass
        return pd.DataFrame(columns=["Rischio", "Probabilità", "Gravità", "Sensibilità", "Controllo", "Significatività"])

    # ------------------------------------------
    # SOTTOSEZIONE 2: Valutazione del Rischio (PRIVATA)
    # ------------------------------------------
    with tabs[1]:
        if render_auth_form(section_prefix="vr"):
            st.subheader("Valutazione del Rischio")
            
            lista_eventi = load_events()
            evento_selezionato_vr = st.selectbox("Seleziona Evento Near Miss", lista_eventi, key="sb_vr")
            
            default_vr_data = pd.DataFrame([
                {"Rischio": "", "Probabilità": 1, "Gravità": 1, "Sensibilità": 1, "Controllo": 1}
            ])
            
            st.write("Compila la tabella di Valutazione del Rischio:")
            edited_vr_df = st.data_editor(
                default_vr_data,
                num_rows="dynamic",
                column_config={
                    "Rischio": st.column_config.TextColumn("Rischio", required=True),
                    "Probabilità": st.column_config.NumberColumn("Probabilità", min_value=1, max_value=3, step=1),
                    "Gravità": st.column_config.NumberColumn("Gravità", min_value=1, max_value=3, step=1),
                    "Sensibilità": st.column_config.NumberColumn("Sensibilità", min_value=1, max_value=3, step=1),
                    "Controllo": st.column_config.NumberColumn("Controllo", min_value=1, max_value=3, step=1),
                },
                key="editor_vr"
            )
            
            def calcola_significativita(df):
                df_calc = df.copy()
                if df_calc.empty:
                    df_calc["Significatività"] = []
                    return df_calc
                
                for col in ["Gravità", "Probabilità", "Sensibilità", "Controllo"]:
                    if col in df_calc.columns:
                        df_calc[col] = pd.to_numeric(df_calc[col], errors="coerce").fillna(1)
                
                df_calc["Controllo"] = df_calc["Controllo"].replace(0, 1)
                df_calc["Significatività"] = (
                    (df_calc["Gravità"] * df_calc["Probabilità"] * df_calc["Sensibilità"]) / df_calc["Controllo"]
                ).round(2)
                return df_calc

            df_vr_calcolato = calcola_significativita(edited_vr_df)
            
            def color_rows(val):
                try:
                    if pd.isna(val):
                        return ''
                    val = float(val)
                    if 1 <= val <= 3:
                        return 'background-color: #d4edda; color: #155724;'
                    elif 4 <= val <= 6:
                        return 'background-color: #fff3cd; color: #856404;'
                    elif 7 <= val <= 9:
                        return 'background-color: #f8d7da; color: #721c24;'
                except Exception:
                    pass
                return ''

            styled_df = df_vr_calcolato.style.map(color_rows, subset=['Significatività'])
            st.dataframe(styled_df, use_container_width=True)
            
            # Generazione nome file corto e info evento
            now_str = datetime.now().strftime("%Y%m%d_%H%M%S")
            short_event_vr = format_event_name_short(evento_selezionato_vr)
            excel_filename_vr = f"{short_event_vr}_{now_str}_Valutazione_Rischio.xlsx"
            
            df_info_evento = pd.DataFrame([{
                "Evento Collegato": evento_selezionato_vr,
                "Data Creazione": datetime.now().strftime("%d-%m-%Y %H:%M:%S")
            }])
            
            buffer_vr = io.BytesIO()
            with pd.ExcelWriter(buffer_vr, engine='openpyxl') as writer:
                df_info_evento.to_excel(writer, index=False, sheet_name="Info Evento")
                df_vr_calcolato.to_excel(writer, index=False, sheet_name="Valutazione Rischio")
            excel_data_vr = buffer_vr.getvalue()
            
            col_btn1, col_btn2 = st.columns(2)
            with col_btn1:
                if st.button("💾 Salva Online", key="btn_gh_vr"):
                    repo_path = f"Piano_Miglioramento/Valutazione_Rischio/{excel_filename_vr}"
                    if save_to_github(repo_path, excel_data_vr, f"Add {excel_filename_vr}"):
                        st.success(f"File salvato con successo su GitHub in: {repo_path}")
            
            with col_btn2:
                st.download_button(
                    label= "Scarica Tabella Dinamica",
                    data=excel_data_vr,
                    file_name=excel_filename_vr,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="btn_dl_excel_vr"
                )

    # ------------------------------------------
    # SOTTOSEZIONE 3: Azioni Piano di Miglioramento (PRIVATA)
    # ------------------------------------------
    with tabs[2]:
        if render_auth_form(section_prefix="am"):
            st.subheader("Azioni Piano di Miglioramento")
            
            lista_eventi = load_events()
            evento_selezionato_am = st.selectbox("Seleziona Evento Collegato", lista_eventi, key="sb_am")
            
            st.markdown("### Azioni Intraprese")
            azioni_immediate = st.text_area("Azioni immediate", height=120, key="ta_azioni_immediate")
            
            st.write("Azioni di miglioramento (correttive o preventive) - Tipologia di intervento")
            tipologie = [
                "Tecnico", "Formazione/Addestramento", "Informazione/Comunicazione/Partecipazione",
                "Definizione/revisione delle procedure e istruzioni lavorative",
                "Verifica applicazione procedure, istruzioni, comportamenti", "Altro"
            ]
            
            default_tipo_df = pd.DataFrame([{"Tipologia d'azione": tipologie[0], "Descrizione": ""}])
            edited_tipo_df = st.data_editor(
                default_tipo_df,
                num_rows="dynamic",
                column_config={
                    "Tipologia d'azione": st.column_config.SelectboxColumn("Tipologia d'azione", options=tipologie, required=True),
                    "Descrizione": st.column_config.TextColumn("Descrizione / Testo libero", required=True)
                },
                key="editor_tipologie"
            )
            
            st.markdown("### Follow up Azioni Intraprese")
            today_date = datetime.today().date()

            default_followup_df = pd.DataFrame([{
                "Azioni di miglioramento (correttive o preventive)": "",
                "Responsabile attuazione": "",
                "Accountable Attuazione": "",
                "Entro il": today_date,
                "Firma presa in carico": "",
                "Data attuazione": today_date,
                "Verifica attuazione Data": today_date,
                "Verifica attuazione Firma": ""
            }])
            
            edited_followup_df = st.data_editor(
                default_followup_df,
                num_rows="dynamic",
                column_config={
                    "Entro il": st.column_config.DateColumn("Entro il", format="DD-MM-YYYY"),
                    "Data attuazione": st.column_config.DateColumn("Data attuazione", format="DD-MM-YYYY"),
                    "Verifica attuazione Data": st.column_config.DateColumn("Verifica attuazione Data", format="DD-MM-YYYY")
                },
                key="editor_followup"
            )
            
            # Formattazione timestamp e nomi file
            now_str = datetime.now().strftime("%Y%m%d_%H%M%S")
            short_event_am = format_event_name_short(evento_selezionato_am)
            
            csv_filename = f"{short_event_am}_{now_str}_Piano_Miglioramento.csv"
            report_excel_filename = f"{short_event_am}_{now_str}_Report.xlsx"

            # Buffer CSV interno
            csv_buffer = io.StringIO()
            csv_buffer.write(f"# Evento Collegato: {evento_selezionato_am}\n")
            csv_buffer.write(f"# Azioni Immediate:\n{azioni_immediate}\n\n")
            csv_buffer.write("# Tipologie Intervento:\n")
            edited_tipo_df.to_csv(csv_buffer, index=False)
            csv_buffer.write("\n# Follow Up:\n")
            edited_followup_df.to_csv(csv_buffer, index=False)
            
            # --- RECUPERO DATI VALUTAZIONE RISCHIO PER LO STESSO EVENTO ---
            df_vr_report = st.session_state.get("editor_vr", pd.DataFrame())
            if isinstance(df_vr_report, pd.DataFrame) and not df_vr_report.empty and df_vr_report.iloc[0].get("Rischio", "") != "":
                df_vr_report = calcola_significativita(df_vr_report)
            else:
                # Se vuoto o non compilato in sessione, cerca file salvato nella cartella Valutazione_Rischio
                df_vr_report = get_existing_vr_data(short_event_am)

            # Dataframe Info Evento
            df_info_evento = pd.DataFrame([{
                "Evento Collegato": evento_selezionato_am,
                "Data Generazione Report": datetime.now().strftime("%d-%m-%Y %H:%M:%S")
            }])

            df_imm_report = pd.DataFrame([{"Azioni Immediate": azioni_immediate}])

            # --- GENERAZIONE EXCEL REPORT MULTI-FOGLIO ---
            report_buffer = io.BytesIO()
            with pd.ExcelWriter(report_buffer, engine='openpyxl') as writer:
                df_info_evento.to_excel(writer, index=False, sheet_name="Info Evento")
                df_vr_report.to_excel(writer, index=False, sheet_name="Valutazione del rischio")
                df_imm_report.to_excel(writer, index=False, sheet_name="Azioni immediate")
                edited_tipo_df.to_excel(writer, index=False, sheet_name="Azioni di miglioramento")
                edited_followup_df.to_excel(writer, index=False, sheet_name="Follow up azioni di miglioramento")
                
            report_excel_data = report_buffer.getvalue()
            
            st.markdown("---")
            col_rep1, col_rep2 = st.columns(2)
            with col_rep1:
                if st.button("Salva Report Online", key="btn_save_report"):
                    csv_repo_path = f"Piano_Miglioramento/Azioni_Piano_Miglioramento/{csv_filename}"
                    report_repo_path = f"Piano_Miglioramento/Report/{report_excel_filename}"
                    
                    saved_csv = save_to_github(csv_repo_path, csv_buffer.getvalue().encode('utf-8'), f"Add {csv_filename}")
                    saved_excel = save_to_github(report_repo_path, report_excel_data, f"Add {report_excel_filename}")
                    
                    if saved_csv and saved_excel:
                        st.success("Report completo e dati del Piano di Miglioramento salvati su GitHub!")

            with col_rep2:
                st.download_button(
                    label="Scarica Report (Excel)",
                    data=report_excel_data,
                    file_name=report_excel_filename,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="btn_dl_report_excel"
                )
# ==================================================================
# --- SEZIONE 9: Stima Costo Economico ---
# ==================================================================
if nav == "Stima Costo Economico":
    st.header("Stima Costo Economico del Near Miss")
    
    # Helper function per formattazione in valuta italiana (es. 1.234,56 €)
    def format_euro(valore):
        try:
            return f"{valore:,.2f} €".replace(",", "X").replace(".", ",").replace("X", ".")
        except Exception:
            return "0,00 €"

    # Helper function per formattare numeri float con virgola per export CSV
    def format_csv_number(valore):
        try:
            return f"{float(valore):.2f}".replace(".", ",")
        except Exception:
            return "0,00"

    # Gestione autenticazione per la Sezione 9 tramite Streamlit Secrets
    if "auth_stima_economico" not in st.session_state:
        st.session_state.auth_stima_economico = False
        
    if not st.session_state.auth_stima_economico:
        st.markdown("🔒 Inserisci la password per accedere all'area di stima del costo economico.")
        pwd_sec9 = st.text_input("Password Sezione 9", type="password", key="pwd_sec9_input")
        if st.button("Verifica Password", use_container_width=True, key="btn_verify_pwd_sec9"):
            correct_pwd = st.secrets.get("PASSWORD_SEZIONE", "hse2026")
            if pwd_sec9 == correct_pwd:
                st.session_state.auth_stima_economico = True
                st.success("Accesso autorizzato!")
                st.rerun()
            else:
                st.error("Password errata.")
    
    if st.session_state.auth_stima_economico:
        # Sottosezioni della Sezione 9
        sotto_sec_9 = st.radio(
            "Seleziona Sottosezione", 
            ["Documentazione di Riferimento", "Calcolo economico NM"], 
            horizontal=True, 
            key="radio_sotto_sec_9"
        )
        
        # --------------------------------------------------
        # SOTTOSEZIONE: Documentazione di Riferimento
        # --------------------------------------------------
        if sotto_sec_9 == "Documentazione di Riferimento":
            st.subheader("Consultazione Documento Stima Economica")
            st.markdown("Consulta o scarica il documento PDF ufficiale relativo alla stima economica del near miss.")
            
            try:
                base_dir = os.path.dirname(os.path.abspath(__file__))
            except NameError:
                base_dir = os.getcwd()
                
            file_pdf_path = os.path.join(base_dir, "Stima_Economica", "Stima Economica del Near Miss.pdf")
            
            if not os.path.exists(file_pdf_path):
                percorsi_alternativi = [
                    os.path.join("Stima_Economica", "Stima Economica del Near Miss.pdf"),
                    os.path.join("stima_economica", "Stima Economica del Near Miss.pdf"),
                    os.path.join(base_dir, "stima_economica", "Stima Economica del Near Miss.pdf")
                ]
                for p in percorsi_alternativi:
                    if os.path.exists(p):
                        file_pdf_path = os.path.abspath(p)
                        break
            
            if os.path.exists(file_pdf_path):
                with open(file_pdf_path, "rb") as f:
                    pdf_bytes = f.read()
                
                st.download_button(
                    label="📥 Scarica / Apri Documento Stima Economica (.pdf)",
                    data=pdf_bytes,
                    file_name="Stima Economica del Near Miss.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
                
                try:
                    import base64
                    base64_pdf = base64.b64encode(pdf_bytes).decode('utf-8')
                    pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="700px" type="application/pdf"></iframe>'
                    st.markdown(pdf_display, unsafe_allow_html=True)
                except Exception as e:
                    st.info("Utilizza il pulsante di download sopra per consultare il documento nel lettore PDF del tuo computer.")
            else:
                st.error("Il file PDF 'Stima Economica del Near Miss.pdf' non è stato trovato nella cartella 'Stima_Economica'.")
                
        # --------------------------------------------------
        # SOTTOSEZIONE: Calcolo Economico NM
        # --------------------------------------------------
        elif sotto_sec_9 == "Calcolo economico NM":
            st.subheader("Calcolo economico NM - Tabella Dinamica e Parametri")
            st.markdown("Configura i parametri di inquadramento (Manodopera) e il fatturato dell'anno precedente (Vendite e Reputazione) per i calcoli automatici.")
            
            opzioni = ["Nessuna (Nuova analisi)"]
                                            
            # Leggi Near Miss
            if 'FILE_NEAR_MISS' in globals() and os.path.exists(FILE_NEAR_MISS):
                try:
                    df_nm = pd.read_csv(FILE_NEAR_MISS, sep=";")
                    for idx, r in df_nm.iterrows():
                        opzioni.append(f"NM | {r.get('Data Segnalazione', 'N/D')} | {r.get('Tipo Evento', 'Evento')}")
                except Exception:
                    pass
                                            
            # Leggi Analisi già fatte
            if 'FILE_ANALISI_NM' in globals() and os.path.exists(FILE_ANALISI_NM):
                try:
                    df_an = pd.read_csv(FILE_ANALISI_NM, sep=";")
                    for idx, r in df_an.iterrows():
                        opzioni.append(f"AN | {r.get('Data Analisi', 'N/D')} | Collegamento: {r.get('Segnalazione Collegata', 'Analisi')}")
                except Exception:
                    pass
                                            
            scelta_rif = st.selectbox("Seleziona evento/analisi collegata:", opzioni)
            
            # Parametri specifici per la gestione della Manodopera
            st.markdown("#### Configurazione Condizioni Area Manodopera")
            col_m1, col_m2, col_m3 = st.columns(3)
            with col_m1:
                inquadramento_mansione = st.selectbox(
                    "Inquadramento", 
                    ["Q", "AS", "A", "B1", "B2S", "B2", "C1S", "C1", "C2", "C3", "D1", "D2", "E"], 
                    key="inquadramento_sel"
                )
            with col_m2:
                inf_1gg = st.selectbox("Infortunio sul lavoro 1gg", ["No", "Sì"], key="inf_1gg_sel")
                inf_2_4gg = st.selectbox("Infortunio sul lavoro 2<4gg", ["No", "Sì"], key="inf_2_4gg_sel")
            with col_m3:
                inf_5gg = st.selectbox("Infortunio sul lavoro >=5gg", ["No", "Sì"], key="inf_5gg_sel")
            
            inquadramento_mapping = {
                "Q": 2882.91, "AS": 2873.55, "A": 2521.80, "B1": 2292.49, "B2S": 2234.61, "B2": 2159.89,
                "C1S": 2034.87, "C1": 1960.11, "C2": 1826.94, "C3": 1732.11, "D1": 1656.25, "D2": 1561.08, "E": 1456.61
            }
            valore_inquadramento = inquadramento_mapping.get(inquadramento_mansione, 2159.89)
            
            perc_indennita_val = 0.0
            ral_mansione_calc = 0.0
            
            if inf_1gg == "Sì":
                perc_indennita_val = 100.0
                ral_mansione_calc = (valore_inquadramento / 30.0) * 3.0 * (perc_indennita_val / 100.0)
            elif inf_2_4gg == "Sì":
                perc_indennita_val = 40.0
                ral_mansione_calc = ((valore_inquadramento / 30.0) * 3.0 * (perc_indennita_val / 100.0)) + (valore_inquadramento / 30.0)
            elif inf_5gg == "Sì":
                perc_indennita_val = 25.0
                ral_mansione_calc = ((valore_inquadramento / 30.0) * 3.0 * (perc_indennita_val / 100.0)) + (valore_inquadramento / 30.0)
            else:
                ral_mansione_calc = valore_inquadramento

            # Parametri specifici per Vendite e Reputazione
            st.markdown("#### Configurazione Condizioni Aree Vendite e Reputazione")
            col_v1, col_v2 = st.columns(2)
            with col_v1:
                fatturato_vendite = st.number_input("Fatturato anno precedente (Vendite) [€]", value=1000000.0, step=10000.0, key="fat_vendite_input")
                vendite_1pct = st.selectbox("Diminuzione vendite 1% (Vendite)", ["No", "Sì"], key="vendite_1_sel")
                vendite_5pct = st.selectbox("Diminuzione vendite 5% (Vendite)", ["No", "Sì"], key="vendite_5_sel")
            with col_v2:
                fatturato_reputazione = st.number_input("Fatturato anno precedente (Reputazione) [€]", value=1000000.0, step=10000.0, key="fat_reputazione_input")
                rep_10pct = st.selectbox("Diminuzione fatturato 10% (Reputazione)", ["No", "Sì"], key="rep_10_sel")
                rep_15pct = st.selectbox("Diminuzione fatturato 15% (Reputazione)", ["No", "Sì"], key="rep_15_sel")

            # Calcoli condizionali Vendite e Reputazione
            val_vendite_1 = fatturato_vendite * (1.0 / 100.0) if vendite_1pct == "Sì" else 0.0
            val_vendite_5 = fatturato_vendite * (5.0 / 100.0) if vendite_5pct == "Sì" else 0.0
            val_rep_10 = fatturato_reputazione * (10.0 / 100.0) if rep_10pct == "Sì" else 0.0
            val_rep_15 = fatturato_reputazione * (15.0 / 100.0) if rep_15pct == "Sì" else 0.0

            # Inizializzazione DataFrame Tabella Dinamica in session_state
            if "df_calcolo_economico_nm" not in st.session_state:
                data_nm = [
                    # Macchinari
                    ["Macchinari", "Ricambio attrezzatura", 0.0],
                    ["Macchinari", "Ricambio pezzo macchinario", 0.0],
                    ["Macchinari", "Cambio attrezzatura", 0.0],
                    ["Macchinari", "Cambio Macchinario", 0.0],
                    # Manodopera
                    ["Manodopera", "Infortunio sul lavoro 1gg", 0.0],
                    ["Manodopera", "Infortunio sul lavoro 2<4gg", 0.0],
                    ["Manodopera", "Infortunio sul lavoro >=5gg", 0.0],
                    ["Manodopera", "Percentuale di indennità", 0.0],
                    ["Manodopera", "RAL per mansione", 0.0],
                    ["Manodopera", "Formazione neoassunto", 0.0],
                    ["Manodopera", "Addestramento personale", 0.0],
                    # Materiali
                    ["Materiali", "Carta", 0.0],
                    ["Materiali", "Amido", 0.0],
                    ["Materiali", "DPI", 0.0],
                    ["Materiali", "Immobiliare", 0.0],
                    ["Materiali", "Vari", 0.0],
                    # Metodo
                    ["Metodo", "Nuova procedura-istruzione", 0.0],
                    ["Metodo", "Periodo adattamento", 0.0],
                    ["Metodo", "Redistribuzione aziendale", 0.0],
                    # Vendite
                    ["Vendite", "Diminuzione delle vendite del 1% rispetto all’anno precedente", 0.0],
                    ["Vendite", "Diminuzione del fatturato del 5% rispetto all’anno precedente", 0.0],
                    # Reputazione
                    ["Reputazione", "Diminuzione del fatturato del 10% rispetto all’anno precedente", 0.0],
                    ["Reputazione", "Diminuzione del fatturato del 15% rispetto all’anno precedente", 0.0],
                    # Sanzioni
                    ["Sanzioni", "Sanzioni amministrative / penali (scaglionate)", 0.0]
                ]
                st.session_state.df_calcolo_economico_nm = pd.DataFrame(data_nm, columns=[
                    "Area d'impatto", "Sottocategoria", "Stima costo (€)"
                ])
            
            df_ce = st.session_state.df_calcolo_economico_nm
            
            # Assegnazione automatica dei valori calcolati
            df_ce.loc[df_ce["Sottocategoria"] == "Percentuale di indennità", "Stima costo (€)"] = perc_indennita_val
            df_ce.loc[df_ce["Sottocategoria"] == "RAL per mansione", "Stima costo (€)"] = ral_mansione_calc
            df_ce.loc[df_ce["Sottocategoria"] == "Diminuzione delle vendite del 1% rispetto all’anno precedente", "Stima costo (€)"] = val_vendite_1
            df_ce.loc[df_ce["Sottocategoria"] == "Diminuzione del fatturato del 5% rispetto all’anno precedente", "Stima costo (€)"] = val_vendite_5
            df_ce.loc[df_ce["Sottocategoria"] == "Diminuzione del fatturato del 10% rispetto all’anno precedente", "Stima costo (€)"] = val_rep_10
            df_ce.loc[df_ce["Sottocategoria"] == "Diminuzione del fatturato del 15% rispetto all’anno precedente", "Stima costo (€)"] = val_rep_15
            
            # Editor della tabella dinamica
            edited_ce = st.data_editor(
                df_ce,
                use_container_width=True,
                num_rows="fixed",
                column_config={
                    "Area d'impatto": st.column_config.TextColumn("Area d'impatto", disabled=True),
                    "Sottocategoria": st.column_config.TextColumn("Sottocategoria", disabled=True),
                    "Stima costo (€)": st.column_config.NumberColumn(
                        "Valore / Costo (€ o %)", 
                        min_value=0.0, 
                        step=10.0, 
                        format="%.2f €"
                    )
                },
                key="editor_calcolo_economico_nm"
            )
            
            st.session_state.df_calcolo_economico_nm = edited_ce
            
            # Esclusione della riga "Percentuale di indennità" dal calcolo monetario
            df_costi_monetari = edited_ce[edited_ce["Sottocategoria"] != "Percentuale di indennità"]
            
            # Calcolo automatico del totale
            st.markdown("### Riepilogo Costi per Area e Totale Generale")
            totale_generale = df_costi_monetari["Stima costo (€)"].sum()
            
            col_tot1, col_tot2 = st.columns(2)
            with col_tot1:
                st.metric(label="💰 STIMA ECONOMICA TOTALE", value=format_euro(totale_generale))
                
            with col_tot2:
                riepilogo_aree = df_costi_monetari.groupby("Area d'impatto")["Stima costo (€)"].sum()
                st.markdown("**Totali parziali per Area (esclusa % indennità):**")
                for area, val in riepilogo_aree.items():
                    st.text(f"- {area}: {format_euro(val)}")
            
            st.markdown("---")
            
            # Preparazione DataFrame per Export con formato italiano (virgola per i decimali)
            df_export = edited_ce.copy()
            df_export.insert(0, "Evento / Analisi Collegata", scelta_rif)
            df_export["Stima costo (€)"] = df_export["Stima costo (€)"].apply(format_csv_number)
            
            riga_totale = pd.DataFrame([[scelta_rif, "TOTALE GENERALE", "SOMMA TUTTI I COSTI", format_csv_number(totale_generale)]], columns=df_export.columns)
            df_export = pd.concat([df_export, riga_totale], ignore_index=True)
            
            # Generazione del nome file dinamico
            import re
            clean_rif = re.sub(r'[\\/*?:"<>|]', "", scelta_rif)
            clean_rif = clean_rif.replace(" ", "_")
            file_name_export = f"{clean_rif}_Stima_Costo_Economico_NM.csv"
            
            # Percorso su GitHub: Stima_Economica/Report/<nome_file>.csv
            github_repo_path = f"Stima_Economica/Report/{file_name_export}"
            
            # Converti in stringa CSV codificata in utf-8
            csv_content_str = df_export.to_csv(index=False, sep=";")
            csv_bytes = csv_content_str.encode("utf-8")
            
            col_save1, col_save2 = st.columns(2)
            with col_save1:
                if st.button("Salva Report Online", key="btn_save_stima_gh"):
                    if 'save_to_github' in globals():
                        if save_to_github(github_repo_path, csv_bytes, f"Add {file_name_export}"):
                            st.success(f"Report salvato con successo su GitHub in: `{github_repo_path}`")
                    else:
                        st.error("Funzione `save_to_github` non trovata nel sistema.")
            
            with col_save2:
                st.download_button(
                    label="Scarica Report CSV",
                    data=csv_bytes,
                    file_name=file_name_export,
                    mime="text/csv",
                    use_container_width=True
                )
# ==================================================================
# --- SEZIONE: Skill Matrix ---
# ==================================================================
if nav == "Skill Matrix":
    st.header("Skill Matrix - Gestione e Autovalutazione")
    
    # Gestione autenticazione per la Sezione Skill Matrix
    if "auth_skill_matrix" not in st.session_state:
        st.session_state.auth_skill_matrix = False

    # Scelta della sottosezione (visibile a tutti)
    sotto_sec_sm = st.radio(
        "Seleziona Sottosezione",
        ["Autovalutazione Skill Matrix", "Skill Matrix"],
        horizontal=True,
        key="radio_sotto_sec_sm"
    )
        
    # Gestione percorsi base
    try:
        base_dir = os.path.dirname(os.path.abspath(__file__))
    except NameError:
        base_dir = os.getcwd()
            
    skill_matrix_dir = os.path.join(base_dir, "Skill_Matrix")
    if not os.path.exists(skill_matrix_dir):
        skill_matrix_dir_alt = os.path.join(base_dir, "APP HSE", "Skill_Matrix")
        if os.path.exists(os.path.join(base_dir, "APP HSE")) or "APP HSE" in base_dir:
            skill_matrix_dir = skill_matrix_dir_alt

    # =========================================================
    # 1. SOTTOSEZIONE PUBBLICA - AUTOVALUTAZIONE SKILL MATRIX
    # =========================================================
    if sotto_sec_sm == "Autovalutazione Skill Matrix":
        st.subheader("Autovalutazione Skill Matrix")
        st.markdown("Consulta o scarica il documento PDF di autovalutazione e compila il form sottostante.")
        
        # Definizione percorsi locali per la lettura del PDF
        base_dir = os.path.dirname(os.path.abspath(__file__))
        skill_matrix_dir = os.path.join(base_dir, "Skill_Matrix")
        
        # Download / Visualizzazione PDF "Skill Matrix Autovalutazione.pdf"
        file_pdf_sm = os.path.join(skill_matrix_dir, "Skill Matrix Autovalutazione.pdf")
        
        if os.path.exists(file_pdf_sm):
            with open(file_pdf_sm, "rb") as f:
                pdf_bytes = f.read()
            
            st.download_button(
                label="Scarica / Apri PDF 'Skill Matrix Autovalutazione'",
                data=pdf_bytes,
                file_name="Skill Matrix Autovalutazione.pdf",
                mime="application/pdf",
                use_container_width=True
            )
        else:
            st.warning("Il file PDF 'Skill Matrix Autovalutazione.pdf' non è stato trovato nella cartella 'Skill_Matrix'.")
        
        st.markdown("---")
        st.markdown("#### Form di Autovalutazione")
        
        # Form di compilazione
        with st.form("form_autovalutazione_skill"):
            col_f1, col_f2 = st.columns(2)
            with col_f1:
                nome_utente = st.text_input("Nome")
            with col_f2:
                cognome_utente = st.text_input("Cognome")
            
            col_f3, col_f4 = st.columns(2)
            with col_f3:
                inquadramento_mansione = st.text_input("Inquadramento-Mansione")
            with col_f4:
                ambito_lavorativo = st.selectbox(
                    "Ambito lavorativo", 
                    ["Uffici", "Produzione", "Manutenzione", "Stoccaggio MP", "Trasporto-Logistica", "Commerciale"]
                )
            
            data_compilazione = st.date_input("Data Autovalutazione", value=datetime.today())
            
            st.markdown("##### AUTOVALUTAZIONE DELLA SKILL MATRIX (Punteggio tra 1 e 5)")
            
            q1 = st.slider("Quanto conosci dei processi produttivi della tua mansione?", 1, 5, 3, key="q1")
            q2 = st.slider("Hai un buon rapporto con i colleghi di reparto?", 1, 5, 3, key="q2")
            q3 = st.slider("Valuta le tue capacità di interfacciarti con fornitori e/o clienti", 1, 5, 3, key="q3")
            q4 = st.slider("Quanto conosci dei processi produttivi del cartone e scatole?", 1, 5, 3, key="q4")
            q5 = st.slider("Quanto conosci il processo di pallettizzazione del prodotto?", 1, 5, 3, key="q5")
            q6 = st.slider("Hai competenze legali o tecniche?", 1, 5, 3, key="q6")
            q7 = st.slider("Quanto sei in grado di individuare i rischi-fabbisogni negli ambienti lavorativi", 1, 5, 3, key="q7")
            q8 = st.slider("Valuta la tua capacità d’adattamento", 1, 5, 3, key="q8")
            q9 = st.slider("Valuta le tue capacità comunicative", 1, 5, 3, key="q9")
            q10 = st.slider("Sei una persona precisa nel lavoro che svolge?", 1, 5, 3, key="q10")
            q11 = st.slider("Riesci a persuadere agli altri per svolgere delle attività o fare cambiamenti?", 1, 5, 3, key="q11")
            q12 = st.slider("Hai un’analisi critico del contesto lavorativo? (sai cosa funzione e cosa si potrebbe migliorare)", 1, 5, 3, key="q12")
            q13 = st.slider("Sei a conoscenza delle turnazioni di lavoro, come vengono comunicate e le dinamiche lavorative all’interno di ogni turno?", 1, 5, 3, key="q13")
            q14 = st.slider("Ci sono altre persone sotto la tua responsabilità o supervisione?", 1, 5, 3, key="q14")
            
            submitted_form = st.form_submit_button("Invia e Salva Autovalutazione", use_container_width=True)
            
            if submitted_form:
                if not nome_utente.strip() or not cognome_utente.strip():
                    st.error("Inserisci obbligatoriamente Nome e Cognome prima di procedere.")
                else:
                    clean_name = re.sub(r'[\\/*?:"<>|]', "", f"{nome_utente.strip()}_{cognome_utente.strip()}")
                    clean_date = re.sub(r'[\\/*?:"<>|]', "", str(data_compilazione))
                    file_name_csv = f"{clean_name}_{clean_date}_Autovalutazione_SkillMatrix.csv"
                    
                    # Costruiamo il percorso del file su GitHub
                    github_path = f"Skill_Matrix/Autovalutazione/{file_name_csv}"
                    
                    dati_form = {
                        "Campo": [
                            "Nome", "Cognome", "Inquadramento-Mansione", "Ambito lavorativo", "Data Autovalutazione",
                            "Processi produttivi mansione", "Rapporto colleghi", "Interfaccia fornitori-clienti",
                            "Processi cartone-scatole", "Processo pallettizzazione", "Competenze legali-tecniche",
                            "Individuazione rischi-fabbisogni", "Capacità d'adattamento", "Capacità comunicative",
                            "Precisione lavoro", "Persuasione", "Analisi critica contesto", "Turnazioni", "Responsabilità supervisione"
                        ],
                        "Valore": [
                            nome_utente.strip(), cognome_utente.strip(), inquadramento_mansione.strip(), ambito_lavorativo, str(data_compilazione),
                            q1, q2, q3, q4, q5, q6, q7, q8, q9, q10, q11, q12, q13, q14
                        ]
                    }
                    df_res = pd.DataFrame(dati_form)
                    csv_content = df_res.to_csv(index=False, sep=";")
                    
                    # Salvataggio tramite API GitHub
                    try:
                        token = st.secrets.get("GITHUB_TOKEN", "")
                        repo_name = st.secrets.get("REPO_NAME", "")
                        
                        if not token or not repo_name:
                            st.error("⚠️ GITHUB_TOKEN o REPO_NAME mancanti nei secrets di Streamlit.")
                        else:
                            g = Github(token)
                            repo = g.get_repo(repo_name)
                            
                            # Controlla se il file esiste già su GitHub per aggiornarlo o crearlo
                            try:
                                file_existing = repo.get_contents(github_path)
                                repo.update_file(
                                    path=github_path,
                                    message=f"Aggiornata autovalutazione: {file_name_csv}",
                                    content=csv_content,
                                    sha=file_existing.sha
                                )
                            except GithubException:
                                repo.create_file(
                                    path=github_path,
                                    message=f"Aggiunta nuova autovalutazione: {file_name_csv}",
                                    content=csv_content
                                )
                                
                            st.success(f"Autovalutazione salvata permanentemente su GitHub nel percorso: `{github_path}`")
                        
                    except Exception as e:
                        st.error(f"Errore nel salvataggio su GitHub: {e}")

    # =========================================================
    # 2. SOTTOSEZIONE RISERVATA - SKILL MATRIX (Lettura da PDF)
    # =========================================================
    elif sotto_sec_sm == "Skill Matrix":
        if not st.session_state.auth_skill_matrix:
            st.markdown("Inserisci la password per accedere alla sezione Skill Matrix.")
            pwd_skill = st.text_input("Password Skill Matrix", type="password", key="pwd_skill_input")
            
            if st.button("Verifica Password", use_container_width=True, key="btn_verify_pwd_skill"):
                password_secrets = st.secrets.get("PASSWORD_SEZIONE", st.secrets.get("SKILL_MATRIX_PASSWORD", "hse2026"))
                
                if pwd_skill and pwd_skill == password_secrets:
                    st.session_state.auth_skill_matrix = True
                    st.success("Accesso autorizzato!")
                    st.rerun()
                else:
                    st.error("Password errata o non valida.")
        else:
            if st.button("🚪 Disconnetti Sezione Riservata", key="btn_logout_skill_matrix"):
                st.session_state.auth_skill_matrix = False
                st.rerun()
                
            st.subheader("Skill Matrix - Panoramica Generale e Tabella Dinamica")
            st.markdown(
                "A ogni competenza si attribuirà un punteggio del 1 al 5. "
                "Le persone saranno anche classificate a seconda dell'area lavorativa d'appartenenza."
            )
            st.markdown("---")
            st.markdown("La tabella sottostante viene aggiornata automaticamente scansionando tutti i report PDF di autovalutazione.")
                
            # Percorsi cartelle
            skill_matrix_dir = os.path.join(base_dir, "Skill_Matrix")
            autoval_dir = os.path.join(skill_matrix_dir, "Autovalutazione")
            file_name_master = "Skill_Matrix_Panoramica_Generale.csv"
            master_local_path = os.path.join(skill_matrix_dir, file_name_master)
            github_master_path = f"Skill_Matrix/{file_name_master}"

            lista_righe_tabella = []

            # --- PARSING AUTOMATICO DEI FILE PDF IN AUTOVALUTAZIONE ---
            if os.path.exists(autoval_dir):
                files_pdf = [f for f in os.listdir(autoval_dir) if f.lower().endswith(".pdf")]
                
                for f_pdf in files_pdf:
                    pdf_path = os.path.join(autoval_dir, f_pdf)
                    try:
                        import pypdf
                        reader = pypdf.PdfReader(pdf_path)
                        testo_completo = ""
                        for page in reader.pages:
                            testo_completo += page.extract_text() + "\n"
                        
                        # Mappatura dei dati dal testo estratto
                        dati_pdf = {}
                        for line in testo_completo.split("\n"):
                            if ":" in line:
                                parti = line.split(":", 1)
                                chiave = parti[0].strip()
                                valore = parti[1].strip()
                                dati_pdf[chiave] = valore

                        def get_pdf_val(key_name, default_val):
                            return dati_pdf.get(key_name, default_val)

                        lista_righe_tabella.append({
                            "Nome": str(get_pdf_val("Nome", "N/D")),
                            "Cognome": str(get_pdf_val("Cognome", "N/D")),
                            "Inquadramento-Mansione": str(get_pdf_val("Inquadramento-Mansione", "")),
                            "Ambito lavorativo": str(get_pdf_val("Ambito lavorativo", "Produzione")),
                            "Data Autovalutazione": str(get_pdf_val("Data Autovalutazione", "")),
                            "Processi produttivi mansione": float(get_pdf_val("Processi produttivi mansione", 3)),
                            "Rapporto colleghi": float(get_pdf_val("Rapporto colleghi", 3)),
                            "Interfaccia fornitori-clienti": float(get_pdf_val("Interfaccia fornitori-clienti", 3)),
                            "Processi cartone-scatole": float(get_pdf_val("Processi cartone-scatole", 3)),
                            "Processo pallettizzazione": float(get_pdf_val("Processo pallettizzazione", 3)),
                            "Competenze legali-tecniche": float(get_pdf_val("Competenze legali-tecniche", 3)),
                            "Individuazione rischi-fabbisogni": float(get_pdf_val("Individuazione rischi-fabbisogni", 3)),
                            "Capacità d'adattamento": float(get_pdf_val("Capacità d'adattamento", 3)),
                            "Capacità comunicative": float(get_pdf_val("Capacità comunicative", 3)),
                            "Precisione lavoro": float(get_pdf_val("Precisione lavoro", 3)),
                            "Persuasione": float(get_pdf_val("Persuasione", 3)),
                            "Analisi critica contesto": float(get_pdf_val("Analisi critica contesto", 3)),
                            "Turnazioni": float(get_pdf_val("Turnazioni", 3)),
                            "Responsabilità supervisione": float(get_pdf_val("Responsabilità supervisione", 3)),
                            "File Sorgente": f_pdf
                        })
                    except Exception:
                        pass

            # Costruzione DataFrame aggiornato
            if lista_righe_tabella:
                df_master_sm = pd.DataFrame(lista_righe_tabella)
            else:
                # Fallback se non ci sono PDF o se è presente il master locale
                if os.path.exists(master_local_path):
                    try:
                        df_master_sm = pd.read_csv(master_local_path, sep=";")
                    except Exception:
                        df_master_sm = pd.DataFrame()
                else:
                    df_master_sm = pd.DataFrame()

            # --- VISUALIZZAZIONE ED EDITING TABELLA ---
            if not df_master_sm.empty:
                st.markdown("#### Tabella Panoramica Modificabile")
                st.info("Visualizza o modifica i dati estratti dai report PDF. Clicca sui pulsanti in basso per salvare su GitHub o scaricare.")
                    
                edited_master_sm = st.data_editor(
                    df_master_sm,
                    use_container_width=True,
                    num_rows="dynamic",
                    column_config={
                        "Ambito lavorativo": st.column_config.SelectboxColumn(
                            "Ambito lavorativo",
                            options=["Uffici", "Produzione", "Manutenzione", "Stoccaggio MP", "Trasporto-Logistica", "Commerciale"],
                            required=True
                        ),
                        "File Sorgente": st.column_config.TextColumn("File Sorgente", disabled=True)
                    },
                    key="editor_skill_matrix_generale"
                )
                    
                col_btn1, col_btn2 = st.columns(2)
                csv_master_data = edited_master_sm.to_csv(index=False, sep=";")
                
                with col_btn1:
                    if st.button("Salva Modifiche Tabella Skill Matrix", use_container_width=True, key="btn_save_master_sm"):
                        # 1. Salvataggio locale
                        try:
                            os.makedirs(skill_matrix_dir, exist_ok=True)
                            with open(master_local_path, "w", encoding="utf-8") as f:
                                f.write(csv_master_data)
                            salvati_local = True
                        except Exception as e:
                            st.error(f"Errore nel salvataggio locale: {e}")
                            salvati_local = False

                        # 2. Salvataggio ed invio su GitHub
                        token = st.secrets.get("GITHUB_TOKEN", "")
                        repo_name = st.secrets.get("REPO_NAME", "")
                        if token and repo_name and salvati_local:
                            try:
                                g = Github(token)
                                repo = g.get_repo(repo_name)
                                try:
                                    file_existing = repo.get_contents(github_master_path)
                                    repo.update_file(
                                        path=github_master_path,
                                        message=f"Aggiornata panoramica generale Skill Matrix da PDF: {file_name_master}",
                                        content=csv_master_data,
                                        sha=file_existing.sha
                                    )
                                except GithubException:
                                    repo.create_file(
                                        path=github_master_path,
                                        message=f"Creata panoramica generale Skill Matrix da PDF: {file_name_master}",
                                        content=csv_master_data
                                    )
                                st.success(f"✅ Dati aggiornati salvati con successo su GitHub in `{github_master_path}`!")
                                st.rerun()
                            except Exception as e:
                                st.error(f"Errore durante il salvataggio su GitHub: {e}")
                        else:
                            st.warning("⚠️ Credentials GitHub assenti. File salvato solo in locale.")
                                
                with col_btn2:
                    st.download_button(
                        label="📥 Scarica Tabella Master in formato CSV",
                        data=csv_master_data,
                        file_name="Skill_Matrix_Panoramica_Generale.csv",
                        mime="text/csv",
                        use_container_width=True,
                        key="btn_download_master_sm_csv"
                    )
            else:
                st.info("Nessun report PDF trovato nella cartella 'Skill_Matrix/Autovalutazione'.")
#------------------------------------------------------------------------------------------------------------------
# SEZIONE 12: RICONOSCIMENTO SEGNALANTI NEAR MISS
#------------------------------------------------------------------------------------------------------------------
if nav == "Riconoscimento":
    st.header("Sezione Riconoscimento")
    st.markdown("Gestione delle classifiche e assegnazione dei punteggi di riconoscimento.")
    
    # ---------------------------------------------------------
    # Autenticazione Password Sezione Riconoscimento tramite Secrets
    # ---------------------------------------------------------
    if "auth_riconoscimento" not in st.session_state:
        st.session_state.auth_riconoscimento = False

    if not st.session_state.auth_riconoscimento:
        st.markdown("🔒 **Area Riservata:** Inserisci la password per accedere.")
        pwd_riconoscimento = st.text_input("Password Riconoscimento", type="password", key="pwd_riconoscimento_input")
        if st.button("Accedi", use_container_width=True, key="btn_auth_riconoscimento"):
            correct_pwd = st.secrets.get("PASSWORD_SEZIONE", "hse2026")
            if pwd_riconoscimento == correct_pwd:
                st.session_state.auth_riconoscimento = True
                st.success("Accesso autorizzato!")
                st.rerun()
            else:
                st.error("Password errata.")
    else:
        # ---------------------------------------------------------
        # Navigazione Sottosezioni
        # ---------------------------------------------------------
        sotto_sec_ric = st.radio(
            "Seleziona Sottosezione",
            ["Classificazione", "Assegnazione Riconoscimento"],
            horizontal=True,
            key="radio_sotto_sec_ric"
        )
        st.markdown("---")
        
        # Gestione percorsi base e GitHub
        try:
            base_dir = os.path.dirname(os.path.abspath(__file__))
        except NameError:
            base_dir = os.getcwd()
            
        rel_github_path = "Riconoscimento/Riconoscimento_Partecipazione_NM.csv"
        file_riconoscimenti_csv = os.path.join(base_dir, "Riconoscimento", "Riconoscimento_Partecipazione_NM.csv")
        
        # ---------------------------------------------------------
        # Funzione di supporto: Lettura/Caricamento dati salvati
        # ---------------------------------------------------------
        def carica_o_inizializza_punteggi():
            # Tentativo 1: Caricamento da file locale in Riconoscimento/Riconoscimento_Partecipazione_NM.csv
            if os.path.exists(file_riconoscimenti_csv):
                try:
                    df_ric = pd.read_csv(file_riconoscimenti_csv, sep=";")
                    if not df_ric.empty:
                        return df_ric
                except Exception:
                    pass
            
            # Recupero nominativi da segnalazioni e Skill Matrix se non esiste ancora il file principale
            segnalatori_set = set()
            file_nm = os.path.join(base_dir, "segnalazioni_near_miss.csv")
            if os.path.exists(file_nm):
                try:
                    df_nm = pd.read_csv(file_nm, sep=";")
                    col_seg = [c for c in df_nm.columns if c.strip().lower() == "segnalatore"]
                    if col_seg:
                        col_name = col_seg[0]
                        segnalatori_validi = df_nm[col_name].dropna().astype(str).str.strip()
                        for s in segnalatori_validi:
                            if s != "" and s.lower() != "nan":
                                segnalatori_set.add(s)
                except Exception as e:
                    st.warning(f"Errore nella lettura di segnalazioni_near_miss.csv: {e}")
                    
            skill_set = set()
            autoval_dir = os.path.join(base_dir, "Skill_Matrix", "Autovalutazione")
            if not os.path.exists(autoval_dir):
                autoval_dir = os.path.join(base_dir, "APP HSE", "Skill_Matrix", "Autovalutazione")
            if os.path.exists(autoval_dir):
                files_csv_sm = [f for f in os.listdir(autoval_dir) if f.endswith(".csv")]
                for f_csv in files_csv_sm:
                    f_path = os.path.join(autoval_dir, f_csv)
                    try:
                        df_sm = pd.read_csv(f_path, sep=";")
                        nome_val = df_sm.loc[df_sm["Campo"] == "Nome", "Valore"]
                        cognome_val = df_sm.loc[df_sm["Campo"] == "Cognome", "Valore"]
                        nome_str = str(nome_val.values[0]).strip() if not nome_val.empty else ""
                        cognome_str = str(cognome_val.values[0]).strip() if not cognome_val.empty else ""
                        full_name = f"{nome_str} {cognome_str}".strip()
                        if full_name and full_name != "N/D N/D":
                            skill_set.add(full_name)
                    except Exception:
                        pass
                        
            tutti_nominativi = list(segnalatori_set.union(skill_set))
            
            rows = []
            for nom in tutti_nominativi:
                is_seg = nom in segnalatori_set
                is_sk = nom in skill_set
                fonte = "Segnalatore & Skill Matrix" if (is_seg and is_sk) else ("Segnalatore" if is_seg else "Skill Matrix")
                
                rows.append({
                    "Nominativo": nom,
                    "Fonte": fonte,
                    "Punti Segnalazione (+50)": 0,
                    "Punti Skill Matrix (+25)": 0,
                    "Punteggio Totale": 0
                })
            
            return pd.DataFrame(rows if rows else [{
                "Nominativo": "Esempio", 
                "Fonte": "N/D", 
                "Punti Segnalazione (+50)": 0, 
                "Punti Skill Matrix (+25)": 0, 
                "Punteggio Totale": 0
            }])

        # Carica il dataframe generale
        df_riconoscimenti = carica_o_inizializza_punteggi()
        
        # Normalizzazione tipi e ricalcolo totale
        df_riconoscimenti["Punti Segnalazione (+50)"] = pd.to_numeric(df_riconoscimenti["Punti Segnalazione (+50)"], errors='coerce').fillna(0).astype(int)
        df_riconoscimenti["Punti Skill Matrix (+25)"] = pd.to_numeric(df_riconoscimenti["Punti Skill Matrix (+25)"], errors='coerce').fillna(0).astype(int)
        df_riconoscimenti["Punteggio Totale"] = df_riconoscimenti["Punti Segnalazione (+50)"] + df_riconoscimenti["Punti Skill Matrix (+25)"]
        
        # Ordinamento decrescente per Punteggio Totale
        df_riconoscimenti = df_riconoscimenti.sort_values(by="Punteggio Totale", ascending=False).reset_index(drop=True)

        # =========================================================
        # 1. SOTTOSEZIONE - CLASSIFICAZIONE
        # =========================================================
        if sotto_sec_ric == "Classificazione":
            st.subheader("Classifica Generale Riconoscimenti")
            st.markdown("Visualizzazione della classifica complessiva letta da `Riconoscimento_Partecipazione_NM`.")
            
            # Podio sui primi 3 con il punteggio più elevato
            st.markdown("### Classifica Generale Riconoscimenti (Top 3)")
            if len(df_riconoscimenti) > 0 and df_riconoscimenti.iloc[0]["Punteggio Totale"] > 0:
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric(
                        "🥇 1° Posto", 
                        f"{df_riconoscimenti.iloc[0]['Nominativo']}", 
                        f"{df_riconoscimenti.iloc[0]['Punteggio Totale']} Punti"
                    )
                with col2:
                    if len(df_riconoscimenti) > 1:
                        st.metric(
                            "🥈 2° Posto", 
                            f"{df_riconoscimenti.iloc[1]['Nominativo']}", 
                            f"{df_riconoscimenti.iloc[1]['Punteggio Totale']} Punti"
                        )
                with col3:
                    if len(df_riconoscimenti) > 2:
                        st.metric(
                            "🥉 3° Posto", 
                            f"{df_riconoscimenti.iloc[2]['Nominativo']}", 
                            f"{df_riconoscimenti.iloc[2]['Punteggio Totale']} Punti"
                        )
            else:
                st.info("Nessun punteggio ancora registrato nel sistema per generare il podio.")
            
            st.markdown("---")
            st.markdown("### Tabella Classifica Complessiva")
            st.dataframe(
                df_riconoscimenti[["Nominativo", "Fonte", "Punti Segnalazione (+50)", "Punti Skill Matrix (+25)", "Punteggio Totale"]],
                use_container_width=True,
                hide_index=True
            )

        # =========================================================
        # 2. SOTTOSEZIONE - ASSEGNAZIONE RICONOSCIMENTO
        # =========================================================
        elif sotto_sec_ric == "Assegnazione Riconoscimento":
            st.subheader("Assegnazione Manuale Punteggi e Modifica Tabella")
            st.info(
                "In questa sezione puoi assegnare manualmente i punti:\n"
                "- **+50 Punti** per le attività di segnalazione Near Miss.\n"
                "- **+25 Punti** per le competenze inserite nella Skill Matrix.\n"
                "I dati inseriti verranno salvati direttamente su GitHub nella cartella `Riconoscimento`."
            )
            
            # Tabella dinamica modificabile (data_editor)
            df_edited = st.data_editor(
                df_riconoscimenti,
                use_container_width=True,
                num_rows="dynamic",
                column_config={
                    "Nominativo": st.column_config.TextColumn("Nome e Cognome / Segnalatore", required=True),
                    "Fonte": st.column_config.TextColumn("Origine Dato", disabled=True),
                    "Punti Segnalazione (+50)": st.column_config.NumberColumn(
                        "Punti Segnalatore",
                        help="Punti assegnati come segnalatore (es. 0, 50, 100...)",
                        min_value=0,
                        step=50
                    ),
                    "Punti Skill Matrix (+25)": st.column_config.NumberColumn(
                        "Punti Skill Matrix",
                        help="Punti assegnati per Skill Matrix (es. 0, 25, 50...)",
                        min_value=0,
                        step=25
                    ),
                    "Punteggio Totale": st.column_config.NumberColumn("Totale Punti", disabled=True)
                },
                key="editor_riconoscimento_table"
            )
            
            # RICALCOLO AUTOMATICO DEL TOTALE
            df_edited["Punti Segnalazione (+50)"] = pd.to_numeric(df_edited["Punti Segnalazione (+50)"], errors='coerce').fillna(0).astype(int)
            df_edited["Punti Skill Matrix (+25)"] = pd.to_numeric(df_edited["Punti Skill Matrix (+25)"], errors='coerce').fillna(0).astype(int)
            df_edited["Punteggio Totale"] = df_edited["Punti Segnalazione (+50)"] + df_edited["Punti Skill Matrix (+25)"]
            
            col_sav1, col_sav2 = st.columns(2)
            
            with col_sav1:
                if st.button("Salva Assegnazione Punteggi Online", use_container_width=True, key="btn_save_riconoscimenti"):
                    # Ordina prima di salvare
                    df_edited = df_edited.sort_values(by="Punteggio Totale", ascending=False).reset_index(drop=True)
                    
                    # Salva in locale
                    os.makedirs(os.path.dirname(file_riconoscimenti_csv), exist_ok=True)
                    df_edited.to_csv(file_riconoscimenti_csv, index=False, sep=";")
                    
                    # Salvataggio remoto su GitHub Repository Online
                    csv_bytes = df_edited.to_csv(index=False, sep=";").encode('utf-8')
                    if 'save_to_github' in globals():
                        if save_to_github(rel_github_path, csv_bytes, "Aggiornamento Riconoscimento_Partecipazione_NM.csv"):
                            st.success(f"Punteggi salvati su GitHub in `{rel_github_path}` e classifica aggiornata!")
                            st.rerun()
                        else:
                            st.warning("Salvato in locale, ma si è verificato un errore durante l'invio a GitHub.")
                    else:
                        st.success("Punteggi salvati in locale con successo!")
                        st.rerun()

            with col_sav2:
                csv_ric_data = df_edited.to_csv(index=False, sep=";").encode('utf-8')
                st.download_button(
                    label="Scarica Classifica Riconoscimenti (CSV)",
                    data=csv_ric_data,
                    file_name="Riconoscimento_Partecipazione_NM.csv",
                    mime="text/csv",
                    use_container_width=True
                )
#------------------------------------------------------------------------------------------------------------------
# SEZIONE 13: Controllo DPI
#------------------------------------------------------------------------------------------------------------------
if nav == "Controllo DPI":
    st.header("Controllo Automatico DPI con Intelligenza Artificiale (YOLO)")
    st.markdown("Seleziona la mansione del lavoratore e carica una foto o scatta uno snapshot per verificare la conformità dei DPI.")
    
    mansione_scelta = st.selectbox("Seleziona Mansione:", list(MANSIONI_DPI.keys()), key="dpi_mansione_sel")
    dpi_richiesti = MANSIONI_DPI[mansione_scelta]
    
    st.info(f"**DPI obbligatori per {mansione_scelta}:** {', '.join(dpi_richiesti)}")
    
    metodo_input_dpi = st.radio("Modalità acquisizione immagine DPI:", ["Carica Immagine", "Scatta Foto con Webcam"], key="dpi_input_mode")
    
    file_foto_dpi = None
    if metodo_input_dpi == "Carica Immagine":
        file_foto_dpi = st.file_uploader("Carica foto lavoratore", type=["png", "jpg", "jpeg"], key="dpi_file_uploader")
    else:
        file_foto_dpi = st.camera_input("Scatta foto al lavoratore", key="dpi_camera_input")
        
    if file_foto_dpi is not None:
        img_pil = Image.open(file_foto_dpi).convert('RGB')
        img_np = np.array(img_pil)
        
        st.image(img_pil, caption="Immagine Acquisita", use_container_width=True)
        
        if st.button("Avvia Analisi DPI con YOLO", use_container_width=True, key="btn_run_yolo_dpi"):
            with st.spinner("Elaborazione rilevamento DPI e invio notifica e-mail in corso..."):
                try:
                    results = model(img_np)
                    res_plotted = results[0].plot() # Immagine BGR/RGB con bounding box
                    
                    st.image(res_plotted, caption="Risultato Rilevamento YOLO", use_container_width=True)
                    
                    # Conversione dell'immagine con i box YOLO in byte JPEG per l'e-mail
                    img_res_pil = Image.fromarray(res_plotted)
                    buffer_img = io.BytesIO()
                    img_res_pil.save(buffer_img, format="JPEG")
                    img_bytes_yolo = buffer_img.getvalue()

                    boxes = results[0].boxes
                    class_indices = boxes.cls.cpu().numpy() if len(boxes) > 0 else []
                    class_names = results[0].names
                    rilievi_oggetti = [class_names[int(c)].lower() for c in class_indices]
                    
                    st.markdown("### Esito Controllo Conformità:")
                    st.write(f"**DPI rilevati dal sistema:** {list(set(rilievi_oggetti)) if rilievi_oggetti else 'Nessun DPI riconosciuto'}")
                    
                    mancanti = [dpi for dpi in dpi_richiesti if not any(dpi.lower() in r.lower() for r in rilievi_oggetti)]
                    esito_conforme = len(mancanti) == 0
                    
                    if esito_conforme:
                        st.success("✅ **CONFORME:** Tutti i DPI obbligatori per la mansione sono indossati correttamente.")
                    else:
                        st.error(f"❌ **NON CONFORME:** Mancano i seguenti DPI obbligatori: {', '.join(mancanti)}")

                    # Invio e-mail automatico completo di immagine allegata
                    esito_mail, msg_mail = invia_email_notifica_dpi(
                        mansione=mansione_scelta,
                        dpi_rilevati=list(set(rilievi_oggetti)),
                        mancanti=mancanti,
                        esito_conforme=esito_conforme,
                        img_bytes=img_bytes_yolo
                    )
                    
                    if esito_mail:
                        st.info(f"📧 Notifica e-mail con immagine YOLO inviata con successo a `{st.secrets.get('EMAIL', '')}`.")
                    else:
                        st.warning(f"⚠️ Impossibile inviare l'e-mail: {msg_mail}")

                except Exception as e:
                    st.error(f"Errore durante l'esecuzione del modello YOLO: {e}")

# ==================================================================================================
# SEZIONE 14: SEGNALAZIONE MANUTENZIONE
# ==================================================================================================
if nav == "Segnalazione Manutenzione":
    st.header("Formulario Segnalazione Manutenzione (NM/NC)")
    st.markdown(
        "Compila il modulo sottostante per inviare una segnalazione di Near Miss / Non Conformità legata alla Manutenzione."
    )

    # ---------------------------------------------------------
    # 1. GESTIONE CARTELLA E DOWNLOAD MODULO VUOTO
    # ---------------------------------------------------------
    DIR_DEST = "Segnalazione_NM_Manutenzione"
    os.makedirs(DIR_DEST, exist_ok=True)

    # Pulsante per scaricare il PDF Vuoto originale se presente nella cartella
    pdf_vuoto_path = "Segnalazione_NM_Manutenzione.pdf"
    if os.path.exists(pdf_vuoto_path):
        with open(pdf_vuoto_path, "rb") as f:
            st.download_button(
                label="Scarica Formulario Vuoto (PDF)",
                data=f.read(),
                file_name="Formulario_Vuoto_Segnalazione_Manutenzione.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
    st.markdown("---")

    # ---------------------------------------------------------
    # 2. COMPILAZIONE FORMULARIO WEB
    # ---------------------------------------------------------
    with st.form(key="form_segnalazione_manutenzione", clear_on_submit=True):
        st.subheader("Informazioni Generali")
        col1, col2 = st.columns(2)
        with col1:
            tipo_evento = st.radio(
                "Tipo evento",
                ["Near Miss", "Non Conformità"],
                horizontal=True,
            )
            segnalatore = st.text_input(
                "Segnalatore (Mansione o Nome e Cognome)",
                placeholder="Es. Mario Rossi",
            )
            manutenzione_in = st.radio(
                "Manutenzione in",
                ["In Azienda", "Azienda Esterna"],
                horizontal=True,
            )
            sesso = st.radio("Sesso", ["Maschio", "Femmina"], horizontal=True)

        with col2:
            fascia_eta = st.selectbox(
                "Fascia di Età",
                ["<18 anni", "18-30 anni", "31-50 anni", "51-67 anni"],
            )
            data_evento = st.date_input("Data Evento", value=date.today())
            luogo = st.radio(
                "Luogo",
                ["In Azienda", "In itinere", "In missione"],
                horizontal=True,
            )
            reparto = st.text_input(
                "Reparto (se In Azienda)",
                placeholder="Es. Officina / Produzione",
            )

        st.markdown("---")
        st.subheader("Descrizione Criticità e Rischio")
        desc_evento = st.text_area(
            "Descrizione dell'evento o della criticità *",
            placeholder="Descrivi dettagliatamente l'accaduto...",
        )
        desc_rischio = st.text_area(
            "Descrizione del rischio associato al Near Miss",
            placeholder="Descrivi i potenziali rischi...",
        )

        st.markdown("---")
        st.subheader("Possibili Cause della Richiesta di Manutenzione")

        opzioni_cause = [
            "Elettrica: Cortocircuito",
            "Elettrica: Sovratensione",
            "Elettrica: Malfunzionamento dei sensori",
            "Elettrica: Problemi ai circuiti di controllo",
            "Elettrica: Interruzione dell'alimentazione",
            "Elettrica: Usura dei cavi elettrici",
            "Elettrica: Calore eccessivo",
            "Elettrica: Umidità nel quadro elettrico",
            "Elettrica: Polvere nel quadro elettrico",
            "Elettrica: Sovraccarico di rete",
            "Elettrica: Dispersione",
            "Elettrica: Problemi messa a terra",
            "Meccanica: Usura macchine",
            "Meccanica: Corrosione",
            "Meccanica: Eccesso di vibrazioni nei macchinari",
            "Meccanica: Urto meccanico",
            "Meccanica: Presenza di umidità nei macchinari",
            "Meccanica: Presenza di polvere nei macchinari",
            "Meccanica: Eccessivo calore nei macchinari",
            "Idrico: Rottura tubazione",
            "Idrico: Guasto caldaia",
            "Idrico: Filtrazione acqua o lubrificanti",
            "Idrico: Problema nello scarico",
            "Idrico: Mancata manutenzione del sistema di filtraggio dell'acqua",
            "Filtri: Usura dei filtri",
            "Filtri: Mancanza di filtri",
            "Filtri: Inadeguatezza del filtro",
            "Installazione: Guasto alle chiusure porte e finestre",
            "Errore dei lavoratori",
            "Mancata manutenzione preventiva",
        ]
        cause_selezionate = st.multiselect(
            "Seleziona una o più cause rilevate:", opzioni_cause
        )
        altro_cause = st.text_input(
            "Altro (specificare altre cause):", placeholder="Inserisci il testo..."
        )

        st.markdown("---")
        st.subheader("Possibili Conseguenze")

        opzioni_conseguenze = [
            "Ustioni",
            "Anomalia/guasto in avviamento/arresto/esercizio (funzionamento)",
            "Scivolamento",
            "Abrasione",
            "Elettrocuzione",
            "Incastramento di un arto",
            "Ostruzione macchinario",
            "Inalazione sostanze tossiche",
            "Ferite o tagli",
            "Presenza di elettricità/linea elettrica accessibile",
            "Presenza imprevista di liquidi (acqua, olio, ...)",
            "Presenza imprevista di gas, vapori",
            "Criticità su impianti generali a supporto dell'area di lavoro",
            "Urti",
        ]
        conseguenze_selezionate = st.multiselect(
            "Seleziona le possibili conseguenze della non manutenzione:",
            opzioni_conseguenze,
        )
        altro_conseguenze = st.text_input(
            "Altre conseguenze (specificare):", placeholder="Inserisci il testo..."
        )

        st.markdown("---")
        st.subheader("Valutazioni e Interventi")
        col_v1, col_v2 = st.columns(2)
        with col_v1:
            gia_presentata = st.radio(
                "La situazione rilevata si è già presentata in passato?",
                ["Sì frequentemente", "Sì raramente", "No"],
            )
            necessita_manutenzione = st.radio(
                "Necessità di manutenzione?", ["Sì", "No"]
            )
            tipo_manutenzione = st.multiselect(
                "Se 'Sì', scegliere la tipologia di manutenzione:",
                [
                    "Manutenzione preventiva",
                    "Manutenzione ordinaria",
                    "Manutenzione a guasto",
                    "Manutenzione straordinaria",
                    "Manutenzione migliorativa",
                ],
            )

        with col_v2:
            intervento_esterno = st.radio(
                "Necessario intervento da azienda esterna?", ["Sì", "No"]
            )
            azioni_miglioramento = st.text_area(
                "Valutazioni / azioni / proposte di miglioramento",
                placeholder="Inserisci eventuali proposte...",
            )

        submitted = st.form_submit_button(
            "Invia Segnalazione Manutenzione", use_container_width=True
        )

    # ---------------------------------------------------------
    # 3. ELABORAZIONE E SALVATAGGIO DEI DATI SU GITHUB
    # ---------------------------------------------------------
    if submitted:
        if not desc_evento.strip():
            st.error(
                "Il campo 'Descrizione dell'evento o della criticità' è obbligatorio!"
            )
        else:
            now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

            # Costruzione riga dati
            nuova_risposta = {
                "Data Ora Invio": now_str,
                "Tipologia": "NM_Manutenzione",
                "Tipo Evento": tipo_evento,
                "Segnalatore": segnalatore,
                "Manutenzione In": manutenzione_in,
                "Sesso": sesso,
                "Fascia Età": fascia_eta,
                "Data Evento": data_evento.strftime("%d/%m/%Y"),
                "Luogo": luogo,
                "Reparto": reparto,
                "Descrizione Evento": desc_evento,
                "Descrizione Rischio": desc_rischio,
                "Cause Selezionate": ", ".join(cause_selezionate),
                "Altre Cause": altro_cause,
                "Conseguenze Selezionate": ", ".join(conseguenze_selezionate),
                "Altre Conseguenze": altro_conseguenze,
                "Già Presentata in Passato": gia_presentata,
                "Necessità Manutenzione": necessita_manutenzione,
                "Tipo Manutenzione": ", ".join(tipo_manutenzione),
                "Intervento Azienda Esterna": intervento_esterno,
                "Azioni / Proposte Miglioramento": azioni_miglioramento,
            }

            try:
                # Inizializzazione della connessione con GitHub API
                github_token = st.secrets["GITHUB_TOKEN"]
                repo_name = st.secrets["REPO_NAME"]
                path_in_repo = "Segnalazione_NM_Manutenzione/manutenzione.csv"

                g = Github(github_token)
                repo = g.get_repo(repo_name)

                # Preparazione stringa della nuova riga in formato CSV (delimitatore ';')
                output = io.StringIO()
                writer = csv.DictWriter(
                    output, fieldnames=nuova_risposta.keys(), delimiter=";"
                )
                writer.writerow(nuova_risposta)
                nuova_riga_csv = output.getvalue()

                try:
                    # Se il file esiste già su GitHub, lo recupera e aggiunge la riga
                    file_content = repo.get_contents(path_in_repo)
                    contenuto_esistente = base64.b64decode(
                        file_content.content
                    ).decode("utf-8-sig")

                    # Aggiunge a capo se non presente alla fine del file esistente
                    if not contenuto_esistente.endswith("\n"):
                        contenuto_esistente += "\n"

                    nuovo_contenuto = contenuto_esistente + nuova_riga_csv

                    repo.update_file(
                        path=path_in_repo,
                        message=f"Nuova segnalazione manutenzione ({now_str})",
                        content=nuovo_contenuto,
                        sha=file_content.sha,
                    )
                except GithubException as e:
                    if e.status == 404:
                        # Se il file non esiste ancora su GitHub, scrive intestazione + prima riga
                        output_init = io.StringIO()
                        writer_init = csv.DictWriter(
                            output_init,
                            fieldnames=nuova_risposta.keys(),
                            delimiter=";",
                        )
                        writer_init.writeheader()
                        writer_init.writerow(nuova_risposta)
                        nuovo_contenuto = output_init.getvalue()

                        repo.create_file(
                            path=path_in_repo,
                            message=f"Creazione manutenzione.csv e prima segnalazione ({now_str})",
                            content=nuovo_contenuto,
                        )
                    else:
                        raise e

                st.session_state["ultima_segnalazione_manutenzione"] = (
                    nuova_risposta
                )
                st.success(
                    "Segnalazione acquisita e salvata con successo su GitHub!"
                )
                st.rerun()

            except Exception as e:
                st.error(
                    f"Si è verificato un errore durante il salvataggio su GitHub: {e}"
                )

    # ---------------------------------------------------------
    # 4. DOWNLOAD PDF DELLA RISPOSTA COMPILATA
    # ---------------------------------------------------------
    if "ultima_segnalazione_manutenzione" in st.session_state:
        st.markdown("---")
        st.subheader("Scarica la tua Segnalazione")

        dati_pdf = st.session_state["ultima_segnalazione_manutenzione"]

        # Generazione PDF dinamico ReportLab
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            "TitleStyle",
            parent=styles["Heading1"],
            fontSize=16,
            spaceAfter=12,
        )

        elements = []
        elements.append(
            Paragraph(
                "MODULO SEGNALAZIONE NEAR MISS / MANUTENZIONE", title_style
            )
        )
        elements.append(Spacer(1, 10))

        for k, v in dati_pdf.items():
            linea = f"<b>{k}:</b> {v if v else 'N/D'}"
            elements.append(Paragraph(linea, styles["Normal"]))
            elements.append(Spacer(1, 4))

        doc.build(elements)
        pdf_data = buffer.getvalue()
        buffer.close()

        st.download_button(
            label="Scarica Risposta Compilata in PDF",
            data=pdf_data,
            file_name=f"Segnalazione_Manutenzione_{dati_pdf['Data Ora Invio'].replace(':', '-').replace(' ', '_')}.pdf",
            mime="application/pdf",
            use_container_width=True,
        )
